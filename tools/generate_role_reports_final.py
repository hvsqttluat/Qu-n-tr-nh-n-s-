# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path.cwd()
OUTPUT_DIR = ROOT / "Bàn giao phần mềm quản lý nhân sự - Nhóm 3" / "Báo cáo chi tiết theo chức năng"


def set_font(
    run,
    name: str = "Times New Roman",
    size: int | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: int = 10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    for style_name in ["Normal", "Body Text", "List Bullet", "List Number"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(11)

    for style_name, size, color in [
        ("Title", 20, "172033"),
        ("Heading 1", 16, "172033"),
        ("Heading 2", 13, "1F2937"),
        ("Heading 3", 12, "374151"),
    ]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style.font.bold = True


def add_title(doc: Document, title: str, owner: str, scope: str):
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, size=20, bold=True, color="172033")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Người trình bày: {owner}")
    set_font(r, size=13, bold=True, color="0F766E")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(scope)
    set_font(r, size=11, color="4B5563")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tài liệu dùng để báo cáo, vừa nói vừa chỉ code trước hội đồng kiểm tra")
    set_font(r, size=10, italic=False, color="6B7280")


def h1(doc: Document, text: str):
    doc.add_paragraph(text, style="Heading 1")


def h2(doc: Document, text: str):
    doc.add_paragraph(text, style="Heading 2")


def p(doc: Document, text: str, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        r = paragraph.add_run(bold_prefix)
        set_font(r, bold=True)
        rest = text[len(bold_prefix) :]
        if rest:
            r = paragraph.add_run(rest)
            set_font(r)
    else:
        r = paragraph.add_run(text)
        set_font(r)
    return paragraph


def bullets(doc: Document, items: list[str]):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        r = paragraph.add_run(item)
        set_font(r)


def numbers(doc: Document, items: list[str]):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(2)
        r = paragraph.add_run(item)
        set_font(r)


def table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float] | None = None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        shade_cell(cell, "D9EAF7")
        set_cell_text(cell, header, bold=True, color="172033", size=10)
        if widths:
            cell.width = Inches(widths[i])

    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9)
            if widths:
                cells[i].width = Inches(widths[i])

    doc.add_paragraph()
    return tbl


def callout(doc: Document, title: str, lines: list[str], fill: str = "EEF6FF"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    first = cell.paragraphs[0]
    r = first.add_run(title)
    set_font(r, size=11, bold=True, color="172033")
    for line in lines:
        para = cell.add_paragraph(style="List Bullet")
        r = para.add_run(line)
        set_font(r, size=10)
    doc.add_paragraph()


def common_architecture(doc: Document):
    h1(doc, "1. Kiến trúc và công nghệ chung")
    table(
        doc,
        ["Thành phần", "Công nghệ / file", "Ý nghĩa khi trình bày"],
        [
            ("Giao diện", "WPF XAML trong MainWindow.xaml và LoginWindow.xaml", "XAML khai báo bố cục, màu, DataGrid, Button, Binding; code-behind chỉ xử lý thao tác giao diện thật cần thiết."),
            ("Ngôn ngữ", "C# trên .NET 10.0-windows, bật UseWPF trong QuanLyNhanSuWpf.csproj", "Ứng dụng desktop Windows, chạy dạng cửa sổ, phù hợp quản lý nội bộ."),
            ("Kết nối CSDL", "Microsoft.Data.SqlClient 5.2.2", "Tầng kho dữ liệu mở SqlConnection, dùng tham số @... để ghi đọc SQL Server."),
            ("Mẫu tổ chức", "MVVM theo hướng thực dụng: View = XAML, ViewModel = ManHinhChinhViewModel.cs, Model = MoHinh.cs, Repository = KhoDuLieuNhanSu.cs / KhoXacThuc.cs", "Dễ nói: người dùng bấm nút -> Command trong ViewModel -> gọi kho dữ liệu -> cập nhật ObservableCollection -> Binding tự đổi màn hình."),
            ("Danh sách động", "ObservableCollection và ICollectionView", "Danh sách nhân viên, đơn nghỉ, bảng lương có thể lọc/sắp xếp và tự cập nhật UI."),
            ("Lệnh giao diện", "ICommand thông qua LenhGiaoDien", "Nút trên màn hình không gọi SQL trực tiếp; nút gọi Command để kiểm tra quyền, validate dữ liệu, rồi mới ghi dữ liệu."),
        ],
        widths=[1.25, 2.45, 3.0],
    )


def add_practice_checklist(doc: Document, owner: str):
    h1(doc, "Checklist luyện nói trước khi báo cáo")
    bullets(
        doc,
        [
            f"{owner} mở sẵn Visual Studio hoặc VS Code tại đúng file được liệt kê trong bảng.",
            "Nói theo luồng: Màn hình -> Command/ViewModel -> Kho dữ liệu -> Bảng SQL -> Binding cập nhật lại giao diện.",
            "Không đọc từng dòng code; chỉ đọc tên hàm, điều kiện chính, câu SQL hoặc phép tính quan trọng.",
            "Khi thầy hỏi sang phần người khác, trả lời ranh giới rõ ràng: phần đó em có dùng dữ liệu, nhưng người phụ trách chính là bạn tương ứng.",
            "Luôn kết thúc mỗi chức năng bằng một câu kiểm chứng: dữ liệu đã lưu vào SQL, danh sách đã Refresh, hoặc thông báo đã cập nhật.",
        ],
    )


def save(doc: Document, filename: str) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def build_long() -> Path:
    doc = Document()
    configure_document(doc)
    add_title(
        doc,
        "Báo cáo phần Long - Đăng nhập, tài khoản, bảo mật, CSDL",
        "Long",
        "Phạm vi: đăng nhập, quản lý tài khoản, bảo mật mật khẩu, cấu trúc và kết nối cơ sở dữ liệu",
    )

    callout(
        doc,
        "Câu chốt phạm vi cho Long",
        [
            "Long phụ trách đường vào hệ thống: từ App.xaml mở LoginWindow, xác thực tài khoản qua SQL Server, tạo phiên đăng nhập, bảo vệ mật khẩu và chuẩn bị CSDL quản trị.",
            "Nếu thầy hỏi dashboard: dashboard không thuộc Long; dashboard là giao diện chính của Luật, còn Long đảm bảo người dùng vào được hệ thống và có dữ liệu SQL an toàn.",
        ],
        fill="EAF7F4",
    )

    common_architecture(doc)

    h1(doc, "2. File cần mở khi thầy hỏi")
    table(
        doc,
        ["Nội dung", "File / dòng chính", "Cách nói ngắn gọn"],
        [
            ("Ứng dụng bắt đầu từ đâu", "QuanLyNhanSuWpf\\App.xaml:4", "StartupUri trỏ vào LoginWindow.xaml, nghĩa là bắt buộc qua màn hình đăng nhập trước khi vào hệ thống."),
            ("Giao diện đăng nhập", "QuanLyNhanSuWpf\\LoginWindow.xaml:205", "Nút Đăng nhập gắn sự kiện DangNhap_Click; mật khẩu nhập bằng PasswordBox để không lộ ký tự."),
            ("Luồng bấm đăng nhập", "QuanLyNhanSuWpf\\LoginWindow.xaml.cs:22, :34, :53", "Click gọi DangNhap(), lấy tài khoản/mật khẩu, gọi KhoXacThuc.DangNhapAsync để xác thực."),
            ("Dịch vụ xác thực", "QuanLyNhanSuWpf\\KhoXacThuc.cs:13-69", "Kiểm tra SQL, tài khoản có tồn tại, trạng thái khóa, mật khẩu hợp lệ, ghi nhật ký và trả về phiên đăng nhập."),
            ("Chống nhập sai nhiều lần", "QuanLyNhanSuWpf\\KhoXacThuc.cs:38-63, :188", "Sai quá 5 lần sẽ tạm khóa, tránh dò mật khẩu thủ công."),
            ("Băm và kiểm tra mật khẩu", "QuanLyNhanSuWpf\\BaoMatMatKhau.cs:13, :25", "Mật khẩu không lưu dạng rõ; lưu hash, salt, số vòng lặp PBKDF2-SHA256."),
            ("Cấu hình kết nối SQL", "QuanLyNhanSuWpf\\CauHinhUngDung.cs:21-23, :41-49", "Ứng dụng thử các máy chủ SQL phổ biến như .\\SQLEXPRESS, localhost, LocalDB."),
            ("Tạo bảng quản trị", "QuanLyNhanSuWpf\\SoDoQuanTriSql.cs:10, :30", "Tự bảo đảm bảng HR_Users và HR_AuditLogs tồn tại trước khi đăng nhập."),
            ("Đồng bộ tài khoản theo nhân viên", "QuanLyNhanSuWpf\\KhoDuLieuNhanSu.cs:608, :620-636", "Có chức năng tạo tài khoản hệ thống cho hồ sơ nhân viên, dùng mật khẩu khởi tạo đã băm."),
        ],
        widths=[1.45, 2.3, 3.0],
    )

    h1(doc, "3. Luồng dữ liệu đăng nhập")
    numbers(
        doc,
        [
            "Người dùng mở app. WPF đọc App.xaml và hiển thị LoginWindow.xaml.",
            "Người dùng nhập tên đăng nhập và mật khẩu, bấm nút Đăng nhập.",
            "LoginWindow.xaml.cs gọi KhoXacThuc.DangNhapAsync(tenDangNhap, matKhau).",
            "KhoXacThuc lấy danh sách chuỗi kết nối từ CauHinhUngDung, mở SqlConnection đến SQL Server.",
            "SoDoQuanTriSql.DamBaoAsync kiểm tra bảng HR_Users, HR_AuditLogs; nếu thiếu thì tạo hoặc bổ sung cột cần thiết.",
            "KhoXacThuc đọc tài khoản trong HR_Users, kiểm tra IsActive, FailedLoginCount, LockoutUntilAt.",
            "BaoMatMatKhau.XacMinhMatKhau băm mật khẩu vừa nhập với salt cũ và so sánh hash.",
            "Nếu sai: tăng số lần sai, có thể tạm khóa, ghi HR_AuditLogs. Nếu đúng: reset số lần sai, cập nhật LastLoginAt, tạo PhienDangNhap.",
            "LoginWindow mở MainWindow và truyền PhienDangNhap vào ViewModel để phân quyền hiển thị/chức năng.",
        ],
    )

    h1(doc, "4. Giải thích từng phần chức năng")
    h2(doc, "4.1. Màn hình đăng nhập")
    bullets(
        doc,
        [
            "XAML chịu trách nhiệm bố cục: tiêu đề, ô tên đăng nhập, ô mật khẩu, nút đăng nhập và vùng thông báo lỗi.",
            "Code-behind không xử lý SQL trực tiếp; nó chỉ gom dữ liệu người dùng nhập, tắt bật trạng thái đang đăng nhập, rồi gọi kho xác thực.",
            "Khi đăng nhập thành công, cửa sổ chính được mở; khi thất bại, thông báo lỗi được hiện ngay tại màn hình.",
        ],
    )
    h2(doc, "4.2. Tài khoản và phân quyền")
    bullets(
        doc,
        [
            "Bảng HR_Users lưu Username, FullName, RoleName, PasswordHash, PasswordSalt, PasswordIterations, IsActive, RequirePasswordChange, FailedLoginCount, LockoutUntilAt, LastLoginAt.",
            "RoleName được đưa vào PhienDangNhap để MainWindow/ManHinhChinhViewModel biết người dùng đang ở vai trò nào.",
            "Chức năng đồng bộ tài khoản nhân sự tạo account theo hồ sơ nhân viên, giúp dữ liệu nhân sự và dữ liệu đăng nhập liên hệ được với nhau.",
        ],
    )
    h2(doc, "4.3. Bảo mật mật khẩu")
    bullets(
        doc,
        [
            "Không lưu mật khẩu gốc. Khi tạo hoặc reset mật khẩu, BaoMatMatKhau.BamMatKhau sinh salt ngẫu nhiên và hash bằng PBKDF2-SHA256.",
            "Khi đăng nhập, hệ thống không giải mã mật khẩu; nó băm mật khẩu người dùng vừa nhập với salt đang lưu rồi so sánh kết quả.",
            "Cơ chế khóa tạm thời sau nhiều lần sai giúp giảm rủi ro dò mật khẩu.",
        ],
    )
    h2(doc, "4.4. Cơ sở dữ liệu")
    bullets(
        doc,
        [
            "SoDoQuanTriSql chịu trách nhiệm bảng bảo mật: HR_Users, HR_AuditLogs.",
            "KhoDuLieuNhanSu chịu trách nhiệm bảng nghiệp vụ: HR_Departments, HR_JobPositions, HR_Employees, HR_Applicants, HR_LeaveRequests, HR_Attendances, HR_Appraisals, HR_Payslips, HR_Contracts.",
            "Cách nói quan trọng: CSDL không phải chỉ là nơi lưu dữ liệu, mà còn là điểm bảo đảm tính nhất quán giữa các module.",
        ],
    )

    h1(doc, "5. Đoạn code cần nhớ và ý nghĩa")
    table(
        doc,
        ["Đoạn code / hàm", "Ý nghĩa", "Câu trình bày mượt"],
        [
            ("StartupUri=\"LoginWindow.xaml\"", "Ép luồng sử dụng đi qua đăng nhập.", "Em đặt cửa sổ đăng nhập là màn hình đầu tiên để người dùng không vào thẳng dữ liệu nhân sự."),
            ("DangNhap_Click -> DangNhap()", "Tách sự kiện click khỏi xử lý bất đồng bộ.", "Click chỉ là điểm kích hoạt; logic thật nằm trong hàm async DangNhap để không làm treo giao diện."),
            ("await khoXacThuc.DangNhapAsync(...)", "Gọi tầng xác thực chuyên trách.", "Em không để UI tự kiểm tra SQL mà chuyển sang KhoXacThuc để dễ bảo trì và kiểm thử."),
            ("BaoMatMatKhau.XacMinhMatKhau(...)", "So sánh hash, salt, iterations.", "Hệ thống không giải mã mật khẩu; chỉ so sánh hash được tạo lại từ mật khẩu người dùng nhập."),
            ("TangSoLanSaiAsync / LockoutUntilAt", "Ghi nhận số lần sai và khóa tạm.", "Nếu nhập sai nhiều lần, tài khoản bị khóa tạm thời để bảo vệ khỏi dò mật khẩu."),
            ("GhiNhatKyAsync(... HR_AuditLogs ...)", "Lưu dấu vết đăng nhập.", "Mọi lần đăng nhập thành công, thất bại, bị khóa đều có nhật ký để quản trị kiểm tra."),
            ("SqlConnectionStringBuilder", "Chuẩn hóa chuỗi kết nối.", "Em dùng builder để tránh nối chuỗi thủ công và bảo đảm chuỗi kết nối hợp lệ."),
        ],
        widths=[2.05, 2.0, 2.7],
    )

    h1(doc, "6. Kịch bản vừa nói vừa chỉ")
    numbers(
        doc,
        [
            "Mở App.xaml và nói: Đây là điểm khởi động, em cho app mở LoginWindow trước để chặn truy cập chưa xác thực.",
            "Mở LoginWindow.xaml, chỉ nút Đăng nhập: Nút này gọi DangNhap_Click, phần giao diện chỉ thu thập thông tin.",
            "Mở LoginWindow.xaml.cs, chỉ hàm DangNhap(): Ở đây em lấy username/password, gọi KhoXacThuc, rồi xử lý thành công hoặc thất bại.",
            "Mở KhoXacThuc.cs, chỉ DangNhapAsync: Hàm này đi theo thứ tự kết nối SQL, bảo đảm schema, đọc tài khoản, kiểm tra khóa, xác minh mật khẩu.",
            "Mở BaoMatMatKhau.cs: Phần bảo mật nằm ở đây, dùng hash + salt, không lưu mật khẩu gốc.",
            "Mở SoDoQuanTriSql.cs: Đây là chỗ tạo HR_Users và HR_AuditLogs nếu CSDL chưa có.",
            "Kết lại: Như vậy Long chịu trách nhiệm cổng vào hệ thống và nền CSDL bảo mật; các bạn còn lại dùng phiên đăng nhập và dữ liệu này để vận hành module.",
        ],
    )

    h1(doc, "7. Câu hỏi thầy có thể hỏi và cách trả lời")
    table(
        doc,
        ["Câu hỏi", "Trả lời nên nói"],
        [
            ("Vì sao không lưu mật khẩu trực tiếp?", "Vì nếu CSDL bị lộ thì mật khẩu gốc không bị lộ. Hệ thống chỉ lưu hash, salt và số vòng lặp; khi đăng nhập thì băm lại mật khẩu nhập vào để so sánh."),
            ("Hash khác mã hóa ở điểm nào?", "Mã hóa có thể giải mã về bản gốc nếu có khóa; hash là một chiều. Với mật khẩu, em dùng hash một chiều kèm salt để an toàn hơn."),
            ("Nếu SQL Server chưa mở thì sao?", "KhoXacThuc thử nhiều chuỗi kết nối. Nếu không kết nối được thì trả thông báo rõ ràng; có cơ chế dự phòng cục bộ theo cấu hình, nhưng khi báo cáo chính em ưu tiên SQL Server."),
            ("Tại sao cần HR_AuditLogs?", "Để lưu dấu vết các sự kiện bảo mật như đăng nhập thành công, sai mật khẩu, khóa tài khoản, reset mật khẩu. Khi có sự cố có thể truy lại."),
            ("Nhập sai bao nhiêu lần thì khóa?", "Trong KhoXacThuc có hằng số SoLanSaiToiDa = 5 và thời gian khóa tạm thời 5 giờ."),
            ("CSDL được tạo bằng tay hay tự tạo?", "Ứng dụng có code bảo đảm schema. SoDoQuanTriSql tạo bảng quản trị, KhoDuLieuNhanSu tạo bảng nghiệp vụ, nên khi triển khai sẽ giảm lỗi thiếu bảng/cột."),
            ("Tài khoản nhân viên lấy từ đâu?", "Có chức năng đồng bộ theo hồ sơ nhân viên. KhoDuLieuNhanSu tạo HR_Users theo mã nhân viên hoặc thông tin nhân sự, mật khẩu khởi tạo vẫn được băm."),
            ("Dashboard có thuộc Long không?", "Không. Dashboard là giao diện chính của Luật. Long chỉ đảm bảo đăng nhập, phiên người dùng, bảo mật và CSDL để dashboard có dữ liệu hợp lệ."),
        ],
        widths=[2.5, 4.1],
    )

    h1(doc, "8. Nếu thầy yêu cầu sửa nhanh")
    table(
        doc,
        ["Yêu cầu sửa", "Mở file", "Cách xử lý"],
        [
            ("Đổi số lần khóa tài khoản", "KhoXacThuc.cs", "Sửa hằng số SoLanSaiToiDa, sau đó build/test lại."),
            ("Đổi thời gian khóa", "KhoXacThuc.cs", "Sửa ThoiGianKhoaTamThoi, ví dụ TimeSpan.FromHours(1)."),
            ("Đổi mật khẩu khởi tạo", "CauHinhUngDung.cs / cấu hình app", "Sửa giá trị mật khẩu khởi tạo hoặc biến cấu hình, không sửa trực tiếp trong SQL."),
            ("Thêm cột nhật ký", "SoDoQuanTriSql.cs", "Thêm ALTER TABLE nếu cột thiếu để nâng cấp CSDL an toàn."),
            ("Thêm vai trò người dùng", "HR_Users + ViewModel phân quyền", "Thêm RoleName mới và bổ sung điều kiện quyền ở ViewModel."),
        ],
        widths=[1.8, 2.0, 3.0],
    )

    add_practice_checklist(doc, "Long")
    return save(doc, "01_Long_Dang_nhap_tai_khoan_bao_mat_CSDL.docx")


def build_luat() -> Path:
    doc = Document()
    configure_document(doc)
    add_title(
        doc,
        "Báo cáo phần Luật - Giao diện chính, dashboard, nhân viên, phòng ban, tuyển dụng",
        "Luật",
        "Phạm vi: màn hình chính, dashboard/tổng quan, hồ sơ nhân viên, cơ cấu phòng ban, tuyển dụng",
    )

    callout(
        doc,
        "Câu trả lời ngay về dashboard",
        [
            "Đúng, dashboard/Tổng quan là phần của Luật vì nó nằm trong Giao diện chính và được khai báo trong MainWindow.xaml.",
            "Cách nói đẹp: Dashboard do Luật phụ trách phần hiển thị và tổng hợp trên giao diện; số liệu như nghỉ phép, lương, thông báo có nguồn từ các module nghiệp vụ của Tuyến.",
        ],
        fill="EAF7F4",
    )

    common_architecture(doc)

    h1(doc, "2. File cần mở khi thầy hỏi")
    table(
        doc,
        ["Nội dung", "File / dòng chính", "Cách nói ngắn gọn"],
        [
            ("Menu điều hướng chính", "QuanLyNhanSuWpf\\MainWindow.xaml:152-171", "Sidebar gồm Tổng quan, Tuyển dụng, Hồ sơ nhân viên, Cơ cấu phòng ban và các module nghiệp vụ khác."),
            ("Dashboard/Tổng quan", "QuanLyNhanSuWpf\\MainWindow.xaml:311-464", "Khai báo KPI, biểu đồ tròn, biểu đồ lương, biểu đồ ứng viên và trung tâm thông báo."),
            ("KPI tổng nhân viên", "MainWindow.xaml:322, :329, :336", "TextBlock Binding tới TongNhanVien, DangLamViec, TamNghi trong ViewModel."),
            ("Biểu đồ dashboard", "QuanLyNhanSuWpf\\BieuDoDashboard.cs:10-46, :142-153", "Custom control tự vẽ biểu đồ theo dữ liệu Binding, không cần thư viện ngoài."),
            ("Command chính", "ManHinhChinhViewModel.cs:97-118", "Các nút thêm/lưu/xóa nhân viên, lưu phòng ban, tạo ứng viên, chuyển ứng viên đều là ICommand."),
            ("Danh sách nhân viên", "MainWindow.xaml:604-654; ViewModel.cs:700-716", "DataGrid lấy ItemsSource từ DanhSachNhanVienView để lọc và sắp xếp."),
            ("Tạo/lưu/xóa nhân viên", "ViewModel.cs:1419, :1529, :1607", "ViewModel validate dữ liệu rồi gọi KhoDuLieuNhanSu để ghi SQL."),
            ("Phòng ban", "MainWindow.xaml:729-779; ViewModel.cs:1721, :1822, :3202-3219", "Lưu phòng ban, gán trưởng phòng, tổng hợp phòng ban điều hành."),
            ("Tuyển dụng", "MainWindow.xaml:803-821; ViewModel.cs:1645, :1859, :1920, :1954", "Tạo ứng viên, chuyển giai đoạn, tiếp nhận thành nhân viên và xuất hợp đồng."),
            ("Kho dữ liệu nhân sự", "KhoDuLieuNhanSu.cs:38, :50, :103, :166, :233", "Repository ghi/đọc HR_Employees, HR_Applicants, HR_Departments."),
        ],
        widths=[1.45, 2.45, 2.85],
    )

    h1(doc, "3. Luồng dữ liệu từ giao diện đến CSDL")
    numbers(
        doc,
        [
            "MainWindow.xaml khai báo giao diện: menu trái, vùng nội dung, DataGrid, form nhập liệu và nút thao tác.",
            "DataContext của MainWindow là ManHinhChinhViewModel, nên các Binding như TongNhanVien, DanhSachNhanVienView, BieuMauNhanVien lấy dữ liệu từ ViewModel.",
            "Người dùng bấm nút, Button.Command gọi ICommand tương ứng, ví dụ LuuLenh hoặc TaoUngVienLenh.",
            "ViewModel kiểm tra quyền, kiểm tra dữ liệu bắt buộc, cập nhật model tạm hoặc gọi KhoDuLieuNhanSu nếu đang dùng SQL.",
            "KhoDuLieuNhanSu thực thi câu lệnh INSERT/UPDATE/DELETE bằng SqlCommand và tham số @... để tránh lỗi nối chuỗi.",
            "Sau khi ghi thành công, ViewModel cập nhật ObservableCollection, gọi Refresh hoặc BaoThayDoi; WPF Binding tự cập nhật lại màn hình.",
            "Dashboard lấy dữ liệu tổng hợp từ các collection hiện có và các thuộc tính tính toán trong ViewModel, ví dụ TongNhanVien, DangLamViec, TamNghi, DuLieuLuong12Thang.",
        ],
    )

    h1(doc, "4. Giải thích từng phần chức năng")
    h2(doc, "4.1. Giao diện chính và điều hướng")
    bullets(
        doc,
        [
            "MainWindow.xaml dùng bố cục WPF: sidebar điều hướng bên trái, vùng nội dung chính bên phải, các Section hiện/ẩn theo mục đang chọn.",
            "Các nút điều hướng được Binding tới trạng thái mục hiện tại trong ViewModel, giúp người dùng đổi màn hình mà không cần mở nhiều cửa sổ rời.",
            "Thiết kế giao diện theo hướng quản trị: nhiều bảng, form nhập nhanh, số liệu tổng hợp và nút thao tác rõ ràng.",
        ],
    )
    h2(doc, "4.2. Dashboard/Tổng quan")
    bullets(
        doc,
        [
            "Dashboard là phần đầu của MainWindow, gồm KPI nhân sự, biểu đồ cơ cấu, biểu đồ lương 12 tháng, biểu đồ ứng viên và thông báo.",
            "Các KPI như TongNhanVien, DangLamViec, TamNghi là property tính toán trong ViewModel; không nhập tay trên giao diện.",
            "Biểu đồ được tách sang BieuDoDashboard.cs để giao diện gọn hơn và có thể tái sử dụng.",
            "Khi trình bày, Luật nên nói rõ: em phụ trách mặt hiển thị và tổng hợp dashboard; dữ liệu chi tiết từ chấm công/lương/thông báo là phần phối hợp với Tuyến.",
        ],
    )
    h2(doc, "4.3. Hồ sơ nhân viên")
    bullets(
        doc,
        [
            "DataGrid hiển thị DanhSachNhanVienView, có lọc theo từ khóa/phòng ban và sắp xếp theo thứ tự chức vụ, phòng ban, họ tên.",
            "Tạo mới nhân viên gọi TaoMoiNhanVien để reset form; lưu gọi LuuNhanVien để validate và ghi SQL; xóa gọi XoaNhanVien và cập nhật lại danh sách.",
            "KhoDuLieuNhanSu.LuuNhanVienAsync quyết định INSERT nếu MaNhanVien = 0, ngược lại UPDATE theo MaNhanVien.",
        ],
    )
    h2(doc, "4.4. Phòng ban")
    bullets(
        doc,
        [
            "Phòng ban có form lưu thông tin và chức năng gán trưởng phòng.",
            "ViewModel tổng hợp dữ liệu phòng ban thành TongHopPhongBanDieuHanh để dashboard/quản trị nhìn được số nhân sự, trưởng phòng và tình trạng từng phòng.",
            "Khi xóa hoặc cập nhật nhân viên, ViewModel có bước cập nhật lại phòng ban để tránh trưởng phòng hoặc quân số bị sai.",
        ],
    )
    h2(doc, "4.5. Tuyển dụng")
    bullets(
        doc,
        [
            "Ứng viên được quản lý riêng trong HR_Applicants, có giai đoạn tuyển dụng và vị trí ứng tuyển.",
            "Khi ứng viên đạt yêu cầu, ChuyenUngVienThanhNhanVien chuyển dữ liệu từ ứng viên sang hồ sơ nhân viên, đồng thời cập nhật trạng thái ứng viên.",
            "XuatHopDongLamViec tạo tài liệu hợp đồng từ dữ liệu ứng viên/nhân viên, hỗ trợ đầu ra cho quy trình tuyển dụng.",
        ],
    )

    h1(doc, "5. Đoạn code cần nhớ và ý nghĩa")
    table(
        doc,
        ["Đoạn code / hàm", "Ý nghĩa", "Câu trình bày mượt"],
        [
            ("TextBlock Text=\"{Binding TongNhanVien}\"", "KPI lấy dữ liệu từ ViewModel.", "Số trên dashboard không cố định trong XAML; nó được Binding từ dữ liệu thật."),
            ("DanhSachNhanVienView = CollectionViewSource.GetDefaultView(...)", "Tạo view lọc/sắp xếp cho DataGrid.", "Em dùng ICollectionView để lọc/sắp xếp mà không làm thay đổi danh sách gốc."),
            ("ThemMoiLenh / LuuLenh / XoaLenh", "Nút thao tác đi qua ICommand.", "Giao diện không gọi hàm lung tung; mỗi nút đi qua một command có điều kiện được phép chạy."),
            ("await khoDuLieu.LuuNhanVienAsync(...)", "Ghi nhân viên xuống SQL.", "ViewModel chỉ điều phối; phần SQL nằm trong KhoDuLieuNhanSu để tách trách nhiệm."),
            ("LuuPhongBanLenh / GanTruongPhongLenh", "Quản lý cơ cấu tổ chức.", "Phòng ban không chỉ là tên phòng, còn có trưởng phòng và liên kết với nhân viên."),
            ("ChuyenUngVienThanhNhanVien()", "Kết nối tuyển dụng với nhân sự.", "Ứng viên đạt yêu cầu được chuyển thành nhân viên, tránh nhập lại dữ liệu."),
            ("BaoThayDoi(nameof(...))", "Thông báo WPF cập nhật Binding.", "Sau khi dữ liệu đổi, ViewModel báo cho giao diện biết để số liệu và bảng tự làm mới."),
        ],
        widths=[2.2, 2.0, 2.55],
    )

    h1(doc, "6. Kịch bản vừa nói vừa chỉ")
    numbers(
        doc,
        [
            "Mở MainWindow.xaml, chỉ sidebar: Đây là giao diện chính, em phụ trách bố cục điều hướng và các màn hình quản lý nhân sự.",
            "Chỉ mục Tổng quan: Phần dashboard thuộc em vì nó nằm trong MainWindow; nó tổng hợp dữ liệu từ toàn hệ thống.",
            "Chỉ KPI TongNhanVien/DangLamViec/TamNghi rồi mở ManHinhChinhViewModel.cs: Các số này Binding tới property tính toán, không phải text tĩnh.",
            "Mở BieuDoDashboard.cs: Các biểu đồ dashboard được tách thành custom control để XAML không quá dài.",
            "Chỉ màn Hồ sơ nhân viên: DataGrid dùng DanhSachNhanVienView để lọc và sắp xếp. Bấm Lưu thì LuuLenh gọi LuuNhanVien.",
            "Mở KhoDuLieuNhanSu.cs dòng LuuNhanVienAsync: Ở đây mới có SQL INSERT/UPDATE, tức là giao diện và dữ liệu được tách tầng.",
            "Chỉ màn Phòng ban: Lưu phòng ban và gán trưởng phòng dùng command riêng, dữ liệu tổng hợp phòng ban được cập nhật lại sau khi thao tác.",
            "Chỉ màn Tuyển dụng: Ứng viên có thể chuyển giai đoạn, tiếp nhận thành nhân viên và xuất hợp đồng.",
            "Kết lại: Phần của Luật là trải nghiệm chính và các module quản lý hồ sơ/cơ cấu/tuyển dụng; dashboard là mặt tổng hợp, phối hợp dữ liệu với các phần khác.",
        ],
    )

    h1(doc, "7. Câu hỏi thầy có thể hỏi và cách trả lời")
    table(
        doc,
        ["Câu hỏi", "Trả lời nên nói"],
        [
            ("Dashboard là của ai?", "Dashboard là của Luật vì thuộc giao diện chính MainWindow. Tuy nhiên dữ liệu trên dashboard được tổng hợp từ nhiều module, trong đó có chấm công, nghỉ phép, lương, thông báo của Tuyến."),
            ("Vì sao dùng Binding thay vì set text trực tiếp?", "Binding giúp giao diện tự cập nhật khi dữ liệu trong ViewModel thay đổi; code sạch hơn và đúng hướng MVVM."),
            ("ICollectionView dùng để làm gì?", "Dùng để lọc, sắp xếp danh sách hiển thị trên DataGrid mà không phá danh sách gốc ObservableCollection."),
            ("Khi lưu nhân viên, dữ liệu đi đâu?", "Button gọi LuuLenh, ViewModel validate và gọi KhoDuLieuNhanSu.LuuNhanVienAsync, sau đó SQL ghi vào HR_Employees và ViewModel refresh danh sách."),
            ("Tại sao tách BieuDoDashboard.cs?", "Biểu đồ cần logic vẽ riêng. Nếu để hết trong XAML hoặc ViewModel sẽ rối, nên tách thành custom control."),
            ("Tuyển dụng liên kết với nhân viên thế nào?", "Ứng viên ở HR_Applicants. Khi đạt yêu cầu, ChuyenUngVienThanhNhanVien tạo hồ sơ HR_Employees và cập nhật trạng thái ứng viên."),
            ("Nếu thêm trường CCCD cho nhân viên thì sửa ở đâu?", "Sửa model NhanVien/BieuMauNhanVien, thêm control trong MainWindow.xaml, cập nhật LuuNhanVienAsync và câu đọc dữ liệu trong KhoDuLieuNhanSu."),
            ("Nếu lọc nhân viên bị sai thì kiểm tra đâu?", "Kiểm tra LocNhanVien và CauHinhDanhSachNhanVienView trong ManHinhChinhViewModel, sau đó kiểm tra Binding từ ô tìm kiếm/phòng ban."),
        ],
        widths=[2.45, 4.25],
    )

    h1(doc, "8. Nếu thầy yêu cầu sửa nhanh")
    table(
        doc,
        ["Yêu cầu sửa", "Mở file", "Cách xử lý"],
        [
            ("Đổi màu hoặc bố cục dashboard", "MainWindow.xaml / App.xaml resources", "Sửa Style, Brush, Grid/Border liên quan, không động vào logic SQL."),
            ("Thêm chỉ số KPI", "MainWindow.xaml + ManHinhChinhViewModel.cs", "Thêm TextBlock Binding và thêm property tính toán trong ViewModel."),
            ("Thêm cột DataGrid nhân viên", "MainWindow.xaml + MoHinh.cs + KhoDuLieuNhanSu.cs", "Thêm property, thêm cột Binding, cập nhật SQL đọc/ghi."),
            ("Thêm trạng thái tuyển dụng", "ManHinhChinhViewModel.cs", "Thêm giá trị trạng thái và điều kiện chuyển giai đoạn trong ChuyenGiaiDoanUngVien."),
            ("Đổi cách sắp xếp nhân viên", "ManHinhChinhViewModel.cs:700-716", "Sửa SortDescriptions theo trường cần ưu tiên."),
        ],
        widths=[1.8, 2.2, 2.8],
    )

    add_practice_checklist(doc, "Luật")
    return save(doc, "02_Luat_Giao_dien_chinh_nhan_vien_phong_ban_tuyen_dung.docx")


def build_tuyen() -> Path:
    doc = Document()
    configure_document(doc)
    add_title(
        doc,
        "Báo cáo phần Tuyến - Chấm công, nghỉ phép, đánh giá, lương, báo cáo, thông báo",
        "Tuyến",
        "Phạm vi: nghiệp vụ vận hành nhân sự sau khi đã có hồ sơ nhân viên",
    )

    callout(
        doc,
        "Câu chốt phạm vi cho Tuyến",
        [
            "Tuyến phụ trách các nghiệp vụ phát sinh hằng ngày: vào/ra ca, nghỉ phép, đánh giá, bảng lương, xuất báo cáo và thông báo.",
            "Dashboard không phải phần chính của Tuyến; Tuyến cung cấp dữ liệu nghiệp vụ để dashboard của Luật tổng hợp và hiển thị.",
        ],
        fill="EAF7F4",
    )

    common_architecture(doc)

    h1(doc, "2. File cần mở khi thầy hỏi")
    table(
        doc,
        ["Nội dung", "File / dòng chính", "Cách nói ngắn gọn"],
        [
            ("Command nghiệp vụ", "ManHinhChinhViewModel.cs:119-138", "Các nút chấm công, nghỉ phép, đánh giá, lương, báo cáo được khai báo thành ICommand."),
            ("Chấm công vào/ra ca", "ViewModel.cs:1987, :2016; KhoDuLieuNhanSu.cs:273, :287", "ViewModel nhận thao tác, KhoDuLieuNhanSu ghi HR_Attendances."),
            ("Điều chỉnh công", "KhoDuLieuNhanSu.cs:309", "Dùng khi cần sửa dữ liệu giờ vào/ra hoặc trạng thái công."),
            ("Nghỉ phép", "ViewModel.cs:2099, :2155; KhoDuLieuNhanSu.cs:327, :352", "Tạo đơn nghỉ, duyệt/từ chối đơn và cập nhật dữ liệu liên quan."),
            ("Tính số ngày nghỉ", "QuyTacNghiepVuNhanSu.cs:18", "Hàm tính số ngày bao gồm cả ngày bắt đầu và ngày kết thúc."),
            ("Đánh giá", "ViewModel.cs:2282, :2335, :2355; KhoDuLieuNhanSu.cs:374, :396, :444, :461", "Tạo, lưu, chốt hoặc xóa phiếu đánh giá nhân viên."),
            ("Bảng lương", "ViewModel.cs:2381, :2409, :2457, :2483; KhoDuLieuNhanSu.cs:481, :547", "Tạo phiếu lương, tính khấu trừ, xác nhận trả lương."),
            ("Quy tắc lương", "QuyTacNghiepVuNhanSu.cs:61, :82-89", "Tính bảo hiểm, khấu trừ nghỉ phép và tổng thực lĩnh."),
            ("Báo cáo", "ViewModel.cs:3815, :3866, :3919, :3967, :4044; BoXuatOffice.cs:17", "Tạo nội dung báo cáo và xuất ra DOCX/XLSX/PPTX/PDF/TXT tùy phần mở rộng."),
            ("Thông báo", "ViewModel.cs:2794, :3309, :3325; KhoDuLieuNhanSu.cs:1940-1978", "Thông báo được tạo từ dữ liệu nghiệp vụ và có thao tác đánh dấu đã đọc/gửi mới."),
        ],
        widths=[1.5, 2.45, 2.8],
    )

    h1(doc, "3. Luồng dữ liệu nghiệp vụ")
    numbers(
        doc,
        [
            "Người dùng chọn nhân viên hoặc bản ghi nghiệp vụ trên MainWindow.",
            "Button trên XAML gọi Command trong ManHinhChinhViewModel, ví dụ GhiNhanVaoCaLenh, ThemNghiPhepLenh, TaoPhieuLuongLenh.",
            "ViewModel kiểm tra điều kiện: đã chọn nhân viên chưa, dữ liệu ngày tháng có hợp lệ không, trạng thái hiện tại có được phép chuyển không.",
            "Nếu đang dùng SQL, ViewModel gọi KhoDuLieuNhanSu để ghi bảng tương ứng: HR_Attendances, HR_LeaveRequests, HR_Appraisals, HR_Payslips.",
            "Các công thức nghiệp vụ như số ngày nghỉ, bảo hiểm, lương thực lĩnh được đặt trong QuyTacNghiepVuNhanSu để không lẫn vào giao diện.",
            "Sau khi ghi, ViewModel cập nhật ObservableCollection và gọi BaoThayDoi/Refresh để DataGrid, dashboard và báo cáo nhận số liệu mới.",
            "Khi xuất báo cáo, ViewModel tạo TaiLieuXuat, BoXuatOffice chọn hàm xuất theo đuôi file và tạo tài liệu cho hội đồng/người dùng.",
        ],
    )

    h1(doc, "4. Giải thích từng phần chức năng")
    h2(doc, "4.1. Chấm công")
    bullets(
        doc,
        [
            "Chấm công bắt đầu từ nhân viên đang chọn. GhiNhanVaoCa tạo bản ghi vào ca; GhiNhanRaCa cập nhật giờ ra cho ca đang mở.",
            "Dữ liệu được ghi xuống HR_Attendances, gồm nhân viên, ngày, giờ vào, giờ ra và trạng thái.",
            "Khi trình bày cần nhấn mạnh: ViewModel chỉ điều phối; phương thức SQL nằm trong KhoDuLieuNhanSu để tránh giao diện phụ thuộc trực tiếp CSDL.",
        ],
    )
    h2(doc, "4.2. Nghỉ phép")
    bullets(
        doc,
        [
            "Người dùng tạo đơn nghỉ với nhân viên, ngày bắt đầu, ngày kết thúc, loại nghỉ và lý do.",
            "QuyTacNghiepVuNhanSu.TinhSoNgayBaoGom tính tổng số ngày để tránh sai lệch khi nghỉ nhiều ngày.",
            "Khi đơn được duyệt/từ chối, ViewModel cập nhật trạng thái và có thể tạo lại phiếu lương để phản ánh ảnh hưởng của nghỉ phép.",
        ],
    )
    h2(doc, "4.3. Đánh giá")
    bullets(
        doc,
        [
            "Phiếu đánh giá gồm nhân viên, kỳ đánh giá, điểm/kết quả, nhận xét và trạng thái.",
            "Lưu đánh giá dùng LuuDanhGiaAsync; chốt đánh giá dùng ChotDanhGiaAsync để khóa kết quả nghiệp vụ.",
            "Câu nên nói: Đánh giá tách khỏi bảng nhân viên để một nhân viên có nhiều kỳ đánh giá khác nhau.",
        ],
    )
    h2(doc, "4.4. Bảng lương")
    bullets(
        doc,
        [
            "Tạo phiếu lương lấy nhân viên, tháng, lương cơ bản, phụ cấp, thưởng, nghỉ phép và khấu trừ.",
            "QuyTacNghiepVuNhanSu.TaoPhieuLuongThang gom công thức nghiệp vụ; ViewModel gọi hàm này để tránh tính rải rác nhiều nơi.",
            "TinhKhauTruBaoHiem tính phần bảo hiểm từ lương cơ bản; tổng thực lĩnh trừ khấu trừ và cộng các khoản thu nhập.",
            "XacNhanTraLuongAsync cập nhật trạng thái đã trả, là điểm kết thúc quy trình lương.",
        ],
    )
    h2(doc, "4.5. Báo cáo")
    bullets(
        doc,
        [
            "Báo cáo được tạo từ dữ liệu hiện có, không nhập tay. ViewModel có các hàm tạo báo cáo nhân sự, chấm công, nghỉ phép, lương.",
            "BoXuatOffice nhận TaiLieuXuat rồi chọn định dạng theo phần mở rộng file: DOCX, XLSX, PPTX, PDF hoặc TXT.",
            "Khi thầy hỏi, mở BoXuatOffice để chứng minh phần xuất file được tách thành bộ riêng.",
        ],
    )
    h2(doc, "4.6. Thông báo")
    bullets(
        doc,
        [
            "Thông báo hệ thống giúp nhắc các sự kiện như nhân sự nghỉ phép, cần xử lý hồ sơ hoặc thông tin vận hành.",
            "ViewModel có lệnh mở/gửi thông báo và đánh dấu đã đọc; số thông báo chưa đọc được dashboard của Luật hiển thị.",
            "Một phần thông báo được tạo từ dữ liệu nghiệp vụ trong KhoDuLieuNhanSu.TaoThongBaoTuDuLieu.",
        ],
    )

    h1(doc, "5. Đoạn code cần nhớ và ý nghĩa")
    table(
        doc,
        ["Đoạn code / hàm", "Ý nghĩa", "Câu trình bày mượt"],
        [
            ("GhiNhanVaoCaAsync / GhiNhanRaCaAsync", "Ghi giờ vào/ra xuống HR_Attendances.", "Chấm công không chỉ đổi giao diện; nó lưu bản ghi thời gian để báo cáo và lương dùng lại."),
            ("ThemNghiPhepAsync", "Tạo đơn nghỉ phép.", "Đơn nghỉ được lưu riêng để có trạng thái chờ duyệt/đã duyệt/từ chối."),
            ("CapNhatTrangThaiNghiPhepAsync", "Duyệt hoặc từ chối đơn.", "Khi trạng thái nghỉ phép đổi, các phần thống kê và lương có thể cập nhật theo."),
            ("TinhSoNgayBaoGom", "Tính tổng ngày nghỉ.", "Em tách công thức ngày nghỉ ra quy tắc nghiệp vụ để dùng thống nhất ở nhiều nơi."),
            ("LuuDanhGiaAsync / ChotDanhGiaAsync", "Lưu và khóa phiếu đánh giá.", "Chốt đánh giá giúp phân biệt bản nháp với kết quả chính thức."),
            ("TaoPhieuLuongThang", "Tính phiếu lương.", "Lương lấy từ nhiều yếu tố nên công thức được gom vào một hàm nghiệp vụ rõ ràng."),
            ("TinhKhauTruBaoHiem", "Tính khấu trừ bảo hiểm.", "Đây là phần khấu trừ có công thức riêng, không nên nhập tay trên giao diện."),
            ("BoXuatOffice.Xuat", "Xuất báo cáo theo định dạng.", "ViewModel chỉ tạo nội dung, bộ xuất file quyết định lưu DOCX/XLSX/PPTX/PDF/TXT."),
            ("GuiThongBaoMoi / DanhDauDaDoc", "Gửi và đánh dấu thông báo.", "Thông báo là cầu nối giữa dữ liệu nghiệp vụ và trải nghiệm người dùng trên dashboard."),
        ],
        widths=[2.2, 2.0, 2.55],
    )

    h1(doc, "6. Kịch bản vừa nói vừa chỉ")
    numbers(
        doc,
        [
            "Mở ManHinhChinhViewModel.cs phần command: Đây là các lệnh nghiệp vụ em phụ trách, mỗi nút trên màn hình gọi một command rõ ràng.",
            "Chỉ GhiNhanVaoCa/GhiNhanRaCa: Khi chọn nhân viên và bấm vào/ra ca, ViewModel kiểm tra dữ liệu rồi gọi KhoDuLieuNhanSu.",
            "Mở KhoDuLieuNhanSu.cs dòng GhiNhanVaoCaAsync: Đây là chỗ ghi HR_Attendances, phục vụ báo cáo chấm công và tính lương.",
            "Chỉ phần nghỉ phép: Tạo đơn nghỉ dùng ThemNghiPhepAsync, duyệt đơn dùng CapNhatTrangThaiNghiPhepAsync.",
            "Mở QuyTacNghiepVuNhanSu.cs: Công thức ngày nghỉ và lương được tách riêng để dùng lại và dễ kiểm tra.",
            "Chỉ phần đánh giá: Phiếu đánh giá có trạng thái lưu/chốt, giúp phân biệt dữ liệu đang soạn với kết quả chính thức.",
            "Chỉ phần bảng lương: TaoPhieuLuongThang gom lương cơ bản, phụ cấp, thưởng, nghỉ phép, bảo hiểm để ra thực lĩnh.",
            "Mở BoXuatOffice.cs: Báo cáo không chỉ xem trên màn hình mà có thể xuất file theo định dạng người dùng chọn.",
            "Kết lại: Phần của Tuyến là chuỗi nghiệp vụ vận hành; dữ liệu này sau đó được dashboard và báo cáo tổng hợp lại.",
        ],
    )

    h1(doc, "7. Câu hỏi thầy có thể hỏi và cách trả lời")
    table(
        doc,
        ["Câu hỏi", "Trả lời nên nói"],
        [
            ("Nếu nhân viên quên chấm ra thì sao?", "Có dữ liệu vào ca nhưng chưa có giờ ra. Hệ thống có chức năng điều chỉnh công để quản trị sửa lại theo thực tế."),
            ("Nghỉ phép ảnh hưởng lương thế nào?", "Khi đơn nghỉ được duyệt, hệ thống có thể tạo/cập nhật phiếu lương tháng; công thức khấu trừ nghỉ phép nằm trong QuyTacNghiepVuNhanSu."),
            ("Tại sao tách QuyTacNghiepVuNhanSu?", "Vì công thức nghiệp vụ cần dùng ở nhiều nơi. Tách riêng giúp kiểm tra, sửa công thức và test dễ hơn."),
            ("Bảng lương tính thực lĩnh ra sao?", "Từ lương cơ bản, phụ cấp, thưởng, trừ bảo hiểm và khấu trừ nghỉ phép. Hàm TaoPhieuLuongThang gom logic này."),
            ("Đánh giá có ghi đè hồ sơ nhân viên không?", "Không. Đánh giá nằm ở bảng riêng HR_Appraisals, nên một nhân viên có nhiều kỳ đánh giá."),
            ("Báo cáo lấy dữ liệu từ đâu?", "Từ các collection đã được nạp từ SQL và các hàm tạo báo cáo trong ViewModel. BoXuatOffice chỉ phụ trách xuất file."),
            ("Thông báo lưu ở đâu?", "Trong bản hiện tại thông báo được tạo từ dữ liệu nghiệp vụ và quản lý trong collection ThongBao; dashboard dùng số thông báo chưa đọc để hiển thị."),
            ("Dashboard có phải phần Tuyến không?", "Không phải phần chính. Tuyến cung cấp dữ liệu nghiệp vụ, còn dashboard hiển thị/tổng hợp là phần của Luật."),
        ],
        widths=[2.45, 4.25],
    )

    h1(doc, "8. Nếu thầy yêu cầu sửa nhanh")
    table(
        doc,
        ["Yêu cầu sửa", "Mở file", "Cách xử lý"],
        [
            ("Đổi công thức tính lương", "QuyTacNghiepVuNhanSu.cs", "Sửa TaoPhieuLuongThang hoặc TinhKhauTruBaoHiem, sau đó chạy test nghiệp vụ."),
            ("Thêm loại nghỉ phép", "MoHinh.cs + MainWindow.xaml + ViewModel", "Thêm giá trị loại nghỉ, thêm lựa chọn UI, kiểm tra cách tính ngày/lương."),
            ("Thêm trạng thái đánh giá", "ManHinhChinhViewModel.cs + KhoDuLieuNhanSu.cs", "Bổ sung trạng thái và điều kiện chuyển trạng thái."),
            ("Thêm định dạng xuất báo cáo", "BoXuatOffice.cs", "Thêm nhánh theo phần mở rộng file mới và hàm xuất tương ứng."),
            ("Thêm thông báo mới", "KhoDuLieuNhanSu.cs:1940-1978 hoặc ViewModel.cs:3325", "Bổ sung điều kiện sinh thông báo hoặc form gửi thông báo."),
        ],
        widths=[1.8, 2.2, 2.8],
    )

    add_practice_checklist(doc, "Tuyến")
    return save(doc, "03_Tuyen_Cham_cong_nghi_phep_danh_gia_luong_bao_cao_thong_bao.docx")


def main():
    paths = [build_long(), build_luat(), build_tuyen()]
    for path in paths:
        print(str(path).encode("unicode_escape").decode("ascii"))


if __name__ == "__main__":
    main()

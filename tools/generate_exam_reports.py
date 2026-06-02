from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "BaoCaoPhanMem"
TODAY = date(2026, 6, 1).strftime("%d/%m/%Y")

PROJECT_NAME = "Phần mềm quản lý nhân sự WPF"
PROJECT_CODE = "QuanLyNhanSuWpf"
TEAM = "Nhóm 3"
COURSE = "Công nghệ phần mềm"
INSTRUCTOR = "Phan Nguyên Hải"
MEMBERS = [
    "Trần Văn Luật",
    "Nguyễn Đình Tuyến",
    "Trần Thanh Long",
]

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 96, 110)
BLACK = RGBColor(0, 0, 0)
TABLE_FILL = "F2F4F7"
ACCENT_FILL = "E8EEF5"


def set_run(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, TABLE_FILL)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run(run, 10, True, BLACK)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run(run, 10, False, BLACK)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run(run, 11, None, BLACK)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run(run, 11, None, BLACK)


def add_paragraph(doc: Document, text: str, bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_label and text.startswith(bold_label):
        r1 = p.add_run(bold_label)
        set_run(r1, 11, True, BLACK)
        r2 = p.add_run(text[len(bold_label):])
        set_run(r2, 11, False, BLACK)
    else:
        r = p.add_run(text)
        set_run(r, 11, False, BLACK)


def configure_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = BLUE
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = BLUE
    styles["Heading 3"].font.name = "Calibri"
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.color.rgb = DARK_BLUE

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.text = f"{PROJECT_CODE} | {title}"
    for run in header.runs:
        set_run(run, 9, False, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = f"{TEAM} - {COURSE}"
    for run in footer.runs:
        set_run(run, 9, False, MUTED)


def cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("TRƯỜNG/ĐƠN VỊ ĐÀO TẠO")
    set_run(r, 12, True, MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title.upper())
    set_run(r, 22, True, BLACK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(subtitle)
    set_run(r, 14, False, MUTED)

    rows = [
        ["Môn học", COURSE],
        ["Đề tài", PROJECT_NAME],
        ["Mã dự án", PROJECT_CODE],
        ["Nhóm thực hiện", TEAM],
        ["Giảng viên hướng dẫn", INSTRUCTOR],
        ["Ngày lập", TODAY],
    ]
    add_table(doc, ["Thông tin", "Nội dung"], rows, [4.0, 11.0])

    doc.add_heading("Thành viên nhóm", level=2)
    add_table(doc, ["STT", "Họ tên", "Vai trò gợi ý"], [
        ["1", MEMBERS[0], "Phân tích yêu cầu, tài liệu"],
        ["2", MEMBERS[1], "Thiết kế, lập trình WPF/CSDL"],
        ["3", MEMBERS[2], "Kiểm thử, đóng gói, thuyết trình"],
    ], [1.5, 8.5, 5.0])
    doc.add_page_break()


def new_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    configure_doc(doc, title)
    cover(doc, title, subtitle)
    return doc


def save_doc(doc: Document, filename: str) -> Path:
    path = OUTPUT / filename
    doc.save(path)
    return path


def doc_feasibility() -> Path:
    doc = new_doc(
        "Báo cáo nghiên cứu tính khả thi",
        "Đề tài: Phần mềm quản lý nhân sự WPF cho doanh nghiệp vừa và nhỏ",
    )
    doc.add_heading("1. Tóm tắt đề tài", level=1)
    add_paragraph(doc, "Dự án xây dựng một ứng dụng desktop chạy trên Windows để hỗ trợ doanh nghiệp quản lý hồ sơ nhân viên, phòng ban, tuyển dụng, chấm công, nghỉ phép, đánh giá, bảng lương, thông báo nội bộ và báo cáo nhân sự.")
    add_paragraph(doc, "Sản phẩm hiện có mã nguồn WPF, kết nối SQL Server khi có môi trường CSDL, có chế độ dữ liệu mẫu cục bộ để demo, có bộ kiểm thử tự động MSTest và script đóng gói bản win-x64.")

    doc.add_heading("2. Vấn đề cần giải quyết", level=1)
    add_bullets(doc, [
        "Thông tin nhân sự phân tán trên file Excel, giấy tờ hoặc nhiều hệ thống nhỏ lẻ.",
        "Quy trình nghỉ phép, chấm công và bảng lương khó theo dõi trạng thái xử lý.",
        "Lãnh đạo thiếu dashboard tổng quan để theo dõi biến động nhân sự và chi phí lương.",
        "Việc phân quyền, audit đăng nhập và sao lưu dữ liệu chưa được chuẩn hóa.",
    ])

    doc.add_heading("3. Mục tiêu và phạm vi", level=1)
    add_table(doc, ["Nhóm mục tiêu", "Nội dung"], [
        ["Nghiệp vụ", "Quản lý vòng đời nhân sự từ tuyển dụng, tiếp nhận, hồ sơ, chấm công, nghỉ phép, đánh giá đến bảng lương."],
        ["Kỹ thuật", "Xây dựng ứng dụng WPF .NET 10, lưu dữ liệu SQL Server, hỗ trợ cấu hình triển khai và fallback dữ liệu mẫu."],
        ["Bảo mật", "Xác thực bằng tài khoản SQL Server, băm mật khẩu PBKDF2-SHA256, phân quyền theo vai trò và ghi audit log."],
        ["Bàn giao", "Có mã nguồn, test tự động, gói zip self-contained win-x64, tài liệu cài đặt, sử dụng, thiết kế và test case."],
    ], [4.0, 11.0])

    doc.add_heading("4. Tính khả thi kỹ thuật", level=1)
    add_table(doc, ["Hạng mục", "Đánh giá", "Minh chứng trong project"], [
        ["Nền tảng", "Khả thi", "WPF trên .NET 10, target net10.0-windows, phù hợp ứng dụng desktop Windows."],
        ["CSDL", "Khả thi", "Microsoft.Data.SqlClient, tự tạo database HRManagementDB và các bảng HR_* khi kết nối được SQL Server."],
        ["Giao diện", "Khả thi", "XAML, DataGrid, style Segoe UI, bố cục sidebar và dashboard theo phong cách Fluent Design."],
        ["Bảo mật", "Khá", "PBKDF2-SHA256, salt riêng, role Admin/Giám đốc/Trưởng phòng/Nhân viên, audit log."],
        ["Kiểm thử", "Khá", "MSTest có 13 test pass cho bảo mật, cấu hình, phân quyền, nghiệp vụ lương/đánh giá và sắp xếp chức vụ."],
        ["Đóng gói", "Khả thi", "PowerShell publish self-contained win-x64, file Inno Setup tạo installer."],
    ], [3.0, 3.0, 9.0])

    doc.add_heading("5. Tính khả thi kinh tế", level=1)
    add_table(doc, ["Khoản mục", "Ước lượng", "Ghi chú"], [
        ["Nhân sự thực hiện", "3 thành viên", "Phù hợp yêu cầu nhóm 3-4 người."],
        ["Công cụ phát triển", "Không phát sinh lớn", "Visual Studio/VS Code, .NET SDK, SQL Server Express/LocalDB, Git."],
        ["Triển khai", "Thấp", "Chạy trên máy Windows, không cần server web nếu dùng SQL Server cục bộ."],
        ["Bảo trì", "Trung bình", "Cần duy trì schema SQL, nghiệp vụ lương/nghỉ phép, backup và phân quyền."],
    ], [4.0, 4.0, 7.0])

    doc.add_heading("6. Tính khả thi vận hành", level=1)
    add_bullets(doc, [
        "Người dùng chính gồm quản trị hệ thống, giám đốc, trưởng phòng và nhân viên.",
        "Các thao tác được gom theo phân hệ rõ ràng: tổng quan, tuyển dụng, hồ sơ, phòng ban, chấm công, nghỉ phép, đánh giá, lương, báo cáo, cài đặt.",
        "Ứng dụng có dữ liệu mẫu để trình diễn trong trường hợp máy thi chưa có SQL Server.",
        "Khi dùng thật cần cấu hình SQL Server và sao lưu định kỳ file backup hoặc database.",
    ])

    doc.add_heading("7. Rủi ro và phương án giảm thiểu", level=1)
    add_table(doc, ["Rủi ro", "Tác động", "Giải pháp"], [
        ["Máy chấm thi thiếu SQL Server", "Không đăng nhập được bằng CSDL thật", "Bật HRM_ALLOW_LOCAL_FALLBACK=true hoặc chuẩn bị sẵn SQL Server Express/LocalDB."],
        ["Thiếu minh chứng Figma", "Không lấy được bonus UI/UX", "Bổ sung link Figma hoặc ảnh export màn hình vào phụ lục."],
        ["Repo Git chưa có commit", "Không chứng minh được quản lý phiên bản", "Commit toàn bộ source, tài liệu và chụp git log/git status trước khi nộp."],
        ["Nghiệp vụ lương thực tế phức tạp", "Sai lệch so với doanh nghiệp thật", "Ghi rõ phạm vi demo, công thức lương cơ bản và hướng mở rộng."],
    ], [4.5, 4.0, 6.5])

    doc.add_heading("8. Kết luận", level=1)
    add_paragraph(doc, "Dự án khả thi để nộp và bảo vệ ở mức đồ án môn Công nghệ phần mềm. Sản phẩm có chương trình desktop chạy được, có CSDL SQL Server, có bảo mật cơ bản, có test tự động và có gói bàn giao. Các điểm cần hoàn thiện trước khi nộp gồm bổ sung tên thành viên, minh chứng Figma nếu muốn nhận bonus, commit Git và chạy kiểm thử thủ công theo file test case.")
    return save_doc(doc, "BaoCao_NghienCuuTinhKhaThi_QuanLyNhanSuWpf.docx")


def doc_plan() -> Path:
    doc = new_doc("Kế hoạch thực hiện dự án", "Kế hoạch phát triển phần mềm quản lý nhân sự WPF")
    doc.add_heading("1. Mục tiêu kế hoạch", level=1)
    add_paragraph(doc, "Kế hoạch này mô tả phạm vi công việc, phân công nhóm, tiến độ, mốc bàn giao, rủi ro và tiêu chí hoàn thành cho dự án phần mềm quản lý nhân sự WPF.")

    doc.add_heading("2. Phân công nhóm", level=1)
    add_table(doc, ["Vai trò", "Người phụ trách", "Nhiệm vụ chính"], [
        ["Trưởng nhóm/BA", MEMBERS[0], "Thu thập yêu cầu, lập tài liệu khả thi, đặc tả yêu cầu, theo dõi tiến độ."],
        ["Lập trình/Thiết kế", MEMBERS[1], "Thiết kế giao diện WPF, CSDL SQL Server, lập trình các phân hệ nghiệp vụ."],
        ["Kiểm thử/Tài liệu", MEMBERS[2], "Viết test case, chạy kiểm thử, lập hướng dẫn sử dụng/cài đặt, chuẩn bị demo."],
    ], [4.0, 5.0, 6.0])

    doc.add_heading("3. Cấu trúc công việc WBS", level=1)
    add_table(doc, ["Mã", "Công việc", "Sản phẩm đầu ra"], [
        ["WBS-01", "Khảo sát và chọn đề tài", "Đề cương đề tài, phạm vi chức năng."],
        ["WBS-02", "Nghiên cứu tính khả thi", "Báo cáo khả thi."],
        ["WBS-03", "Lập kế hoạch dự án", "Kế hoạch, phân công, lịch trình."],
        ["WBS-04", "Đặc tả yêu cầu", "SRS, actor, use case, yêu cầu phi chức năng."],
        ["WBS-05", "Thiết kế phần mềm", "Kiến trúc WPF, module, luồng xử lý."],
        ["WBS-06", "Thiết kế CSDL", "Schema SQL Server, bảng, quan hệ, dữ liệu mẫu."],
        ["WBS-07", "Lập trình", "Source code WPF, SQL, kiểm thử tự động."],
        ["WBS-08", "Kiểm thử và sửa lỗi", "Test case, biên bản test, build pass."],
        ["WBS-09", "Đóng gói và bàn giao", "Zip win-x64, tài liệu sử dụng/cài đặt, báo cáo Git."],
    ], [2.0, 5.0, 8.0])

    doc.add_heading("4. Tiến độ dự kiến", level=1)
    add_table(doc, ["Tuần", "Nội dung", "Kết quả"], [
        ["1", "Chọn đề tài, khảo sát quy trình nhân sự, lập phạm vi", "Đề tài được duyệt, danh sách chức năng sơ bộ."],
        ["2", "Viết khả thi, kế hoạch, đặc tả yêu cầu", "Bộ tài liệu phân tích yêu cầu."],
        ["3", "Thiết kế UI, CSDL, kiến trúc chương trình", "Tài liệu thiết kế phần mềm và CSDL."],
        ["4", "Lập trình WPF, kết nối SQL Server, phân quyền", "Ứng dụng chạy được các phân hệ chính."],
        ["5", "Kiểm thử, đóng gói, hoàn thiện báo cáo", "Build Release, test pass, tài liệu nộp."],
    ], [2.0, 7.0, 6.0])

    doc.add_heading("5. Công cụ và môi trường", level=1)
    add_table(doc, ["Nhóm công cụ", "Công cụ sử dụng"], [
        ["IDE/ngôn ngữ", "C#, XAML, .NET SDK 10, Visual Studio hoặc VS Code."],
        ["Giao diện", "Windows Presentation Foundation (WPF), custom style theo hướng Fluent Design."],
        ["CSDL", "SQL Server LocalDB/Express, Microsoft.Data.SqlClient."],
        ["Kiểm thử", "MSTest 4.0.2, dotnet test."],
        ["Quản lý phiên bản", "Git, GitHub Actions workflow desktop-ci."],
        ["Đóng gói", "dotnet publish self-contained win-x64, PowerShell, Inno Setup."],
    ], [4.0, 11.0])

    doc.add_heading("6. Tiêu chí hoàn thành", level=1)
    add_bullets(doc, [
        "Ứng dụng build Release không lỗi và chạy được trên Windows.",
        "Đăng nhập được bằng tài khoản mặc định khi có SQL Server hoặc chế độ fallback khi demo.",
        "Các phân hệ chính có dữ liệu, thao tác thêm/sửa/xóa hoặc xử lý trạng thái phù hợp.",
        "Bộ tài liệu bắt buộc theo yêu cầu môn học được hoàn thiện.",
        "Test tự động pass và test case thủ công được cập nhật kết quả trước khi nộp.",
        "Repo Git có commit rõ ràng, có thể chứng minh quá trình làm việc.",
    ])
    return save_doc(doc, "KeHoach_ThucHienDuAn_QuanLyNhanSuWpf.docx")


def doc_srs() -> Path:
    doc = new_doc("Tài liệu đặc tả yêu cầu", "Software Requirements Specification cho phần mềm quản lý nhân sự WPF")
    doc.add_heading("1. Giới thiệu", level=1)
    add_paragraph(doc, "Tài liệu đặc tả yêu cầu mô tả các chức năng, actor, dữ liệu, ràng buộc và tiêu chí chấp nhận của phần mềm quản lý nhân sự WPF. Đây là căn cứ để thiết kế, lập trình, kiểm thử và nghiệm thu sản phẩm.")

    doc.add_heading("2. Phạm vi hệ thống", level=1)
    add_bullets(doc, [
        "Quản lý hồ sơ nhân viên, phòng ban, vị trí công việc và cơ cấu tổ chức.",
        "Quản lý tuyển dụng từ ứng viên đến tiếp nhận thành nhân viên.",
        "Ghi nhận chấm công, nghỉ phép, đánh giá năng lực và bảng lương.",
        "Cung cấp dashboard, báo cáo, thông báo nội bộ và quản trị tài khoản.",
        "Không bao gồm tích hợp ngân hàng, bảo hiểm xã hội, email/SMS thật hoặc ứng dụng mobile.",
    ])

    doc.add_heading("3. Actor và quyền", level=1)
    add_table(doc, ["Actor", "Mô tả", "Phạm vi quyền"], [
        ["Admin", "Quản trị hệ thống", "Toàn quyền tài khoản, dữ liệu, nghiệp vụ, báo cáo và cấu hình."],
        ["Giám đốc", "Người điều hành", "Xem toàn hệ thống, xử lý phòng ban, nhân sự, nghỉ phép, lương, đánh giá, báo cáo."],
        ["Trưởng phòng", "Quản lý đội nhóm", "Xem và xử lý nhân sự trong phạm vi phòng ban; trưởng phòng nhân sự được xử lý bảng lương."],
        ["Nhân viên", "Người dùng tự phục vụ", "Xem hồ sơ cá nhân, chấm công, tạo nghỉ phép, xem phiếu lương, thông báo và đánh giá liên quan."],
    ], [3.0, 4.0, 8.0])

    doc.add_heading("4. Yêu cầu chức năng", level=1)
    add_table(doc, ["Mã", "Tên yêu cầu", "Mô tả", "Ưu tiên"], [
        ["FR-01", "Đăng nhập", "Người dùng đăng nhập bằng tài khoản lưu trong SQL Server, mật khẩu băm PBKDF2-SHA256.", "Cao"],
        ["FR-02", "Phân quyền", "Ẩn/khóa chức năng theo vai trò Admin, Giám đốc, Trưởng phòng, Nhân viên.", "Cao"],
        ["FR-03", "Dashboard", "Hiển thị tổng nhân viên, trạng thái làm việc, thông báo, biểu đồ nhân sự và lương.", "Cao"],
        ["FR-04", "Hồ sơ nhân viên", "Thêm, sửa, xóa, tìm kiếm, lọc nhân viên; quản lý CCCD, liên hệ khẩn cấp, tài khoản ngân hàng.", "Cao"],
        ["FR-05", "Phòng ban", "Tạo/sửa/xóa phòng ban, gán trưởng phòng, thống kê nhân sự theo phòng ban.", "Cao"],
        ["FR-06", "Tuyển dụng", "Tạo ứng viên, chuyển giai đoạn, tiếp nhận ứng viên thành nhân viên, xuất hợp đồng làm việc.", "Cao"],
        ["FR-07", "Chấm công", "Vào ca, ra ca, tính số giờ làm, cho phép người có quyền điều chỉnh công.", "Cao"],
        ["FR-08", "Nghỉ phép", "Tạo đơn nghỉ, tính số ngày, duyệt hoặc từ chối đơn nghỉ phép.", "Cao"],
        ["FR-09", "Đánh giá", "Ghi nhận đánh giá theo kỳ quý, điểm, nhận xét và trạng thái chốt.", "Trung bình"],
        ["FR-10", "Bảng lương", "Tính lương theo lương cơ bản, giờ công, phụ cấp, khấu trừ nghỉ phép và xác nhận trả lương.", "Cao"],
        ["FR-11", "Thông báo", "Tạo, lọc, đánh dấu đã đọc, gắn tệp và hiển thị thông báo theo phân hệ.", "Trung bình"],
        ["FR-12", "Báo cáo", "Xuất báo cáo nhân viên, chấm công, nghỉ phép, bảng lương sang DOCX/XLSX/PPTX/TXT.", "Trung bình"],
        ["FR-13", "Tài khoản", "Tạo tài khoản nhân viên, khóa/mở tài khoản, đặt lại mật khẩu.", "Cao"],
        ["FR-14", "Sao lưu", "Xuất và phục hồi dữ liệu bằng file .hrmbackup.json.", "Trung bình"],
        ["FR-15", "Audit log", "Ghi nhật ký đăng nhập và thao tác tài khoản quan trọng.", "Trung bình"],
    ], [1.6, 3.3, 8.0, 2.0])

    doc.add_heading("5. Yêu cầu phi chức năng", level=1)
    add_table(doc, ["Mã", "Nhóm", "Yêu cầu"], [
        ["NFR-01", "Hiệu năng", "Màn hình danh sách dùng DataGrid có ảo hóa dòng/cột, thao tác phản hồi trong vài giây với dữ liệu demo/mức vừa."],
        ["NFR-02", "Bảo mật", "Không lưu mật khẩu dạng rõ; dùng salt riêng, PBKDF2-SHA256, kiểm tra thời gian hằng định."],
        ["NFR-03", "Khả dụng", "Có fallback dữ liệu mẫu để demo khi SQL Server chưa sẵn sàng, nhưng triển khai thật ưu tiên SQL Server."],
        ["NFR-04", "Dễ dùng", "Giao diện tiếng Việt, menu theo phân hệ, nút thao tác rõ ràng, dashboard trực quan."],
        ["NFR-05", "Bảo trì", "Tách lớp model, viewmodel, kho dữ liệu, xác thực, cấu hình và tiện ích xuất Office."],
        ["NFR-06", "Triển khai", "Chạy trên Windows 10/11, có script publish self-contained win-x64 và file Inno Setup."],
    ], [1.8, 3.0, 10.2])

    doc.add_heading("6. Use case chính", level=1)
    add_table(doc, ["Use case", "Actor", "Luồng thành công"], [
        ["UC-01 Đăng nhập", "Tất cả", "Nhập tên đăng nhập/mật khẩu, hệ thống xác thực, mở màn hình chính theo vai trò."],
        ["UC-02 Tiếp nhận ứng viên", "Admin/Giám đốc/Trưởng phòng", "Tạo ứng viên, cập nhật giai đoạn, chuyển thành nhân viên khi ký hợp đồng."],
        ["UC-03 Chấm công", "Tất cả", "Chọn nhân viên trong phạm vi, vào ca, ra ca, hệ thống tính số giờ làm."],
        ["UC-04 Duyệt nghỉ phép", "Admin/Giám đốc/Trưởng phòng", "Xem đơn chờ duyệt, duyệt/từ chối, dashboard và thông báo cập nhật."],
        ["UC-05 Tính lương", "Admin/Giám đốc/TP nhân sự", "Chọn nhân viên, tính lương kỳ tháng, xem phiếu lương, xác nhận trả lương."],
        ["UC-06 Xuất báo cáo", "Admin/Giám đốc/Trưởng phòng", "Lọc dữ liệu, chọn loại báo cáo, lưu file DOCX/XLSX/PPTX/TXT."],
    ], [4.0, 4.0, 7.0])
    return save_doc(doc, "TaiLieu_DacTaYeuCau_QuanLyNhanSuWpf.docx")


def doc_design() -> Path:
    doc = new_doc("Tài liệu thiết kế phần mềm", "Thiết kế kiến trúc và module cho ứng dụng WPF")
    doc.add_heading("1. Tổng quan kiến trúc", level=1)
    add_paragraph(doc, "Ứng dụng được thiết kế theo hướng MVVM đơn giản: lớp giao diện WPF XAML hiển thị và binding dữ liệu, ViewModel điều phối lệnh và trạng thái màn hình, lớp kho dữ liệu làm việc với SQL Server hoặc dữ liệu mẫu cục bộ.")
    add_table(doc, ["Lớp", "Thành phần", "Vai trò"], [
        ["Presentation", "LoginWindow.xaml, MainWindow.xaml, App.xaml", "Định nghĩa giao diện, style, navigation, DataGrid, form nhập liệu và dashboard."],
        ["ViewModel", "ManHinhChinhViewModel.cs, LenhGiaoDien", "Quản lý state, command, lọc dữ liệu, phân quyền, gọi nghiệp vụ."],
        ["Domain Model", "MoHinh.cs, QuyTacNghiepVuNhanSu.cs", "Định nghĩa đối tượng nghiệp vụ và quy tắc lương, kỳ đánh giá, mã nhân viên."],
        ["Data Access", "KhoDuLieuNhanSu.cs, KhoXacThuc.cs, SoDoQuanTriSql.cs", "Kết nối SQL Server, tạo schema, CRUD nghiệp vụ, xác thực và audit."],
        ["Infrastructure", "CauHinhUngDung.cs, BaoMatMatKhau.cs, BoXuatOffice.cs", "Cấu hình, bảo mật mật khẩu, xuất tài liệu Office."],
        ["Testing/Release", "QuanLyNhanSuWpf.Tests, tools/package-release.ps1, installer/*.iss", "Kiểm thử tự động, publish và đóng gói cài đặt."],
    ], [3.2, 5.0, 6.8])

    doc.add_heading("2. Sơ đồ luồng tổng quát", level=1)
    add_numbered(doc, [
        "Người dùng mở ứng dụng, hệ thống hiển thị LoginWindow.",
        "KhoXacThuc đọc cấu hình kết nối, đảm bảo database và bảng quản trị, xác thực tài khoản.",
        "Sau khi đăng nhập, MainWindow nhận PhienDangNhap và khởi tạo ManHinhChinhViewModel.",
        "ViewModel tải dữ liệu qua KhoDuLieuNhanSu, ưu tiên SQL Server, fallback dữ liệu mẫu nếu được cấu hình.",
        "Người dùng thao tác theo phân hệ, command trong ViewModel gọi lớp kho dữ liệu để cập nhật CSDL.",
        "Dashboard, thông báo, biểu đồ và báo cáo được làm mới sau mỗi thao tác nghiệp vụ.",
    ])

    doc.add_heading("3. Thiết kế module", level=1)
    add_table(doc, ["Module", "Màn hình/chức năng", "Ghi chú thiết kế"], [
        ["Đăng nhập", "LoginWindow", "Xác thực SQL Server, thông báo lỗi rõ ràng, ghi audit."],
        ["Dashboard", "Tổng quan vận hành", "KPI nhân sự, biểu đồ trạng thái, biểu đồ lương, danh sách thông báo."],
        ["Nhân viên", "Hồ sơ nhân viên", "DataGrid kết hợp form, lọc theo từ khóa/phòng ban, CRUD có kiểm tra quyền."],
        ["Phòng ban", "Cơ cấu phòng ban", "Quản lý phòng ban, trưởng phòng, số nhân viên, quỹ lương."],
        ["Tuyển dụng", "Ứng viên", "Quản lý giai đoạn ứng viên, chuyển ứng viên thành nhân viên."],
        ["Chấm công", "Vào ca/Ra ca/Điều chỉnh", "Chống mở nhiều ca, tính số giờ theo CheckIn/CheckOut."],
        ["Nghỉ phép", "Tạo và duyệt đơn", "Tính số ngày theo khoảng ngày, trạng thái Pending/Approved/Rejected."],
        ["Đánh giá", "Kỳ đánh giá quý", "Tạo/chốt đánh giá, lưu điểm và nhận xét."],
        ["Bảng lương", "Tính lương/Xác nhận trả", "Tạo phiếu lương tháng, phụ cấp 5%, khấu trừ nghỉ phép đã duyệt."],
        ["Thông báo", "Tạo, lọc, đánh dấu đã đọc", "Thông báo nội bộ trong phiên, có phân hệ, mức độ và tệp đính kèm."],
        ["Báo cáo", "Xuất Office", "Xuất dữ liệu sang DOCX/XLSX/PPTX/TXT bằng helper nội bộ."],
        ["Cài đặt", "Tài khoản, backup/restore", "Quản trị tài khoản và sao lưu file .hrmbackup.json."],
    ], [3.0, 4.0, 8.0])

    doc.add_heading("4. Thiết kế giao diện", level=1)
    add_bullets(doc, [
        "Phong cách giao diện: WPF desktop hiện đại, màu xanh chủ đạo, font Segoe UI, card thống kê, sidebar phân hệ.",
        "Màn hình đăng nhập gồm phần nhận diện hệ thống và form tài khoản/mật khẩu.",
        "Màn hình chính có sidebar, thanh thông tin phiên, hero dashboard, thông báo nhanh và vùng nội dung theo phân hệ.",
        "Các danh sách dùng DataGrid read-only, thao tác thông qua form và nút lệnh để giảm sai sót.",
        "Thiết kế lấy cảm hứng Fluent Design; chưa có minh chứng dùng thư viện Microsoft Fluent UI chính thức hoặc MAUI.",
    ])

    doc.add_heading("5. Thiết kế bảo mật", level=1)
    add_table(doc, ["Cơ chế", "Thiết kế"], [
        ["Mật khẩu", "PBKDF2-SHA256, salt 16 byte, hash 32 byte, 210.000 vòng lặp, so sánh FixedTimeEquals."],
        ["Tài khoản", "Bảng HR_Users lưu username, role, hash, salt, trạng thái khóa, số lần sai, lần đăng nhập cuối."],
        ["Phân quyền", "Các property CoQuyen* trong ViewModel quyết định menu, command và phạm vi dữ liệu."],
        ["Audit", "Bảng HR_AuditLogs ghi đăng nhập, tạo tài khoản, khóa/mở, reset mật khẩu và máy thực hiện."],
        ["Cấu hình", "Ưu tiên biến môi trường HRM_CONNECTION_STRING, HRM_INITIAL_PASSWORD, HRM_ALLOW_LOCAL_FALLBACK."],
    ], [3.0, 12.0])

    doc.add_heading("6. Thiết kế kiểm thử và triển khai", level=1)
    add_bullets(doc, [
        "Kiểm thử tự động dùng MSTest cho bảo mật, cấu hình, phân quyền, nghiệp vụ nhân sự và sắp xếp chức vụ.",
        "Build và test chạy bằng dotnet build/dotnet test trên solution QuanLyNhanSuWpf.sln.",
        "GitHub Actions workflow desktop-ci chạy restore, build, test, publish trên windows-latest.",
        "Script tools/package-release.ps1 tạo bản publish self-contained win-x64 và nén file bàn giao.",
    ])
    return save_doc(doc, "TaiLieu_ThietKePhanMem_QuanLyNhanSuWpf.docx")


def doc_db() -> Path:
    doc = new_doc("Tài liệu thiết kế cơ sở dữ liệu", "Thiết kế CSDL SQL Server HRManagementDB")
    doc.add_heading("1. Tổng quan", level=1)
    add_paragraph(doc, "Cơ sở dữ liệu chính của ứng dụng là HRManagementDB trên SQL Server. Khi kết nối thành công, ứng dụng tự tạo database và các bảng nghiệp vụ HR_* nếu chưa tồn tại.")

    doc.add_heading("2. Danh sách bảng", level=1)
    add_table(doc, ["Bảng", "Mục đích", "Khóa chính"], [
        ["HR_Departments", "Lưu phòng ban và trưởng phòng", "DepartmentID"],
        ["HR_JobPositions", "Lưu vị trí công việc thuộc phòng ban", "PositionID"],
        ["HR_Employees", "Lưu hồ sơ nhân viên", "EmployeeID"],
        ["HR_Applicants", "Lưu ứng viên tuyển dụng", "ApplicantID"],
        ["HR_LeaveRequests", "Lưu đơn nghỉ phép", "LeaveID"],
        ["HR_Attendances", "Lưu chấm công vào/ra ca", "AttendanceID"],
        ["HR_Appraisals", "Lưu đánh giá năng lực", "AppraisalID"],
        ["HR_Payslips", "Lưu phiếu lương", "PayslipID"],
        ["HR_Contracts", "Lưu hợp đồng và lương cơ bản", "ContractID"],
        ["HR_Users", "Lưu tài khoản đăng nhập và vai trò", "UserID"],
        ["HR_AuditLogs", "Lưu nhật ký hệ thống", "AuditID"],
    ], [4.0, 8.0, 3.0])

    doc.add_heading("3. Mô tả trường chính", level=1)
    table_rows = [
        ["HR_Departments", "DepartmentID, Name, ManagerID", "ManagerID trỏ đến nhân viên làm trưởng phòng."],
        ["HR_JobPositions", "PositionID, DepartmentID, Name, ExpectedSalary, Status", "Mỗi vị trí thuộc một phòng ban."],
        ["HR_Employees", "EmployeeID, EmployeeCode, ApplicantID, FullName, DepartmentID, PositionID, ManagerID, JoinDate, IsActive, EmergencyContact, BankAccount, IdentityNumber", "Bảng trung tâm của hệ thống nhân sự."],
        ["HR_Applicants", "ApplicantID, PositionID, FullName, Email, Phone, CVFile_Url, Stage", "Stage gồm các giai đoạn tuyển dụng như New, Interview, Offer, Signed."],
        ["HR_LeaveRequests", "LeaveID, EmployeeID, LeaveType, StartDate, EndDate, TotalDays, Status, ApproverID, Reason", "Status gồm Pending, Approved, Rejected."],
        ["HR_Attendances", "AttendanceID, EmployeeID, CheckInTime, CheckOutTime, WorkHours", "WorkHours tính theo khoảng vào/ra ca."],
        ["HR_Appraisals", "AppraisalID, EmployeeID, ReviewerID, ReviewPeriod, Score, Feedback, Status", "ReviewPeriod theo quý, ví dụ 2026-Q2."],
        ["HR_Payslips", "PayslipID, EmployeeID, PayPeriod, BasicSalary, WorkDays, TotalAllowances, TotalDeductions, NetSalary, Status", "PayPeriod theo tháng, ví dụ 2026-05."],
        ["HR_Contracts", "ContractID, EmployeeID, ContractType, StartDate, EndDate, BasicSalary, Status", "Nguồn lấy lương cơ bản khi tính lương."],
        ["HR_Users", "UserID, Username, FullName, RoleName, PasswordHash, PasswordSalt, PasswordIterations, IsActive, RequirePasswordChange, FailedLoginCount, CreatedAt, LastLoginAt", "Không lưu mật khẩu dạng rõ."],
        ["HR_AuditLogs", "AuditID, ActorUsername, ActionName, EntityName, EntityKey, Detail, MachineName, CreatedAt", "Ghi vết hành động quan trọng."],
    ]
    add_table(doc, ["Bảng", "Trường chính", "Ghi chú"], table_rows, [3.5, 7.5, 4.0])

    doc.add_heading("4. Quan hệ dữ liệu", level=1)
    add_table(doc, ["Quan hệ", "Diễn giải"], [
        ["HR_Departments 1-n HR_JobPositions", "Một phòng ban có nhiều vị trí công việc."],
        ["HR_Departments 1-n HR_Employees", "Một phòng ban có nhiều nhân viên."],
        ["HR_JobPositions 1-n HR_Applicants", "Một vị trí có nhiều ứng viên."],
        ["HR_Applicants 0..1 - 0..1 HR_Employees", "Ứng viên có thể được tiếp nhận thành nhân viên."],
        ["HR_Employees 1-n HR_Attendances", "Một nhân viên có nhiều bản ghi chấm công."],
        ["HR_Employees 1-n HR_LeaveRequests", "Một nhân viên có nhiều đơn nghỉ phép."],
        ["HR_Employees 1-n HR_Appraisals", "Một nhân viên có nhiều đánh giá, reviewer cũng là nhân viên."],
        ["HR_Employees 1-n HR_Payslips", "Một nhân viên có nhiều phiếu lương theo kỳ."],
        ["HR_Employees 1-n HR_Contracts", "Một nhân viên có nhiều hợp đồng theo thời gian."],
    ], [6.0, 9.0])

    doc.add_heading("5. Ghi chú thiết kế", level=1)
    add_bullets(doc, [
        "Schema trong code hiện chủ yếu tạo bảng và cột, chưa khai báo đầy đủ FOREIGN KEY trong SQL. Ứng dụng đảm bảo quan hệ ở tầng nghiệp vụ.",
        "Thông báo nội bộ hiện nằm trong bộ nhớ ứng dụng, chưa có bảng HR_Notifications. Nếu triển khai thật nên bổ sung bảng thông báo để lưu lâu dài.",
        "Dữ liệu mật như PasswordHash và PasswordSalt cần được bảo vệ bằng phân quyền SQL Server và sao lưu an toàn.",
        "Cần backup database định kỳ và kiểm thử phục hồi trước khi dùng trong môi trường thật.",
    ])
    return save_doc(doc, "TaiLieu_ThietKeCSDL_QuanLyNhanSuWpf.docx")


def doc_install() -> Path:
    doc = new_doc("Tài liệu hướng dẫn cài đặt", "Hướng dẫn cài đặt và chạy phần mềm quản lý nhân sự WPF")
    doc.add_heading("1. Yêu cầu môi trường", level=1)
    add_table(doc, ["Thành phần", "Yêu cầu"], [
        ["Hệ điều hành", "Windows 10/11."],
        ["Runtime/build", ".NET SDK 10 nếu build từ source; bản publish self-contained không cần cài .NET runtime riêng."],
        ["CSDL", "SQL Server LocalDB, SQL Server Express hoặc SQL Server đầy đủ."],
        ["Quyền máy", "Tài khoản Windows có quyền tạo database trên SQL Server hoặc quyền dùng database đã cấp."],
        ["Công cụ tùy chọn", "Inno Setup nếu muốn tạo file cài đặt .exe."],
    ], [4.0, 11.0])

    doc.add_heading("2. Cấu hình", level=1)
    add_paragraph(doc, "Có thể cấu hình bằng biến môi trường hoặc file appsettings.json đặt cạnh file QuanLyNhanSuWpf.exe.")
    add_table(doc, ["Khóa", "Ý nghĩa", "Ví dụ"], [
        ["HRM_CONNECTION_STRING", "Chuỗi kết nối SQL Server chính", "Server=.\\SQLEXPRESS;Database=HRManagementDB;Trusted_Connection=True;TrustServerCertificate=True;"],
        ["HRM_INITIAL_PASSWORD", "Mật khẩu khởi tạo tài khoản mặc định", "Admin@2026!"],
        ["HRM_ALLOW_LOCAL_FALLBACK", "Cho phép đăng nhập demo khi không có SQL Server", "true hoặc false"],
    ], [4.0, 6.0, 5.0])

    doc.add_heading("3. Chạy từ source", level=1)
    add_numbered(doc, [
        "Mở PowerShell tại thư mục C:\\Users\\LUAT\\Documents\\CNPM nhóm 3.",
        "Chạy lệnh: dotnet restore .\\QuanLyNhanSuWpf\\QuanLyNhanSuWpf.sln.",
        "Chạy lệnh: dotnet build .\\QuanLyNhanSuWpf\\QuanLyNhanSuWpf.sln -c Release.",
        "Chạy lệnh: dotnet test .\\QuanLyNhanSuWpf\\QuanLyNhanSuWpf.sln -c Release.",
        "Chạy ứng dụng: dotnet run --project .\\QuanLyNhanSuWpf\\QuanLyNhanSuWpf.csproj.",
    ])

    doc.add_heading("4. Tài khoản mặc định", level=1)
    add_table(doc, ["Tên đăng nhập", "Vai trò", "Mật khẩu"], [
        ["admin", "Admin", "Giá trị HRM_INITIAL_PASSWORD"],
        ["giamdoc", "Giám đốc", "Giá trị HRM_INITIAL_PASSWORD"],
        ["truongphong", "Trưởng phòng", "Giá trị HRM_INITIAL_PASSWORD"],
        ["nhanvien", "Nhân viên", "Giá trị HRM_INITIAL_PASSWORD"],
    ], [4.0, 4.0, 7.0])

    doc.add_heading("5. Đóng gói bàn giao", level=1)
    add_numbered(doc, [
        "Chạy PowerShell: powershell -ExecutionPolicy Bypass -File .\\tools\\package-release.ps1.",
        "Kiểm tra thư mục artifacts\\QuanLyNhanSuWpf chứa file chạy.",
        "Nộp file artifacts\\QuanLyNhanSuWpf-win-x64.zip nếu yêu cầu bản chạy portable.",
        "Nếu cần installer, biên dịch installer\\QuanLyNhanSuWpf.iss bằng Inno Setup sau khi publish.",
    ])

    doc.add_heading("6. Xử lý lỗi thường gặp", level=1)
    add_table(doc, ["Lỗi", "Nguyên nhân", "Cách xử lý"], [
        ["Không kết nối SQL Server", "SQL Server chưa cài/chưa chạy hoặc sai connection string", "Cài SQL Server Express/LocalDB, sửa HRM_CONNECTION_STRING hoặc bật fallback cho demo."],
        ["Đăng nhập thất bại", "Sai mật khẩu khởi tạo hoặc tài khoản bị khóa", "Kiểm tra HRM_INITIAL_PASSWORD, dùng admin reset hoặc tạo lại database demo."],
        ["Build lỗi thiếu SDK", "Máy chưa có .NET SDK 10", "Cài .NET SDK 10 hoặc dùng bản publish self-contained."],
        ["Không tạo được installer", "Chưa publish hoặc chưa cài Inno Setup", "Chạy package-release.ps1 trước, sau đó mở file .iss bằng Inno Setup."],
    ], [4.0, 5.5, 5.5])
    return save_doc(doc, "TaiLieu_HuongDanCaiDat_QuanLyNhanSuWpf.docx")


def doc_user() -> Path:
    doc = new_doc("Tài liệu hướng dẫn sử dụng", "Hướng dẫn thao tác cho phần mềm quản lý nhân sự WPF")
    doc.add_heading("1. Đăng nhập và giao diện chính", level=1)
    add_numbered(doc, [
        "Mở QuanLyNhanSuWpf.exe.",
        "Nhập tên đăng nhập và mật khẩu được cấp.",
        "Sau khi đăng nhập, kiểm tra vai trò và phạm vi dữ liệu ở thanh bên trái.",
        "Chọn phân hệ trên menu để thao tác: Tổng quan, Tuyển dụng, Hồ sơ, Phòng ban, Chấm công, Nghỉ phép, Đánh giá, Bảng lương, Báo cáo, Cài đặt.",
    ])

    doc.add_heading("2. Tổng quan vận hành", level=1)
    add_bullets(doc, [
        "Xem tổng số nhân viên, số đang làm, số tạm nghỉ và số thông báo chưa đọc.",
        "Theo dõi biểu đồ cơ cấu nhân sự và lương 12 tháng.",
        "Dùng vùng thông báo để đọc thông tin hệ thống, lọc thông báo chưa đọc hoặc tạo thông báo mới nếu có quyền.",
    ])

    doc.add_heading("3. Hồ sơ nhân viên và phòng ban", level=1)
    add_table(doc, ["Nghiệp vụ", "Cách thao tác"], [
        ["Thêm nhân viên", "Vào Hồ sơ nhân viên, bấm Tạo mới, nhập mã, họ tên, phòng ban, vị trí, ngày vào làm, thông tin cá nhân, bấm Lưu hồ sơ."],
        ["Sửa nhân viên", "Chọn dòng nhân viên, bấm Nạp dòng chọn, chỉnh form và bấm Lưu hồ sơ."],
        ["Xóa nhân viên", "Chọn nhân viên và bấm Xóa nếu tài khoản có quyền."],
        ["Tìm kiếm/lọc", "Nhập từ khóa hoặc chọn phòng ban để lọc DataGrid."],
        ["Quản lý phòng ban", "Vào Cơ cấu phòng ban để tạo/sửa/xóa phòng ban và gán trưởng phòng."],
    ], [4.0, 11.0])

    doc.add_heading("4. Tuyển dụng", level=1)
    add_numbered(doc, [
        "Chọn phân hệ Tuyển dụng.",
        "Nhập thông tin ứng viên gồm họ tên, email, điện thoại, vị trí ứng tuyển.",
        "Cập nhật giai đoạn tuyển dụng khi ứng viên đi qua từng bước.",
        "Khi ứng viên đạt yêu cầu, dùng chức năng chuyển thành nhân viên để tạo hồ sơ nhân sự.",
        "Có thể xuất hợp đồng làm việc theo mẫu dữ liệu hiện có.",
    ])

    doc.add_heading("5. Chấm công và nghỉ phép", level=1)
    add_table(doc, ["Chức năng", "Mô tả"], [
        ["Vào ca", "Ghi nhận thời điểm bắt đầu làm việc cho nhân viên trong phạm vi."],
        ["Ra ca", "Ghi nhận thời điểm kết thúc và tự tính số giờ làm."],
        ["Điều chỉnh công", "Người có quyền có thể đặt số giờ chuẩn cho bản ghi được chọn."],
        ["Tạo nghỉ phép", "Nhập nhân viên, loại nghỉ, từ ngày, đến ngày, lý do."],
        ["Duyệt/từ chối nghỉ phép", "Admin, Giám đốc, Trưởng phòng được xử lý đơn chờ duyệt."],
    ], [4.0, 11.0])

    doc.add_heading("6. Đánh giá và bảng lương", level=1)
    add_bullets(doc, [
        "Đánh giá theo kỳ quý, nhập điểm và nhận xét, sau đó chốt đánh giá.",
        "Bảng lương tính theo kỳ tháng, lương cơ bản từ hợp đồng, phụ cấp 5%, khấu trừ ngày nghỉ đã duyệt và số giờ công.",
        "Người có quyền có thể xem phiếu lương chi tiết và xác nhận đã trả lương.",
    ])

    doc.add_heading("7. Báo cáo, tài khoản và sao lưu", level=1)
    add_table(doc, ["Nhóm chức năng", "Hướng dẫn"], [
        ["Báo cáo", "Chọn loại báo cáo nhân viên/chấm công/nghỉ phép/lương, chọn nơi lưu file và định dạng."],
        ["Tài khoản", "Admin tạo tài khoản nhân viên, khóa/mở tài khoản hoặc đặt lại mật khẩu."],
        ["Sao lưu", "Dùng Xuất dữ liệu để tạo file .hrmbackup.json và Phục hồi dữ liệu để nhập lại khi cần."],
        ["Đăng xuất", "Bấm Đăng xuất ở góc phải vùng dashboard để quay lại màn hình đăng nhập."],
    ], [4.0, 11.0])
    return save_doc(doc, "TaiLieu_HuongDanSuDung_QuanLyNhanSuWpf.docx")


def doc_git_checklist() -> Path:
    doc = new_doc("Báo cáo Git và checklist yêu cầu", "Đối chiếu yêu cầu môn học với project QuanLyNhanSuWpf")
    doc.add_heading("1. Công nghệ sử dụng", level=1)
    add_table(doc, ["Nhóm", "Công nghệ"], [
        ["Loại phần mềm", "Ứng dụng desktop Windows, không phải web."],
        ["Ngôn ngữ", "C# 14/.NET 10, XAML."],
        ["Framework UI", "Windows Presentation Foundation (WPF), giao diện custom theo phong cách Fluent Design."],
        ["CSDL", "SQL Server HRManagementDB, Microsoft.Data.SqlClient 5.2.2."],
        ["Kiểm thử", "MSTest 4.0.2, dotnet test."],
        ["Đóng gói", "dotnet publish self-contained win-x64, PowerShell, Inno Setup."],
        ["Quản lý phiên bản", "Git local repo, GitHub Actions workflow desktop-ci. Hiện cần tạo commit trước khi nộp."],
    ], [4.0, 11.0])

    doc.add_heading("2. Checklist yêu cầu bắt buộc", level=1)
    add_table(doc, ["Yêu cầu", "Trạng thái", "Minh chứng"], [
        ["Lập nhóm 3-4 người", "Cần điền tên", "Trang bìa tài liệu đang để Nhóm 3 và chỗ điền thành viên."],
        ["Chọn đề tài không trùng", "Đạt", "Đề tài: Phần mềm quản lý nhân sự WPF."],
        ["Tài liệu nghiên cứu tính khả thi", "Đạt", "BaoCao_NghienCuuTinhKhaThi_QuanLyNhanSuWpf.docx."],
        ["Kế hoạch dự án", "Đạt", "KeHoach_ThucHienDuAn_QuanLyNhanSuWpf.docx."],
        ["Đặc tả yêu cầu", "Đạt", "TaiLieu_DacTaYeuCau_QuanLyNhanSuWpf.docx."],
        ["Tài liệu thiết kế", "Đạt", "TaiLieu_ThietKePhanMem_QuanLyNhanSuWpf.docx."],
        ["Tài liệu kiểm thử", "Đạt khung", "TaiLieu_TestCase_QuanLyNhanSuWpf.xlsx, cần cập nhật kết quả test thủ công trước khi nộp."],
        ["Lập chương trình", "Đạt", "Solution WPF build Release thành công, 13/13 test tự động pass."],
    ], [5.0, 3.0, 7.0])

    doc.add_heading("3. Checklist bonus", level=1)
    add_table(doc, ["Bonus", "Trạng thái", "Ghi chú"], [
        ["Sử dụng quản lý phiên bản", "Cần chốt", "Repo Git đã khởi tạo nhưng hiện chưa có commit. Cần git add/commit và chụp git log."],
        ["Figma UI/UX", "Chưa có minh chứng", "Chưa thấy link/file Figma trong project. Cần bổ sung nếu muốn lấy bonus."],
        ["Tài liệu thiết kế CSDL", "Đạt", "Có tài liệu CSDL và schema SQL Server trong code."],
        ["Hướng dẫn sử dụng/cài đặt", "Đạt", "Có 2 file hướng dẫn riêng."],
        ["Fluent UI hoặc MAUI", "Đạt một phần", "Giao diện WPF theo phong cách Fluent, nhưng không phải MAUI hoặc thư viện Fluent UI chính thức."],
        ["Notification", "Đạt mức demo", "Có module thông báo nội bộ trong WPF, tạo/lọc/đánh dấu đã đọc/đính kèm file."],
    ], [5.0, 3.0, 7.0])

    doc.add_heading("4. Lệnh kiểm tra đã chạy", level=1)
    add_table(doc, ["Lệnh", "Kết quả"], [
        ["dotnet build .\\QuanLyNhanSuWpf\\QuanLyNhanSuWpf.sln -c Release", "Build succeeded, 0 warning, 0 error."],
        ["dotnet test .\\QuanLyNhanSuWpf\\QuanLyNhanSuWpf.sln -c Release --no-build", "Passed 13/13 tests."],
        ["git status --short", "Toàn bộ project đang ở trạng thái untracked do repo chưa có commit."],
    ], [7.0, 8.0])

    doc.add_heading("5. Việc cần làm trước khi nộp", level=1)
    add_numbered(doc, [
        "Điền tên thật của thành viên và giảng viên hướng dẫn vào trang bìa các tài liệu.",
        "Chạy ứng dụng thủ công theo file test case, cập nhật trạng thái Pass/Fail và ghi chú lỗi nếu có.",
        "Tạo commit Git đầu tiên, chụp git log/git status hoặc đẩy GitHub nếu được yêu cầu.",
        "Bổ sung link hoặc ảnh Figma nếu nhóm có thiết kế UI/UX.",
        "Chuẩn bị máy demo có SQL Server hoặc bật fallback cục bộ để tránh lỗi môi trường khi bảo vệ.",
    ])
    return save_doc(doc, "BaoCao_GitVaChecklist_QuanLyNhanSuWpf.docx")


def make_test_cases() -> list[list[str]]:
    rows: list[list[str]] = []

    def add(tc_id, module, name, pre, steps, data, expected, kind="Manual", status="Cần chạy"):
        rows.append([tc_id, module, name, pre, steps, data, expected, kind, status])

    add("TC-AUTH-01", "Đăng nhập", "Đăng nhập admin thành công", "Có SQL Server hoặc fallback", "Mở app, nhập admin và mật khẩu khởi tạo, bấm Đăng nhập", "admin/Admin@2026!", "Mở màn hình chính, vai trò Admin, thấy đầy đủ menu")
    add("TC-AUTH-02", "Đăng nhập", "Sai mật khẩu", "Có tài khoản admin", "Nhập admin với mật khẩu sai", "admin/SaiMatKhau", "Hiện thông báo tên đăng nhập hoặc mật khẩu chưa đúng")
    add("TC-AUTH-03", "Đăng nhập", "Tài khoản nhân viên bị giới hạn quyền", "Đăng nhập nhân viên", "Đăng nhập nhanvien, quan sát menu", "nhanvien/Admin@2026!", "Không vào được Cài đặt tài khoản, không xử lý bảng lương")
    add("TC-AUTH-04", "Đăng nhập", "Audit đăng nhập", "SQL Server hoạt động", "Đăng nhập thành công/sai, kiểm tra HR_AuditLogs", "admin", "Có bản ghi LoginSuccess hoặc LoginFailed")

    add("TC-EMP-01", "Hồ sơ nhân viên", "Thêm nhân viên mới", "Vai trò có quyền quản lý hồ sơ", "Bấm Tạo mới, nhập dữ liệu, Lưu hồ sơ", "NV999, Nguyễn Văn Test", "Nhân viên xuất hiện trong danh sách")
    add("TC-EMP-02", "Hồ sơ nhân viên", "Sửa hồ sơ", "Có nhân viên trong danh sách", "Chọn nhân viên, Nạp dòng chọn, sửa phòng ban/chức vụ, Lưu", "NV999", "Dữ liệu cập nhật đúng")
    add("TC-EMP-03", "Hồ sơ nhân viên", "Tìm kiếm nhân viên", "Có dữ liệu", "Nhập từ khóa họ tên/mã/phòng ban", "NV001", "Danh sách lọc đúng nhân viên liên quan")
    add("TC-EMP-04", "Hồ sơ nhân viên", "Xóa nhân viên", "Có nhân viên test", "Chọn nhân viên test, bấm Xóa", "NV999", "Nhân viên bị xóa và dữ liệu liên quan được xử lý")

    add("TC-DEPT-01", "Phòng ban", "Tạo phòng ban", "Vai trò Admin/Giám đốc", "Vào Cơ cấu phòng ban, nhập tên, Lưu", "Phòng Kiểm thử", "Phòng ban mới xuất hiện")
    add("TC-DEPT-02", "Phòng ban", "Gán trưởng phòng", "Có phòng ban và nhân viên", "Chọn phòng ban, chọn trưởng phòng, bấm Gán trưởng phòng", "Phòng Nhân sự", "Trưởng phòng cập nhật đúng")
    add("TC-DEPT-03", "Phòng ban", "Không xóa phòng ban có nhân viên", "Phòng ban đang có nhân viên", "Chọn phòng ban có nhân viên, bấm Xóa", "Phòng Nhân sự", "Hệ thống cảnh báo không thể xóa")

    add("TC-REC-01", "Tuyển dụng", "Tạo ứng viên", "Có vị trí công việc", "Nhập họ tên, email, điện thoại, vị trí, bấm lưu", "Ứng viên A", "Ứng viên xuất hiện ở danh sách")
    add("TC-REC-02", "Tuyển dụng", "Chuyển giai đoạn ứng viên", "Có ứng viên", "Chọn ứng viên, chọn giai đoạn mới, cập nhật", "Interview/Offer", "Giai đoạn thay đổi đúng")
    add("TC-REC-03", "Tuyển dụng", "Tiếp nhận ứng viên thành nhân viên", "Ứng viên chưa được tiếp nhận", "Chọn ứng viên, bấm chuyển thành nhân viên", "Ứng viên A", "Có hồ sơ nhân viên mới, ứng viên chuyển Signed")
    add("TC-REC-04", "Tuyển dụng", "Xuất hợp đồng làm việc", "Có ứng viên được chọn", "Chọn ứng viên, bấm Xuất hợp đồng", "Ứng viên A", "File hợp đồng DOCX/TXT được tạo")

    add("TC-ATT-01", "Chấm công", "Vào ca", "Có nhân viên trong phạm vi", "Chọn nhân viên, bấm Vào ca", "NV001", "Có bản ghi CheckInTime, CheckOutTime rỗng")
    add("TC-ATT-02", "Chấm công", "Không cho vào ca khi chưa ra ca", "Nhân viên có ca mở", "Bấm Vào ca lần nữa", "NV001", "Hệ thống báo nhân viên đang có ca chưa ra")
    add("TC-ATT-03", "Chấm công", "Ra ca", "Nhân viên có ca mở", "Bấm Ra ca", "NV001", "Có CheckOutTime và WorkHours > 0")
    add("TC-ATT-04", "Chấm công", "Điều chỉnh công", "Vai trò có quyền", "Chọn dòng công, bấm Điều chỉnh công", "NV001", "WorkHours được đặt về 8")

    add("TC-LEAVE-01", "Nghỉ phép", "Tạo đơn nghỉ phép", "Có nhân viên", "Nhập loại nghỉ, từ ngày, đến ngày, lý do, bấm tạo", "Nghỉ phép năm", "Đơn nghỉ ở trạng thái Chờ duyệt")
    add("TC-LEAVE-02", "Nghỉ phép", "Duyệt đơn nghỉ", "Có đơn chờ duyệt", "Chọn đơn, bấm Duyệt", "Đơn nghỉ", "Trạng thái chuyển Đã duyệt")
    add("TC-LEAVE-03", "Nghỉ phép", "Từ chối đơn nghỉ", "Có đơn chờ duyệt", "Chọn đơn, bấm Từ chối", "Đơn nghỉ", "Trạng thái chuyển Từ chối")
    add("TC-LEAVE-04", "Nghỉ phép", "Tính ngày nghỉ qua tháng", "Có đơn nghỉ giao giữa 2 tháng", "Tạo đơn từ cuối tháng sang đầu tháng", "29/04-02/05", "Số ngày tính đúng theo khoảng ngày")

    add("TC-REV-01", "Đánh giá", "Tạo đánh giá quý", "Vai trò có quyền", "Chọn nhân viên, nhập điểm, nhận xét, tạo đánh giá", "2026-Q2", "Đánh giá xuất hiện trong danh sách")
    add("TC-REV-02", "Đánh giá", "Chốt đánh giá", "Có đánh giá nháp", "Chọn đánh giá, bấm Chốt", "Đánh giá", "Trạng thái chuyển hoàn thành/chốt")
    add("TC-REV-03", "Đánh giá", "Kỳ đánh giá tự tính theo quý", "Có màn hình đánh giá", "Tạo đánh giá trong quý 2", "30/05/2026", "Kỳ hiển thị là 2026-Q2")

    add("TC-PAY-01", "Bảng lương", "Tính lương tháng", "Có nhân viên và hợp đồng", "Chọn nhân viên, bấm Tính lương", "NV001", "Tạo phiếu lương kỳ yyyy-MM")
    add("TC-PAY-02", "Bảng lương", "Khấu trừ nghỉ phép đã duyệt", "Có ngày nghỉ đã duyệt", "Tính lương sau khi duyệt nghỉ", "2 ngày nghỉ", "Khấu trừ đúng theo lương cơ bản/ngày công chuẩn")
    add("TC-PAY-03", "Bảng lương", "Xem phiếu lương", "Có phiếu lương", "Chọn phiếu, bấm Xem", "Phiếu lương", "Hiển thị chi tiết lương, phụ cấp, khấu trừ, thực lãnh")
    add("TC-PAY-04", "Bảng lương", "Xác nhận trả lương", "Có phiếu lương nháp", "Chọn phiếu, bấm Xác nhận trả", "Phiếu lương", "Trạng thái chuyển Đã trả")

    add("TC-NOTI-01", "Thông báo", "Tạo thông báo", "Vai trò Admin/Giám đốc/Trưởng phòng", "Bấm Tạo thông báo, nhập tiêu đề/nội dung, gửi", "Thông báo test", "Thông báo mới xuất hiện và chưa đọc")
    add("TC-NOTI-02", "Thông báo", "Lọc thông báo chưa đọc", "Có thông báo đã đọc/chưa đọc", "Tick Chỉ hiện chưa đọc", "Thông báo", "Danh sách chỉ còn thông báo chưa đọc")
    add("TC-NOTI-03", "Thông báo", "Đánh dấu đã đọc", "Có thông báo chưa đọc", "Chọn đánh dấu đã đọc", "Thông báo", "Số thông báo chưa đọc giảm")
    add("TC-NOTI-04", "Thông báo", "Đính kèm file", "Có file test", "Chọn tệp khi tạo thông báo, gửi, mở tệp", "file.txt", "Tên tệp hiển thị và mở được")

    add("TC-REP-01", "Báo cáo", "Xuất báo cáo nhân viên DOCX", "Có dữ liệu nhân viên", "Vào Báo cáo, chọn nhân viên, lưu DOCX", "NhanVien.docx", "File được tạo và mở được")
    add("TC-REP-02", "Báo cáo", "Xuất báo cáo chấm công XLSX", "Có dữ liệu chấm công", "Chọn báo cáo chấm công, lưu XLSX", "ChamCong.xlsx", "File được tạo và có dữ liệu")
    add("TC-REP-03", "Báo cáo", "Xuất báo cáo lương", "Vai trò có quyền lương", "Chọn báo cáo lương, lưu file", "Luong.xlsx", "File chứa bảng lương theo bộ lọc")
    add("TC-SET-01", "Cài đặt", "Tạo tài khoản nhân viên", "Vai trò Admin", "Bấm tạo tài khoản mẫu", "nhanvienXX", "Tài khoản xuất hiện trong danh sách")
    add("TC-SET-02", "Cài đặt", "Khóa/mở tài khoản", "Có tài khoản", "Chọn tài khoản, bấm khóa/mở", "nhanvien", "Trạng thái thay đổi đúng")
    add("TC-SET-03", "Cài đặt", "Reset mật khẩu", "Có tài khoản", "Chọn tài khoản, bấm đặt lại mật khẩu", "nhanvien", "Mật khẩu reset về HRM_INITIAL_PASSWORD")
    add("TC-SET-04", "Cài đặt", "Sao lưu dữ liệu", "Có dữ liệu", "Bấm Sao lưu, chọn nơi lưu", ".hrmbackup.json", "File backup được tạo")
    add("TC-SET-05", "Cài đặt", "Phục hồi dữ liệu", "Có file backup", "Bấm Phục hồi, chọn file", ".hrmbackup.json", "Dữ liệu được nạp lại")

    add("TC-AUTO-01", "Unit test", "Salt khác nhau cho cùng mật khẩu", "Đã build Release", "Chạy dotnet test", "BaoMatMatKhauTests", "Test pass", "Automated", "Pass")
    add("TC-AUTO-02", "Unit test", "Xác minh mật khẩu đúng/sai", "Đã build Release", "Chạy dotnet test", "BaoMatMatKhauTests", "Test pass", "Automated", "Pass")
    add("TC-AUTO-03", "Unit test", "Chuỗi kết nối luôn có nguồn mặc định", "Đã build Release", "Chạy dotnet test", "CauHinhUngDungTests", "Test pass", "Automated", "Pass")
    add("TC-AUTO-04", "Unit test", "Mật khẩu khởi tạo đủ mạnh", "Đã build Release", "Chạy dotnet test", "CauHinhUngDungTests", "Test pass", "Automated", "Pass")
    add("TC-AUTO-05", "Unit test", "Nhân viên chỉ xem dữ liệu cá nhân", "Đã build Release", "Chạy dotnet test", "PhanQuyenPhamViTests", "Test pass", "Automated", "Pass")
    add("TC-AUTO-06", "Unit test", "Trưởng phòng nhân sự chỉ xem phòng nhân sự", "Đã build Release", "Chạy dotnet test", "PhanQuyenPhamViTests", "Test pass", "Automated", "Pass")
    add("TC-AUTO-07", "Unit test", "Giám đốc xem toàn hệ thống", "Đã build Release", "Chạy dotnet test", "PhanQuyenPhamViTests", "Test pass", "Automated", "Pass")
    add("TC-AUTO-08", "Unit test", "Tính kỳ đánh giá theo quý", "Đã build Release", "Chạy dotnet test", "NghiepVuNhanSuTests", "Test pass", "Automated", "Pass")
    add("TC-AUTO-09", "Unit test", "Tính phiếu lương theo công và nghỉ phép", "Đã build Release", "Chạy dotnet test", "NghiepVuNhanSuTests", "Test pass", "Automated", "Pass")
    add("TC-AUTO-10", "Unit test", "Tính số ngày giao với tháng lương", "Đã build Release", "Chạy dotnet test", "NghiepVuNhanSuTests", "Test pass", "Automated", "Pass")
    add("TC-AUTO-11", "Unit test", "Tạo mã nhân viên tiếp theo", "Đã build Release", "Chạy dotnet test", "NghiepVuNhanSuTests", "Test pass", "Automated", "Pass")
    add("TC-AUTO-12", "Unit test", "Sắp xếp chức vụ quản lý trước nhân viên", "Đã build Release", "Chạy dotnet test", "SapXepChucVuTests", "Test pass", "Automated", "Pass")
    add("TC-AUTO-13", "Unit test", "Cấp bậc chức vụ thay đổi theo vị trí", "Đã build Release", "Chạy dotnet test", "SapXepChucVuTests", "Test pass", "Automated", "Pass")
    return rows


def autosize(ws) -> None:
    for column_cells in ws.columns:
        max_len = 0
        column = get_column_letter(column_cells[0].column)
        for cell in column_cells:
            value = str(cell.value or "")
            max_len = max(max_len, min(len(value), 80))
        ws.column_dimensions[column].width = max(12, min(max_len + 2, 55))


def workbook_testcases() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "TongHop"
    title_font = Font(name="Calibri", size=16, bold=True, color="172033")
    header_font = Font(name="Calibri", size=11, bold=True, color="172033")
    body_font = Font(name="Calibri", size=10, color="172033")
    header_fill = PatternFill("solid", fgColor="E8EEF5")
    pass_fill = PatternFill("solid", fgColor="E8F5E9")
    pending_fill = PatternFill("solid", fgColor="FFF7E6")
    border = Border(
        left=Side(style="thin", color="D8E0EA"),
        right=Side(style="thin", color="D8E0EA"),
        top=Side(style="thin", color="D8E0EA"),
        bottom=Side(style="thin", color="D8E0EA"),
    )

    cases = make_test_cases()
    ws["A1"] = "TÀI LIỆU TEST CASE - PHẦN MỀM QUẢN LÝ NHÂN SỰ WPF"
    ws["A1"].font = title_font
    ws.merge_cells("A1:I1")
    ws["A2"] = f"Dự án: {PROJECT_CODE}"
    ws["A3"] = f"Nhóm: {TEAM}"
    ws["A4"] = f"Ngày lập: {TODAY}"
    ws["A6"] = "Tổng số test case"
    ws["B6"] = len(cases)
    ws["A7"] = "Automated Pass"
    ws["B7"] = sum(1 for r in cases if r[7] == "Automated" and r[8] == "Pass")
    ws["A8"] = "Manual cần chạy"
    ws["B8"] = sum(1 for r in cases if r[7] == "Manual")
    ws["A10"] = "Ghi chú"
    ws["B10"] = "Các test tự động đã pass qua dotnet test. Các test thủ công cần chạy trên app trước khi nộp chính thức."
    ws.merge_cells("B10:I10")

    summary_headers = ["Phân hệ", "Số TC", "Automated Pass", "Manual cần chạy"]
    for col, text in enumerate(summary_headers, 1):
        cell = ws.cell(12, col, text)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    modules = sorted(set(r[1] for r in cases))
    row_idx = 13
    for module in modules:
        module_cases = [r for r in cases if r[1] == module]
        values = [
            module,
            len(module_cases),
            sum(1 for r in module_cases if r[7] == "Automated" and r[8] == "Pass"),
            sum(1 for r in module_cases if r[7] == "Manual"),
        ]
        for col, value in enumerate(values, 1):
            cell = ws.cell(row_idx, col, value)
            cell.font = body_font
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
        row_idx += 1
    autosize(ws)

    headers = ["ID", "Phân hệ", "Mục tiêu", "Tiền điều kiện", "Bước thực hiện", "Dữ liệu", "Kết quả mong đợi", "Loại", "Trạng thái"]
    groups: dict[str, list[list[str]]] = {}
    for case in cases:
        groups.setdefault(case[1], []).append(case)

    for module, module_cases in groups.items():
        sheet_name = module[:31]
        ws = wb.create_sheet(sheet_name)
        ws.freeze_panes = "A2"
        for col, text in enumerate(headers, 1):
            cell = ws.cell(1, col, text)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row_idx, case in enumerate(module_cases, 2):
            for col_idx, value in enumerate(case, 1):
                cell = ws.cell(row_idx, col_idx, value)
                cell.font = body_font
                cell.border = border
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                if col_idx == 9:
                    cell.fill = pass_fill if value == "Pass" else pending_fill
        widths = [16, 18, 32, 34, 48, 24, 44, 14, 14]
        for col_idx, width in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(col_idx)].width = width
        for row in ws.iter_rows(min_row=2):
            ws.row_dimensions[row[0].row].height = 60
        ws.auto_filter.ref = ws.dimensions

    path = OUTPUT / "TaiLieu_TestCase_QuanLyNhanSuWpf.xlsx"
    wb.save(path)
    return path


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    outputs = [
        doc_feasibility(),
        doc_plan(),
        doc_srs(),
        doc_design(),
        doc_db(),
        doc_install(),
        doc_user(),
        doc_git_checklist(),
        workbook_testcases(),
    ]
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()

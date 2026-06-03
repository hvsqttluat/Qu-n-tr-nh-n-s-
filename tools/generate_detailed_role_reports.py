from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Users\LUAT\Documents\CNPM nhóm 3")
OUT_DIR = ROOT / "Bàn giao phần mềm quản lý nhân sự - Nhóm 3" / "Báo cáo chi tiết theo chức năng"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def add_title(doc: Document, title: str, subtitle: str, owner: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Người trình bày: {owner}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    keep_with_next(p)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(31, 78, 121 if level == 1 else 89)
        run.font.size = Pt(16 if level == 1 else 13)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12.5)
        rest = text[len(bold_prefix) :]
        r = p.add_run(rest)
    else:
        r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12.2)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12.2)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF2F8")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(31, 78, 121)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    doc.add_paragraph()


def add_code_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Khi nói đến"
    hdr[1].text = "Mở file / vị trí"
    hdr[2].text = "Nói với thầy"
    widths = [4.0, 6.2, 7.0]
    for i, cell in enumerate(hdr):
        set_cell_shading(cell, "1F4E79")
        set_cell_width(cell, widths[i])
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10.5)


def add_qa(doc: Document, qas: list[tuple[str, str]]) -> None:
    for question, answer in qas:
        add_para(doc, "Thầy hỏi: " + question, bold_prefix="Thầy hỏi:")
        add_para(doc, "Trả lời: " + answer, bold_prefix="Trả lời:")


def setup_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12.5)
    return doc


LONG = {
    "filename": "Long - Đăng nhập, tài khoản, bảo mật, CSDL.docx",
    "owner": "Long",
    "title": "Kịch bản báo cáo chi tiết - Đăng nhập, tài khoản, bảo mật, CSDL",
    "subtitle": "Phần chức năng khởi động hệ thống, xác thực người dùng và nền dữ liệu",
    "sections": [
        {
            "heading": "1. Lời đọc mở đầu và phạm vi phụ trách",
            "paragraphs": [
                "Thưa thầy, phần em phụ trách là nhóm chức năng đăng nhập, tài khoản người dùng, bảo mật và cơ sở dữ liệu. Đây là phần được chạy đầu tiên khi mở phần mềm, vì trước khi người dùng sử dụng các phân hệ nhân sự thì hệ thống phải xác thực xem người đó là ai, có tài khoản hợp lệ hay không và vai trò của người đó trong hệ thống là gì.",
                "Ứng dụng của nhóm em được xây dựng bằng WPF C#. Về mặt giao diện, màn hình đầu tiên là LoginWindow.xaml. Về mặt xử lý, khi bấm nút Đăng nhập, file LoginWindow.xaml.cs sẽ lấy tên đăng nhập và mật khẩu người dùng nhập vào, sau đó gọi lớp KhoXacThuc để kiểm tra trong SQL Server. Nếu tài khoản hợp lệ, chương trình tạo phiên đăng nhập và mở MainWindow. Nếu sai tài khoản hoặc sai mật khẩu, chương trình giữ nguyên màn hình đăng nhập và hiển thị thông báo lỗi.",
                "Điểm em muốn nhấn mạnh trong phần này là hệ thống không lưu mật khẩu gốc. Mật khẩu được băm bằng thuật toán PBKDF2-SHA256, có salt riêng và số vòng lặp. Khi đăng nhập, mật khẩu nhập vào được băm lại rồi so sánh với hash đã lưu. Cách làm này an toàn hơn so với lưu mật khẩu trực tiếp trong cơ sở dữ liệu.",
                "Ngoài ra, ứng dụng còn có phần tự đảm bảo cơ sở dữ liệu. Khi kết nối SQL Server, chương trình kiểm tra database HRManagementDB đã tồn tại hay chưa. Nếu chưa có thì tạo database, sau đó tạo các bảng nghiệp vụ như HR_Employees, HR_Departments, HR_Attendances, HR_LeaveRequests, HR_Payslips, HR_Users và HR_AuditLogs. Vì vậy khi bàn giao, người nhận chỉ cần có SQL Server phù hợp, ứng dụng có thể tự khởi tạo cấu trúc dữ liệu ban đầu.",
            ],
            "bullets": [
                "Mục tiêu chính: cho phép người dùng đăng nhập an toàn trước khi vào phần mềm.",
                "Dữ liệu chính: bảng HR_Users lưu tài khoản, bảng HR_AuditLogs lưu nhật ký đăng nhập.",
                "Code chính cần nhớ: App.xaml, LoginWindow.xaml, LoginWindow.xaml.cs, KhoXacThuc.cs, BaoMatMatKhau.cs, CauHinhUngDung.cs, SoDoQuanTriSql.cs.",
            ],
        },
        {
            "heading": "2. Vừa nói vừa demo trên phần mềm",
            "paragraphs": [
                "Khi demo, em mở phần mềm lên trước. Em chỉ vào màn hình đăng nhập và nói: đây là màn hình đầu tiên của ứng dụng. Trong App.xaml, thuộc tính StartupUri đang trỏ tới LoginWindow.xaml, nên khi chạy chương trình WPF sẽ mở màn hình đăng nhập trước.",
                "Sau đó em chỉ vào ô Tên đăng nhập và Mật khẩu. Hai ô này được thiết kế trong LoginWindow.xaml. Người dùng nhập dữ liệu tại đây. Khi bấm nút Đăng nhập, sự kiện Click sẽ gọi hàm DangNhap_Click trong LoginWindow.xaml.cs. Hàm này không tự xử lý toàn bộ bảo mật mà gọi tiếp phương thức DangNhapAsync của lớp KhoXacThuc.",
                "Nếu đăng nhập thành công, LoginWindow.xaml.cs tạo đối tượng MainWindow và truyền phiên đăng nhập sang. Việc truyền phiên đăng nhập này giúp màn hình chính biết tên người dùng, vai trò và quyền sử dụng các phân hệ. Sau đó chương trình gọi Show() để mở MainWindow và đóng cửa sổ đăng nhập.",
                "Nếu đăng nhập thất bại, ví dụ sai mật khẩu, lớp KhoXacThuc trả về kết quả không thành công. Khi đó LoginWindow.xaml.cs gán nội dung thông báo vào ThongBaoTextBlock, focus lại ô mật khẩu và không mở màn hình chính. Cách này giúp người dùng biết lỗi nhưng không làm lộ chi tiết nhạy cảm.",
            ],
            "numbered": [
                "Mở app và chỉ màn hình đăng nhập.",
                "Nói App.xaml quyết định màn hình khởi động.",
                "Nhập tài khoản, mật khẩu.",
                "Bấm Đăng nhập và chỉ luồng gọi từ LoginWindow.xaml.cs sang KhoXacThuc.cs.",
                "Đăng nhập đúng thì mở MainWindow, đăng nhập sai thì hiển thị lỗi.",
            ],
        },
        {
            "heading": "3. Luồng dữ liệu và vị trí code cần chỉ",
            "paragraphs": [
                "Luồng dữ liệu của phần đăng nhập đi theo thứ tự rất rõ: người dùng nhập dữ liệu trên giao diện, code-behind nhận dữ liệu, lớp xác thực kiểm tra SQL Server, lớp bảo mật kiểm tra hash mật khẩu, sau đó kết quả quay ngược về giao diện. Nếu hợp lệ thì mở màn hình chính, nếu không hợp lệ thì hiển thị lỗi.",
                "Em có thể mô tả bằng lời như sau: Từ TextBox và PasswordBox trên LoginWindow.xaml, dữ liệu được lấy trong hàm DangNhap của LoginWindow.xaml.cs. Sau đó tên đăng nhập và mật khẩu được truyền vào KhoXacThuc.DangNhapAsync. Trong KhoXacThuc, chương trình mở kết nối SQL Server, đọc tài khoản từ bảng HR_Users, kiểm tra IsActive, kiểm tra PasswordHash bằng BaoMatMatKhau.XacMinhMatKhau, cập nhật số lần đăng nhập sai hoặc thời điểm đăng nhập cuối cùng, rồi ghi nhật ký vào HR_AuditLogs.",
            ],
            "code_table": [
                ("Màn hình chạy đầu tiên", "App.xaml dòng 4", "StartupUri trỏ tới LoginWindow.xaml nên app mở màn hình đăng nhập trước."),
                ("Giao diện đăng nhập", "LoginWindow.xaml dòng 93 và 175", "Dòng 93 là ảnh nền, dòng 175 là tiêu đề form đăng nhập."),
                ("Bấm nút đăng nhập", "LoginWindow.xaml.cs dòng 22", "Sự kiện Click gọi DangNhap_Click, sau đó gọi hàm DangNhap."),
                ("Gọi xác thực", "LoginWindow.xaml.cs dòng 53", "Code gọi khoXacThuc.DangNhapAsync để kiểm tra tài khoản."),
                ("Mở màn hình chính", "LoginWindow.xaml.cs dòng 72", "Nếu thành công thì tạo MainWindow và truyền phiên đăng nhập."),
                ("Logic đăng nhập", "KhoXacThuc.cs dòng 11", "DangNhapAsync là hàm trung tâm xử lý đăng nhập."),
                ("Kiểm tra mật khẩu", "KhoXacThuc.cs dòng 39", "So sánh mật khẩu nhập vào với hash đã lưu."),
                ("Hash mật khẩu", "BaoMatMatKhau.cs dòng 21", "Dùng PBKDF2-SHA256 để băm mật khẩu."),
                ("Tạo bảng tài khoản", "SoDoQuanTriSql.cs dòng 10", "Tạo bảng HR_Users nếu chưa tồn tại."),
                ("Ghi nhật ký", "KhoXacThuc.cs dòng 179", "Ghi LoginSuccess hoặc LoginFailed vào HR_AuditLogs."),
            ],
        },
        {
            "heading": "4. Bảo mật và cơ sở dữ liệu nói sao cho chắc",
            "paragraphs": [
                "Khi thầy hỏi về bảo mật, em không nên nói chung chung là mật khẩu được mã hóa. Từ đúng hơn là mật khẩu được băm. Mã hóa thường có thể giải mã ngược nếu có khóa, còn hash là một chiều. Hệ thống lưu PasswordHash, PasswordSalt và PasswordIterations, không lưu mật khẩu gốc.",
                "Salt là chuỗi ngẫu nhiên được thêm vào quá trình băm mật khẩu. Nhờ salt, hai người dùng có cùng mật khẩu vẫn tạo ra hash khác nhau. PasswordIterations là số vòng lặp của PBKDF2. Số vòng lặp càng nhiều thì việc thử đoán mật khẩu càng tốn thời gian hơn.",
                "Về cơ sở dữ liệu, ứng dụng dùng SQL Server. Chuỗi kết nối có thể lấy từ appsettings hoặc biến môi trường HRM_CONNECTION_STRING. Lớp CauHinhUngDung có nhiệm vụ chọn chuỗi kết nối và đảm bảo database tồn tại. Lớp KhoDuLieuNhanSu tạo các bảng nghiệp vụ, còn SoDoQuanTriSql tạo các bảng quản trị tài khoản và audit.",
                "Khi nói phần này, em nên nhấn vào tính khả thi: nếu triển khai trên máy khác, chỉ cần cài SQL Server đúng cấu hình, ứng dụng có thể tự tạo database HRManagementDB và nạp dữ liệu mẫu phục vụ demo. Đây là điểm giúp phần mềm dễ bàn giao và dễ chạy thử.",
            ],
            "bullets": [
                "Không nói: hệ thống lưu mật khẩu đã mã hóa. Nên nói: hệ thống lưu hash mật khẩu.",
                "Không nói: salt là mật khẩu phụ. Nên nói: salt là dữ liệu ngẫu nhiên giúp hash khác nhau.",
                "Không nói: app tự tạo mọi thứ không cần SQL Server. Nên nói: app tự tạo database khi máy đã có SQL Server và chuỗi kết nối phù hợp.",
            ],
        },
        {
            "heading": "5. Câu hỏi thầy có thể hỏi và cách trả lời",
            "qas": [
                ("Vì sao phải có màn hình đăng nhập?", "Để xác định người dùng trước khi vào hệ thống. Sau khi đăng nhập, ứng dụng biết người dùng là ai, vai trò gì và được phép dùng những phân hệ nào."),
                ("Hash mật khẩu là gì?", "Hash là kết quả băm một chiều của mật khẩu. Hệ thống không lưu mật khẩu gốc. Khi đăng nhập, mật khẩu nhập vào được băm lại rồi so sánh với hash đã lưu."),
                ("Salt là gì?", "Salt là chuỗi ngẫu nhiên thêm vào khi băm mật khẩu. Nó làm cho hai tài khoản có cùng mật khẩu vẫn có hash khác nhau, giảm rủi ro dò mật khẩu hàng loạt."),
                ("PBKDF2-SHA256 là gì?", "PBKDF2 là thuật toán dẫn xuất khóa dùng nhiều vòng lặp. SHA256 là hàm băm bên trong. Mục đích là làm quá trình thử mật khẩu tốn thời gian hơn."),
                ("Nếu người dùng nhập sai mật khẩu thì sao?", "KhoXacThuc trả về kết quả thất bại, ghi audit LoginFailed, tăng số lần sai và LoginWindow hiển thị thông báo lỗi thay vì mở MainWindow."),
                ("App tự tạo CSDL ở đâu?", "Phần kiểm tra và tạo database nằm trong CauHinhUngDung.DamBaoCoSoDuLieuAsync. Phần tạo bảng nghiệp vụ nằm trong KhoDuLieuNhanSu, còn bảng tài khoản và audit nằm trong SoDoQuanTriSql."),
                ("Bảng HR_AuditLogs dùng để làm gì?", "Dùng để ghi nhật ký thao tác quan trọng như đăng nhập thành công, đăng nhập thất bại, tạo tài khoản, đổi trạng thái tài khoản. Khi cần kiểm tra lịch sử hệ thống thì xem bảng này."),
            ],
        },
        {
            "heading": "6. Nếu thầy yêu cầu sửa code tại chỗ",
            "paragraphs": [
                "Trường hợp dễ nhất là thầy yêu cầu đổi câu thông báo lỗi. Khi đó em mở LoginWindow.xaml.cs, tìm đoạn gán ThongBaoTextBlock.Text trong hàm DangNhap. Ví dụ đổi câu Vui lòng nhập tên đăng nhập và mật khẩu thành Vui lòng điền đầy đủ thông tin đăng nhập. Đây là sửa giao diện thông báo, không ảnh hưởng bảo mật.",
                "Nếu thầy yêu cầu đổi giao diện đăng nhập, em mở LoginWindow.xaml. Ví dụ đổi tiêu đề Đăng nhập tài khoản, đổi text mô tả, đổi màu nút trong style NutDangNhapHienDai, hoặc đổi ảnh nền tại ImageBrush dòng 93. Nếu đổi ảnh nền, cần để ảnh trong thư mục Assets và khai báo Resource trong QuanLyNhanSuWpf.csproj.",
                "Nếu thầy yêu cầu thêm vai trò tài khoản, em phải nói đây là sửa logic nghiệp vụ. Cần cập nhật dữ liệu RoleName trong HR_Users, cập nhật quyền trong ViewModel nếu quyền đó ảnh hưởng đến phân hệ, và kiểm tra các property CoQuyen... trong ManHinhChinhViewModel.",
                "Nếu thầy yêu cầu thay số vòng lặp hash mật khẩu, em mở BaoMatMatKhau.cs và tìm hằng SoVongLapMacDinh. Sau đó giải thích rằng tăng vòng lặp sẽ tăng độ an toàn nhưng cũng làm đăng nhập chậm hơn một chút. Đây là thay đổi cần kiểm thử lại đăng nhập và tạo tài khoản mới.",
            ],
            "numbered": [
                "Sửa text/màu giao diện: LoginWindow.xaml.",
                "Sửa thông báo đăng nhập: LoginWindow.xaml.cs.",
                "Sửa kiểm tra tài khoản: KhoXacThuc.cs.",
                "Sửa thuật toán hash hoặc số vòng lặp: BaoMatMatKhau.cs.",
                "Sửa cấu hình SQL Server: appsettings.json hoặc CauHinhUngDung.cs.",
                "Sửa bảng tài khoản/audit: SoDoQuanTriSql.cs.",
            ],
        },
    ],
}


LUAT = {
    "filename": "Luật - Giao diện chính, nhân viên, phòng ban, tuyển dụng.docx",
    "owner": "Luật",
    "title": "Kịch bản báo cáo chi tiết - Giao diện chính, nhân viên, phòng ban, tuyển dụng",
    "subtitle": "Phần chức năng giao diện WPF, điều hướng, CRUD nhân sự và tuyển dụng",
    "sections": [
        {
            "heading": "1. Lời đọc mở đầu và phạm vi phụ trách",
            "paragraphs": [
                "Thưa thầy, phần em phụ trách là giao diện chính của phần mềm và nhóm chức năng quản lý nhân sự cơ bản gồm nhân viên, phòng ban và tuyển dụng. Sau khi người dùng đăng nhập thành công, hệ thống mở MainWindow. Đây là màn hình làm việc chính, có thanh điều hướng bên trái, vùng nội dung ở giữa và các phân hệ được hiển thị theo lựa chọn của người dùng.",
                "Về giao diện, em sử dụng XAML trong WPF để xây dựng bố cục. Các thành phần chính gồm Grid để chia vùng, Border để tạo khối nội dung, DataGrid để hiển thị danh sách, TextBox để nhập dữ liệu, ComboBox để chọn phòng ban hoặc chức vụ, DatePicker để chọn ngày và Button để thao tác. Các style dùng chung như nút chính, nút phụ, ô nhập, thẻ thống kê được khai báo trong App.xaml để toàn bộ giao diện thống nhất.",
                "Về logic C#, màn hình chính sử dụng mô hình gần với MVVM. MainWindow.xaml chủ yếu mô tả giao diện, còn dữ liệu, trạng thái đang chọn, danh sách nhân viên, danh sách phòng ban và các command xử lý nút nằm trong ManHinhChinhViewModel.cs. Khi bấm nút Lưu phòng ban hoặc Tạo ứng viên, XAML không tự xử lý dữ liệu mà gọi command trong ViewModel.",
                "Phần em sẽ demo sâu nhất là Phòng ban, vì đây là phân hệ gọn và thể hiện đủ các thao tác CRUD: tạo mới, nạp dòng chọn, lưu phòng ban và xóa phòng ban. Ngoài ra em sẽ chỉ thêm phân hệ Nhân viên và Tuyển dụng để thầy thấy luồng nghiệp vụ nhân sự đầy đủ: có phòng ban, có hồ sơ nhân viên, có ứng viên và có thể chuyển ứng viên thành nhân viên.",
            ],
            "bullets": [
                "Giao diện chính: MainWindow.xaml.",
                "Logic điều hướng, binding, command: ManHinhChinhViewModel.cs.",
                "Model dữ liệu: MoHinh.cs.",
                "Lưu đọc SQL Server: KhoDuLieuNhanSu.cs.",
                "Thiết kế trình bày từng phân hệ: ThietKeGiaoDienWindow.xaml.",
            ],
        },
        {
            "heading": "2. Vừa nói vừa demo trên phần mềm",
            "paragraphs": [
                "Khi demo, em đăng nhập vào phần mềm rồi chỉ thanh điều hướng bên trái. Em nói: đây là sidebar điều hướng các phân hệ. Mỗi mục như Tổng quan, Nhân viên, Phòng ban, Ứng viên, Chấm công đều được thiết kế trong MainWindow.xaml. Khi người dùng bấm một mục, Command ChonMucLenh trong ViewModel sẽ đổi giá trị MucDangChon. Các Grid trong XAML dùng Binding MucDangChon để hiện đúng phân hệ và ẩn các phân hệ còn lại.",
                "Tiếp theo em bấm Nhân viên. Em chỉ bảng DataGrid bên trái hoặc giữa màn hình, ô tìm kiếm, bộ lọc phòng ban và form nhập thông tin nhân viên. Em nói: dữ liệu nhân viên được binding vào DanhSachNhanVienView, khi người dùng tìm kiếm hoặc chọn phòng ban thì ViewModel refresh danh sách hiển thị.",
                "Sau đó em bấm Phòng ban. Đây là phần em nên demo chính. Em chỉ các nút Tạo mới, Nạp dòng chọn, Lưu phòng ban và Xóa. Em giải thích rằng nút Lưu phòng ban trong XAML bind tới LuuPhongBanLenh. Khi bấm, ViewModel lấy dữ liệu từ BieuMauPhongBan, kiểm tra tên phòng ban, cập nhật ObservableCollection phòng ban và lưu qua KhoDuLieuNhanSu nếu đang kết nối SQL Server.",
                "Cuối cùng em bấm Ứng viên. Em nói phân hệ này quản lý danh sách ứng viên, thông tin liên hệ, vị trí ứng tuyển và trạng thái tuyển dụng. Khi ứng viên đạt yêu cầu, hệ thống có chức năng tiếp nhận thành nhân viên, tức là tạo hồ sơ nhân viên từ dữ liệu ứng viên. Đây là luồng nối giữa tuyển dụng và quản lý hồ sơ nhân sự.",
            ],
            "numbered": [
                "Đăng nhập vào MainWindow.",
                "Chỉ sidebar và nói về điều hướng phân hệ.",
                "Bấm Nhân viên, chỉ DataGrid, ô tìm kiếm, form nhập.",
                "Bấm Phòng ban, demo tạo mới hoặc nạp dòng chọn và lưu.",
                "Bấm Ứng viên, nói về tạo ứng viên và chuyển thành nhân viên.",
            ],
        },
        {
            "heading": "3. Luồng dữ liệu giao diện và code cần chỉ",
            "paragraphs": [
                "Luồng dữ liệu ở phần giao diện chính đi theo hướng: người dùng thao tác trên XAML, Binding chuyển dữ liệu vào ViewModel, Command trong ViewModel xử lý nghiệp vụ, dữ liệu được cập nhật trong các ObservableCollection và sau đó giao diện tự cập nhật lại. Nếu có SQL Server thì KhoDuLieuNhanSu lưu dữ liệu xuống database.",
                "Ví dụ ở phân hệ Phòng ban, người dùng nhập tên phòng ban trong TextBox. TextBox này binding tới BieuMauPhongBan.TenPhongBan. Khi bấm Lưu phòng ban, nút Button gọi LuuPhongBanLenh. Command này chạy hàm xử lý trong ManHinhChinhViewModel.cs. Sau khi lưu xong, danh sách DuLieu.PhongBan thay đổi, DataGrid đang binding vào DuLieu.PhongBan nên giao diện tự hiển thị dữ liệu mới.",
                "Ví dụ ở điều hướng, các RadioButton trong sidebar binding Command ChonMucLenh và truyền CommandParameter là tên phân hệ. Khi người dùng bấm Phòng ban, ViewModel đổi MucDangChon thành Phòng ban. Grid của phân hệ Phòng ban có Visibility binding MucDangChon với converter HienThiTheoMuc, nên nó hiện lên; các Grid khác ẩn đi.",
            ],
            "code_table": [
                ("Sidebar điều hướng", "MainWindow.xaml dòng 144-156", "Các mục Nhân viên, Phòng ban, Chấm công... bind tới ChonMucLenh."),
                ("Màn Nhân viên", "MainWindow.xaml dòng 583", "Grid này chỉ hiện khi MucDangChon là Nhân viên."),
                ("Nút nhân viên", "MainWindow.xaml dòng 598-604", "Các nút tạo mới, nạp dòng, lưu, xóa gọi command trong ViewModel."),
                ("Bảng nhân viên", "MainWindow.xaml dòng 636", "DataGrid bind vào DanhSachNhanVienView."),
                ("Màn Phòng ban", "MainWindow.xaml dòng 694", "Đây là phân hệ em nên demo chính."),
                ("Nút phòng ban", "MainWindow.xaml dòng 708-711", "Tạo mới, nạp dòng chọn, lưu phòng ban, xóa phòng ban."),
                ("Bảng phòng ban", "MainWindow.xaml dòng 731", "DataGrid bind vào DuLieu.PhongBan."),
                ("Form phòng ban", "MainWindow.xaml dòng 747-755", "Các ô nhập bind vào BieuMauPhongBan."),
                ("Màn Ứng viên", "MainWindow.xaml dòng 762", "Phân hệ tuyển dụng và chuyển ứng viên thành nhân viên."),
                ("Thiết kế UI", "ThietKeGiaoDienWindow.xaml dòng 1", "File dùng để trình bày bản thiết kế từng phân hệ trong Visual Studio Designer."),
            ],
        },
        {
            "heading": "4. Thiết kế giao diện đã làm những gì",
            "paragraphs": [
                "Về thiết kế giao diện, em đã làm ba lớp. Lớp thứ nhất là giao diện chạy thật của phần mềm, gồm LoginWindow.xaml và MainWindow.xaml. LoginWindow là màn hình đăng nhập, đã được thiết kế lại với nền ảnh văn phòng, form đăng nhập nổi và các nút rõ ràng. MainWindow là màn hình chính, có sidebar, vùng dashboard, các phân hệ nghiệp vụ và form nhập liệu.",
                "Lớp thứ hai là hệ thống style dùng chung trong App.xaml. Những thành phần lặp lại như Button, TextBox, DataGrid, Border thẻ thống kê được định nghĩa style để giao diện thống nhất. Nhờ vậy khi muốn đổi màu nút hoặc kiểu ô nhập, có thể sửa một chỗ trong App.xaml thay vì sửa từng nút.",
                "Lớp thứ ba là file thiết kế trình bày ThietKeGiaoDienWindow.xaml. File này không phải luồng chạy chính, nhưng dùng để mở bằng Visual Studio Designer khi báo cáo thiết kế. Trong file này có mục lục thiết kế và các tab/mục tương ứng từng phân hệ như Đăng nhập, Tổng quan, Tuyển dụng, Nhân viên, Phòng ban, Chấm công, Nghỉ phép, Đánh giá, Bảng lương, Báo cáo, Tài khoản và Thông báo.",
                "Khi thầy hỏi thiết kế giao diện ở đâu, em sẽ mở MainWindow.xaml để chỉ giao diện chạy thật, mở App.xaml để chỉ style dùng chung, và mở ThietKeGiaoDienWindow.xaml nếu thầy muốn xem bản thiết kế từng phân hệ trong Visual Studio Designer.",
            ],
            "bullets": [
                "LoginWindow.xaml: màn hình đăng nhập chạy thật.",
                "MainWindow.xaml: màn hình chính và các phân hệ chạy thật.",
                "App.xaml: style dùng chung cho giao diện.",
                "ThietKeGiaoDienWindow.xaml: bản thiết kế trình bày từng phân hệ.",
                "Assets/login-background.png: ảnh nền đăng nhập được nhúng vào project.",
            ],
        },
        {
            "heading": "5. Câu hỏi thầy có thể hỏi và cách trả lời",
            "qas": [
                ("Binding là gì?", "Binding là cơ chế nối dữ liệu giữa giao diện XAML và thuộc tính trong ViewModel. Ví dụ TextBox nhập tên phòng ban binding tới BieuMauPhongBan.TenPhongBan, khi người dùng nhập thì ViewModel nhận dữ liệu."),
                ("Command là gì?", "Command là cách gắn thao tác nút bấm với hàm xử lý trong ViewModel. Ví dụ nút Lưu phòng ban không viết logic trong XAML mà bind tới LuuPhongBanLenh."),
                ("MVVM là gì?", "MVVM là mô hình tách giao diện, dữ liệu hiển thị và logic xử lý. View là XAML, ViewModel là ManHinhChinhViewModel.cs, Model là các lớp dữ liệu trong MoHinh.cs."),
                ("Dữ liệu trên DataGrid lấy từ đâu?", "DataGrid lấy từ các collection trong ViewModel. Ví dụ bảng phòng ban binding vào DuLieu.PhongBan, bảng nhân viên binding vào DanhSachNhanVienView."),
                ("Nếu muốn thêm một trường cho phòng ban thì sửa ở đâu?", "Cần sửa Model trong MoHinh.cs, thêm ô nhập và cột hiển thị trong MainWindow.xaml, sửa ViewModel để nhận dữ liệu, và nếu lưu SQL thì sửa bảng/insert/update trong KhoDuLieuNhanSu.cs."),
                ("Vì sao chọn WPF/XAML?", "WPF phù hợp làm ứng dụng desktop Windows, hỗ trợ XAML để thiết kế giao diện, binding dữ liệu, DataGrid, style và tách giao diện với logic C# khá rõ."),
            ],
        },
        {
            "heading": "6. Nếu thầy yêu cầu sửa code tại chỗ",
            "paragraphs": [
                "Nếu thầy yêu cầu sửa giao diện đơn giản, ví dụ đổi chữ tiêu đề của phân hệ Phòng ban, em mở MainWindow.xaml, tìm Grid của Phòng ban ở dòng 694 và sửa TextBlock tiêu đề gần khu vực đó. Đây là sửa XAML, không ảnh hưởng logic.",
                "Nếu thầy yêu cầu đổi tên cột trong bảng phòng ban, em mở MainWindow.xaml dòng 731-734, tìm DataGridTextColumn Header rồi sửa tên cột. Ví dụ đổi Header Tên phòng ban thành Tên đơn vị. Đây là bài sửa rất dễ và nên nhận nếu thầy hỏi.",
                "Nếu thầy yêu cầu thêm nút, ví dụ thêm nút Làm mới ở phân hệ phòng ban, em thêm một Button trong khu vực nút dòng 708-711. Nếu nút chỉ gọi lại dữ liệu hiện có thì bind tới một command đã có hoặc tạo command mới trong ManHinhChinhViewModel.cs. Khi thêm command mới, cần khai báo property ICommand, khởi tạo trong constructor ViewModel và viết hàm xử lý.",
                "Nếu thầy yêu cầu thêm trường Ghi chú cho phòng ban, em cần nói đây là sửa đủ 4 lớp: model, giao diện, ViewModel và CSDL. Model thêm thuộc tính GhiChu trong lớp PhongBan. Giao diện thêm TextBox trong form phòng ban và cột DataGrid. ViewModel khi lưu phòng ban phải lấy BieuMauPhongBan.GhiChu. CSDL phải thêm cột GhiChu vào bảng HR_Departments và sửa câu INSERT/UPDATE trong KhoDuLieuNhanSu.cs.",
            ],
            "numbered": [
                "Sửa chữ/nút/cột: MainWindow.xaml.",
                "Sửa style chung: App.xaml.",
                "Sửa logic nút bấm: ManHinhChinhViewModel.cs.",
                "Sửa dữ liệu model: MoHinh.cs.",
                "Sửa lưu SQL: KhoDuLieuNhanSu.cs.",
                "Sửa bản thiết kế trình bày: ThietKeGiaoDienWindow.xaml.",
            ],
        },
    ],
}


TUYEN = {
    "filename": "Tuyến - Chấm công, nghỉ phép, đánh giá, bảng lương, báo cáo, thông báo.docx",
    "owner": "Tuyến",
    "title": "Kịch bản báo cáo chi tiết - Chấm công, nghỉ phép, đánh giá, bảng lương, báo cáo, thông báo",
    "subtitle": "Phần chức năng vận hành nhân sự sau khi có hồ sơ nhân viên",
    "sections": [
        {
            "heading": "1. Lời đọc mở đầu và phạm vi phụ trách",
            "paragraphs": [
                "Thưa thầy, phần em phụ trách là nhóm chức năng vận hành nhân sự sau khi hệ thống đã có dữ liệu nhân viên. Nhóm chức năng này gồm chấm công, nghỉ phép, đánh giá, bảng lương, báo cáo và thông báo. Đây là các chức năng dùng hằng ngày trong doanh nghiệp để theo dõi quá trình làm việc của nhân viên.",
                "Phân hệ Chấm công giúp ghi nhận ngày làm, giờ vào, giờ ra, tổng giờ công và trạng thái ca làm. Phân hệ Nghỉ phép cho phép tạo đơn nghỉ, theo dõi trạng thái và duyệt hoặc từ chối đơn. Phân hệ Đánh giá ghi nhận kỳ đánh giá, người đánh giá, điểm số, nhận xét và trạng thái. Phân hệ Bảng lương tổng hợp lương cơ bản, phụ cấp, thưởng, khấu trừ và thực lĩnh. Phân hệ Báo cáo dùng để xuất dữ liệu nhân viên, chấm công, nghỉ phép và lương. Phân hệ Thông báo là phần bonus giúp gửi thông báo nội bộ.",
                "Về mặt kỹ thuật, giao diện các phân hệ này nằm trong MainWindow.xaml. Các danh sách hiển thị bằng DataGrid, các bộ lọc dùng TextBox, DatePicker, ComboBox. Logic xử lý nằm trong ManHinhChinhViewModel.cs, còn các quy tắc nghiệp vụ hỗ trợ nằm trong QuyTacNghiepVuNhanSu.cs và phần xuất báo cáo dùng BoXuatOffice.cs.",
                "Em sẽ demo theo luồng nghiệp vụ: trước tiên xem chấm công, sau đó tạo hoặc duyệt nghỉ phép, tiếp theo xem đánh giá, bảng lương, xuất báo cáo và cuối cùng là thông báo nội bộ.",
            ],
            "bullets": [
                "Chấm công: theo dõi giờ vào, giờ ra, tổng giờ công, ca cần rà soát.",
                "Nghỉ phép: tạo đơn, duyệt, từ chối, tổng hợp ngày nghỉ.",
                "Đánh giá: nhập điểm, nhận xét, chốt kết quả.",
                "Bảng lương: tính tổng thu nhập, khấu trừ và thực lĩnh.",
                "Báo cáo: xuất dữ liệu ra tài liệu Office.",
                "Thông báo: gửi thông báo nội bộ, có mức độ và tệp đính kèm.",
            ],
        },
        {
            "heading": "2. Vừa nói vừa demo trên phần mềm",
            "paragraphs": [
                "Khi demo, em bắt đầu từ phân hệ Chấm công. Em chỉ các thẻ thống kê ở phía trên như số bản ghi chấm công, ca đang mở, tổng giờ, số ca cần rà soát. Sau đó em chỉ bộ lọc theo từ khóa, từ ngày, đến ngày, phòng ban và trạng thái. Em nói: người dùng có thể lọc dữ liệu để xem chấm công theo phòng ban hoặc theo khoảng thời gian.",
                "Tiếp theo em bấm Nghỉ phép. Em chỉ form gửi đơn nghỉ ở bên trái hoặc phía trên, gồm nhân viên, loại nghỉ, từ ngày, đến ngày. Ở danh sách bên phải có nút Duyệt đơn và Từ chối. Em nói: khi đơn được duyệt thì dữ liệu nghỉ phép ảnh hưởng đến báo cáo tổng quan, số ngày nghỉ và các thống kê trong kỳ.",
                "Sau đó em bấm Đánh giá. Em chỉ các ô chọn nhân viên, người đánh giá, kỳ đánh giá, điểm, trạng thái và nhận xét. Em nói: phân hệ này ghi nhận kết quả đánh giá nhân viên theo kỳ, phục vụ theo dõi hiệu suất và có thể liên quan đến báo cáo cá nhân xuất sắc.",
                "Tiếp theo em bấm Bảng lương. Em nói: bảng lương được lập theo kỳ, mỗi phiếu gồm lương cơ bản, phụ cấp, thưởng, khấu trừ và thực lĩnh. Dữ liệu lương có thể lọc theo kỳ, phòng ban hoặc từ khóa nhân viên.",
                "Cuối cùng em bấm Báo cáo và Thông báo. Ở Báo cáo, em chọn loại báo cáo như nhân viên, chấm công, nghỉ phép hoặc lương rồi xuất file. Ở Thông báo, em chỉ phần tạo thông báo, mức độ thông báo và tệp đính kèm. Đây là phần bonus Notification của nhóm.",
            ],
            "numbered": [
                "Bấm Chấm công, chỉ thống kê và bộ lọc.",
                "Bấm Nghỉ phép, demo gửi đơn và duyệt đơn.",
                "Bấm Đánh giá, chỉ điểm và nhận xét.",
                "Bấm Bảng lương, nói cách tổng hợp lương.",
                "Bấm Báo cáo, nói xuất dữ liệu Office.",
                "Bấm Thông báo, nói bonus notification.",
            ],
        },
        {
            "heading": "3. Luồng dữ liệu và vị trí code cần chỉ",
            "paragraphs": [
                "Luồng dữ liệu trong nhóm chức năng vận hành bắt đầu từ dữ liệu nhân viên. Nhân viên là dữ liệu gốc. Từ nhân viên, hệ thống tạo hoặc đọc dữ liệu chấm công, nghỉ phép, đánh giá và bảng lương. Các dữ liệu này được lưu trong các collection của ViewModel và có thể đồng bộ xuống SQL Server qua KhoDuLieuNhanSu.",
                "Ở Chấm công, DataGrid hiển thị DanhSachChamCongView. Khi người dùng thay đổi từ ngày, đến ngày, phòng ban hoặc trạng thái, ViewModel refresh bộ lọc và cập nhật các thống kê như tổng giờ công, số ca cần rà soát. Ở Nghỉ phép, khi người dùng gửi đơn, ViewModel tạo một bản ghi nghỉ phép mới. Khi quản lý bấm Duyệt đơn hoặc Từ chối, trạng thái bản ghi thay đổi và các thống kê tổng quan được cập nhật.",
                "Ở Báo cáo, ViewModel không xuất trực tiếp từng control trên giao diện. Thay vào đó, ViewModel tạo một đối tượng TaiLieuOffice chứa tiêu đề, dòng mô tả, bảng dữ liệu, sau đó gọi BoXuatOffice.Xuat để ghi ra file. Cách này giúp phần xuất báo cáo tách khỏi giao diện.",
            ],
            "code_table": [
                ("Menu Chấm công", "MainWindow.xaml dòng 148", "Sidebar có mục Chấm công và kiểm tra quyền CoQuyenChamCong."),
                ("Màn Chấm công", "MainWindow.xaml dòng 811", "Grid hiển thị thống kê, bộ lọc và bảng chấm công."),
                ("Bảng chấm công", "MainWindow.xaml dòng 946", "DataGrid bind vào DanhSachChamCongView."),
                ("Màn Nghỉ phép", "MainWindow.xaml dòng 962", "Form tạo đơn nghỉ và danh sách đơn nghỉ phép."),
                ("Duyệt/Từ chối", "MainWindow.xaml dòng 995", "Hai nút gọi command duyệt hoặc từ chối nghỉ phép."),
                ("Màn Đánh giá", "MainWindow.xaml dòng 1052", "Form nhập đánh giá, điểm, trạng thái, nhận xét."),
                ("Tạo báo cáo chấm công", "ManHinhChinhViewModel.cs dòng 3120", "Tạo dữ liệu báo cáo chấm công theo bộ lọc."),
                ("Tạo báo cáo nghỉ phép", "ManHinhChinhViewModel.cs dòng 3173", "Tạo báo cáo nghỉ phép theo kỳ."),
                ("Tạo báo cáo lương", "ManHinhChinhViewModel.cs dòng 3221", "Tạo báo cáo lương theo kỳ và phòng ban."),
                ("Xuất Office", "ManHinhChinhViewModel.cs dòng 3274", "Gọi BoXuatOffice.Xuat để xuất tài liệu."),
                ("Gửi thông báo", "ManHinhChinhViewModel.cs dòng 2848", "Tạo thông báo mới và đưa vào danh sách thông báo."),
            ],
        },
        {
            "heading": "4. Giải thích từng phân hệ theo nghiệp vụ",
            "paragraphs": [
                "Chấm công là phân hệ dùng để biết nhân viên đi làm ngày nào, giờ vào giờ ra ra sao và tổng thời gian làm việc trong kỳ là bao nhiêu. Trong giao diện có các bộ lọc để tránh bảng quá dài. Khi lọc theo phòng ban, hệ thống chỉ hiển thị nhân viên thuộc phòng ban đó. Khi lọc theo trạng thái, người dùng có thể rà soát những ca thiếu giờ ra hoặc cần kiểm tra.",
                "Nghỉ phép là phân hệ xử lý vòng đời của một đơn nghỉ. Đầu tiên đơn được tạo với trạng thái chờ duyệt. Sau đó người có quyền có thể duyệt hoặc từ chối. Nếu đơn được duyệt, dữ liệu này được dùng ở dashboard và báo cáo để tính số nhân sự nghỉ trong kỳ.",
                "Đánh giá là phân hệ phục vụ quản lý hiệu suất. Mỗi đánh giá có nhân viên, người đánh giá, kỳ đánh giá, điểm và nhận xét. Điểm đánh giá có thể được dùng để hiển thị cá nhân xuất sắc hoặc hỗ trợ báo cáo tổng hợp theo phòng ban.",
                "Bảng lương là phân hệ liên quan đến thu nhập. Một phiếu lương gồm lương cơ bản, phụ cấp, thưởng, khấu trừ và thực lĩnh. Khi nói với thầy, em không cần đi quá sâu vào công thức phức tạp, chỉ cần nói thực lĩnh được tính dựa trên tổng khoản cộng trừ trong phiếu lương.",
                "Báo cáo là phần tổng hợp dữ liệu. Người dùng không phải tự lọc thủ công rồi copy dữ liệu, mà phần mềm tạo tài liệu báo cáo theo loại báo cáo được chọn. Đây là điểm thể hiện phần mềm có khả năng hỗ trợ công việc thực tế.",
                "Thông báo là phần bonus. Người có quyền có thể tạo thông báo theo phân hệ, mức độ và nội dung. Danh sách thông báo có thể lọc theo chưa đọc, mở tệp đính kèm và đánh dấu đã đọc.",
            ],
        },
        {
            "heading": "5. Câu hỏi thầy có thể hỏi và cách trả lời",
            "qas": [
                ("Dữ liệu chấm công lấy từ đâu?", "Dữ liệu chấm công nằm trong danh sách chấm công của ViewModel và có thể lưu xuống bảng HR_Attendances trong SQL Server. Giao diện DataGrid bind vào DanhSachChamCongView."),
                ("Tại sao cần bộ lọc chấm công?", "Vì dữ liệu chấm công thường nhiều theo ngày và nhân viên. Bộ lọc giúp xem theo thời gian, phòng ban, trạng thái hoặc từ khóa để dễ rà soát."),
                ("Nghỉ phép ảnh hưởng gì đến báo cáo?", "Đơn nghỉ đã duyệt được dùng để tính số nhân sự nghỉ trong kỳ, số ngày nghỉ và thống kê trên dashboard hoặc báo cáo nghỉ phép."),
                ("Bảng lương tính như thế nào?", "Mỗi phiếu lương có các khoản cộng như lương cơ bản, phụ cấp, thưởng và khoản trừ. Thực lĩnh là kết quả sau khi cộng trừ các khoản đó."),
                ("Xuất báo cáo hoạt động ra sao?", "ViewModel tạo đối tượng TaiLieuOffice chứa tiêu đề, mô tả và bảng dữ liệu. Sau đó gọi BoXuatOffice.Xuat để ghi file báo cáo."),
                ("Thông báo có phải notification không?", "Có. Đây là phần thông báo nội bộ trong ứng dụng. Người dùng có thể tạo thông báo, đặt mức độ, chọn phân hệ và đính kèm tệp."),
                ("Nếu dữ liệu thay đổi thì giao diện cập nhật thế nào?", "Các danh sách dùng ObservableCollection và ViewModel gọi BaoThayDoi hoặc Refresh view, nên DataGrid và thống kê tự cập nhật theo binding."),
            ],
        },
        {
            "heading": "6. Nếu thầy yêu cầu sửa code tại chỗ",
            "paragraphs": [
                "Nếu thầy yêu cầu đổi tên cột trong bảng chấm công, em mở MainWindow.xaml tại khu vực DataGrid chấm công dòng 946 và sửa Header của DataGridTextColumn. Đây là sửa đơn giản nhất, không ảnh hưởng logic.",
                "Nếu thầy yêu cầu thêm bộ lọc trạng thái nghỉ phép, em cần thêm ComboBox trong MainWindow.xaml ở phân hệ Nghỉ phép, thêm property trạng thái đang chọn trong ManHinhChinhViewModel.cs, rồi sửa hàm LocNghiPhep để lọc theo trạng thái đó.",
                "Nếu thầy yêu cầu đổi công thức lương, em tìm phần xử lý bảng lương trong ManHinhChinhViewModel.cs và lớp model Phiếu lương trong MoHinh.cs. Nếu chỉ đổi cách hiển thị thực lĩnh thì sửa property hiển thị hoặc hàm tính thực lĩnh. Nếu lưu thêm trường mới thì sửa cả KhoDuLieuNhanSu.cs và bảng HR_Payslips.",
                "Nếu thầy yêu cầu thêm một loại báo cáo mới, em cần tạo hàm TaoTaiLieuBaoCaoMoi trong ManHinhChinhViewModel.cs theo mẫu các hàm TaoTaiLieuBaoCaoChamCong, TaoTaiLieuBaoCaoNghiPhep, TaoTaiLieuBaoCaoLuong. Sau đó thêm nút hoặc lựa chọn trong giao diện Báo cáo và gọi hàm XuatTaiLieu.",
                "Nếu thầy yêu cầu thêm mức độ thông báo mới, em tìm danh sách CacMucDoThongBao trong ViewModel, thêm giá trị mới như Khẩn cấp hoặc Nhắc việc, sau đó kiểm tra ComboBox mức độ thông báo trong MainWindow.xaml.",
            ],
            "numbered": [
                "Sửa bảng, tiêu đề cột, bộ lọc: MainWindow.xaml.",
                "Sửa logic lọc/chấm công/nghỉ phép/lương: ManHinhChinhViewModel.cs.",
                "Sửa quy tắc ngày nghỉ hoặc trạng thái: QuyTacNghiepVuNhanSu.cs.",
                "Sửa model dữ liệu: MoHinh.cs.",
                "Sửa lưu SQL: KhoDuLieuNhanSu.cs.",
                "Sửa xuất file: BoXuatOffice.cs hoặc hàm tạo TaiLieuOffice trong ViewModel.",
            ],
        },
    ],
}


def build_doc(data: dict) -> Path:
    doc = setup_doc()
    add_title(doc, data["title"], data["subtitle"], data["owner"])
    add_callout(
        doc,
        "Cách dùng",
        "Tài liệu này viết theo kiểu có thể đọc gần như nguyên văn. Khi luyện nói, đọc phần lời chính, sau đó mở đúng file code ở bảng vị trí để chỉ cho thầy.",
    )

    for idx, section in enumerate(data["sections"]):
        if idx and idx in {1, 2, 4, 5}:
            doc.add_page_break()
        add_heading(doc, section["heading"], 1)
        for text in section.get("paragraphs", []):
            add_para(doc, text)
        if "bullets" in section:
            add_bullets(doc, section["bullets"])
        if "numbered" in section:
            add_numbered(doc, section["numbered"])
        if "code_table" in section:
            add_code_table(doc, section["code_table"])
        if "qas" in section:
            add_qa(doc, section["qas"])

    section = doc.sections[0]
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("Báo cáo chức năng - Phần mềm quản lý nhân sự WPF - Nhóm 3")
    footer_run.font.name = "Times New Roman"
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / data["filename"]
    doc.save(path)
    return path


def main() -> None:
    for index, data in enumerate((LONG, LUAT, TUYEN), start=1):
        build_doc(data)
        print(f"created report {index}")


if __name__ == "__main__":
    main()

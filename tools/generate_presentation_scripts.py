from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "BaoCaoTheoMauGoc" / "KichBan_BaoCao_3Nguoi"


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run(r, 18, True, (31, 78, 121))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run(r, 11, False, (71, 85, 105))
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    set_run(r, 14 if level == 1 else 12, True, (15, 118, 110) if level == 1 else (37, 99, 235))


def add_para(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run(r, 11, True)
        rest = text[len(bold_prefix) :]
        if rest:
            r = p.add_run(rest)
            set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run(r)


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run(r)


def add_table(doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run(run, 10, True, (255, 255, 255))
        cell._tc.get_or_add_tcPr().append(
            __import__("docx").oxml.parse_xml(
                r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="1F4E79"/>'
            )
        )

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    set_run(run, 10)
    doc.add_paragraph()


def add_common_footer(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("QuanLyNhanSuWpf - Kịch bản báo cáo nhóm")
    set_run(r, 9, False, (100, 116, 139))


def new_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    add_common_footer(doc)
    add_title(doc, title, subtitle)
    return doc


person1_script = [
    "Em xin trình bày phần thứ nhất là phân tích đề tài, tài liệu yêu cầu và thiết kế tổng thể của hệ thống.",
    "Đề tài nhóm em chọn là phần mềm quản lý nhân sự cho doanh nghiệp, xây dựng dưới dạng ứng dụng desktop WPF. Lý do chọn đề tài này là vì nghiệp vụ nhân sự có nhiều chức năng rõ ràng để áp dụng các bước của môn Công nghệ phần mềm, ví dụ như quản lý hồ sơ nhân viên, phòng ban, tuyển dụng, chấm công, nghỉ phép, đánh giá, bảng lương, báo cáo và tài khoản người dùng.",
    "Trước khi lập trình, nhóm em lập tài liệu nghiên cứu tính khả thi. Trong tài liệu này, nhóm đánh giá dự án theo các góc độ kỹ thuật, kinh tế, thời gian và vận hành. Về kỹ thuật, dự án khả thi vì sử dụng C# .NET, WPF và SQL Server, đều là công nghệ phổ biến và phù hợp với ứng dụng desktop. Về kinh tế, các công cụ như Visual Studio, .NET SDK, SQL Server Express hoặc LocalDB đều có thể dùng trong phạm vi học tập. Về thời gian, nhóm chia công việc theo từng giai đoạn: khảo sát, phân tích, thiết kế, lập trình, kiểm thử và bàn giao.",
    "Sau đó nhóm lập kế hoạch dự án. Kế hoạch có phân công công việc cho từng thành viên, xác định các đầu việc chính và sản phẩm bàn giao. Các đầu việc gồm nghiên cứu khả thi, đặc tả yêu cầu, thiết kế phần mềm, thiết kế cơ sở dữ liệu, lập trình, kiểm thử và chuẩn bị báo cáo.",
    "Ở tài liệu đặc tả yêu cầu, nhóm xác định các actor chính gồm Admin, Giám đốc, Trưởng phòng và Nhân viên. Mỗi actor có phạm vi sử dụng khác nhau. Admin quản trị tài khoản và toàn bộ hệ thống. Giám đốc xem dữ liệu toàn công ty. Trưởng phòng quản lý dữ liệu trong phạm vi phòng ban. Nhân viên chủ yếu xem và thao tác thông tin cá nhân như chấm công, nghỉ phép, đánh giá và phiếu lương.",
    "Về yêu cầu chức năng, hệ thống có các phân hệ chính: đăng nhập, tổng quan dashboard, quản lý hồ sơ nhân viên, quản lý phòng ban, tuyển dụng ứng viên, chấm công, nghỉ phép, đánh giá, bảng lương, báo cáo và cài đặt tài khoản. Ngoài ra còn có thông báo nội bộ, sao lưu phục hồi dữ liệu và xuất báo cáo.",
    "Về thiết kế phần mềm, nhóm thiết kế ứng dụng theo hướng WPF kết hợp ViewModel. Giao diện được viết bằng XAML, còn dữ liệu và lệnh thao tác nằm trong C#. Các lớp mô hình dữ liệu như NhanVien, PhongBan, NghiPhep, ChamCong, DanhGia, PhieuLuong và TaiKhoanHeThong được đặt trong file MoHinh.cs. ViewModel chính điều phối dữ liệu, bộ lọc, menu và thao tác người dùng.",
    "Về thiết kế cơ sở dữ liệu, hệ thống sử dụng SQL Server với database HRManagementDB. Các bảng nghiệp vụ được đặt theo tiền tố HR, ví dụ HR_Employees, HR_Departments, HR_JobPositions, HR_Attendances, HR_LeaveRequests, HR_Appraisals, HR_Payslips, HR_Users và HR_AuditLogs. Thiết kế này giúp dữ liệu nhân sự, tài khoản và nhật ký hệ thống được tách rõ theo chức năng.",
    "Tóm lại, phần em phụ trách giúp dự án bám đúng quy trình Công nghệ phần mềm: có nghiên cứu khả thi, có kế hoạch, có đặc tả yêu cầu, có thiết kế phần mềm, có thiết kế cơ sở dữ liệu và có cơ sở để nhóm triển khai chương trình.",
]


person2_script = [
    "Em xin trình bày phần thứ hai là thiết kế giao diện và xây dựng các chức năng chính của chương trình.",
    "Về công nghệ giao diện, nhóm em sử dụng WPF trên nền .NET 10. Giao diện được viết bằng XAML, còn phần xử lý dữ liệu và sự kiện được viết bằng C#. Trong file App.xaml, nhóm khai báo các tài nguyên dùng chung như màu nền, màu chữ, kiểu nút, ô nhập liệu, DataGrid và có nạp theme PresentationFramework.Fluent để giao diện theo phong cách Fluent Design.",
    "Đầu tiên là màn hình đăng nhập, nằm trong file LoginWindow.xaml. Màn hình này được chia thành hai vùng. Bên trái là nhận diện hệ thống quản trị nhân sự, bên phải là form đăng nhập gồm tên đăng nhập, mật khẩu, thông báo lỗi và các nút đăng nhập, thoát. Phần xử lý đăng nhập nằm trong LoginWindow.xaml.cs, khi người dùng nhập tài khoản thì chương trình gọi lớp xác thực để kiểm tra.",
    "Màn hình chính nằm trong file MainWindow.xaml. Em thiết kế màn hình chính theo bố cục có thanh điều hướng bên trái và vùng nội dung bên phải. Thanh điều hướng gồm các phân hệ: Tổng quan, Tuyển dụng, Hồ sơ nhân viên, Phòng ban, Chấm công, Nghỉ phép, Đánh giá, Bảng lương, Báo cáo và Cài đặt tài khoản.",
    "Ở trang Tổng quan, giao diện hiển thị các chỉ số nhanh như tổng nhân viên, quân số đang làm, thông báo, chấm công, nghỉ phép và bảng lương. Ngoài ra còn có các biểu đồ dashboard để người dùng nắm tình hình nhanh hơn. Phần biểu đồ được xử lý trong file BieuDoDashboard.cs.",
    "Ở phân hệ Nhân viên, giao diện có bộ lọc, bảng danh sách bằng DataGrid và form nhập liệu để thêm, sửa, lưu hoặc xóa hồ sơ. Các thông tin gồm mã nhân viên, họ tên, phòng ban, vị trí, ngày sinh, ngày vào làm, trạng thái, liên hệ khẩn cấp, tài khoản ngân hàng và căn cước.",
    "Ở phân hệ Phòng ban, người dùng có thể quản lý danh sách phòng ban, thêm phòng ban mới, sửa phòng ban, xóa phòng ban và gán trưởng phòng. Phần này giúp hệ thống có cơ cấu tổ chức rõ ràng để phục vụ phân quyền và thống kê.",
    "Ở phân hệ Ứng viên, hệ thống hỗ trợ nhập ứng viên, chuyển giai đoạn tuyển dụng và chuyển ứng viên thành nhân viên. Khi ứng viên được tiếp nhận, dữ liệu sẽ được cập nhật sang hồ sơ nhân viên.",
    "Ở phân hệ Chấm công, người dùng có thể ghi nhận vào ca, ra ca, điều chỉnh công và xem trạng thái công như đang trong ca, đủ công, thiếu giờ hoặc tăng ca. Phân hệ Nghỉ phép cho phép tạo đơn nghỉ, duyệt hoặc từ chối đơn nghỉ. Các thao tác này ảnh hưởng tới dashboard và tính lương.",
    "Ở phân hệ Đánh giá và Bảng lương, người có quyền có thể tạo đánh giá, chốt đánh giá, tính lương và xác nhận đã trả lương. Bảng lương được trình bày bằng DataGrid với các cột lương cơ bản, phụ cấp, khấu trừ, thực lãnh và trạng thái.",
    "Ở phân hệ Báo cáo, hệ thống hỗ trợ xuất báo cáo hồ sơ nhân viên, báo cáo chấm công, báo cáo nghỉ phép và báo cáo lương. Người dùng có thể chọn định dạng Word, Excel, PDF, PowerPoint hoặc văn bản tùy nhu cầu.",
    "Ngoài ra, hệ thống có trung tâm thông báo nội bộ. Người có quyền có thể tạo thông báo, chọn phân hệ, mức độ, nhập nội dung, đính kèm tệp và gửi. Người dùng có thể lọc thông báo chưa đọc và đánh dấu đã đọc.",
    "Tóm lại, phần giao diện không chỉ để hiển thị đẹp, mà được thiết kế theo đúng luồng nghiệp vụ: đăng nhập, xem tổng quan, quản lý nhân viên, xử lý chấm công, nghỉ phép, đánh giá, bảng lương, báo cáo và thông báo nội bộ.",
]


person3_script = [
    "Em xin trình bày phần thứ ba là bảo mật, phân quyền, kiểm thử, triển khai và bàn giao hệ thống.",
    "Đầu tiên là phần bảo mật mật khẩu. Code nằm trong file BaoMatMatKhau.cs. Nhóm em không lưu mật khẩu trực tiếp dưới dạng văn bản thường, mà sử dụng thuật toán PBKDF2-SHA256 để băm mật khẩu. Mỗi tài khoản có salt riêng được tạo bằng RandomNumberGenerator. Khi đăng nhập, hệ thống băm lại mật khẩu người dùng nhập vào rồi so sánh với hash đã lưu.",
    "Tiếp theo là luồng đăng nhập. Phần này nằm trong file KhoXacThuc.cs, hàm DangNhapAsync. Khi người dùng đăng nhập, chương trình kết nối SQL Server, kiểm tra tài khoản có tồn tại không, tài khoản có bị khóa không, sau đó xác minh mật khẩu. Nếu đăng nhập thành công thì hệ thống tạo phiên đăng nhập. Nếu thất bại thì hiển thị thông báo phù hợp.",
    "Phần tài khoản và nhật ký hệ thống được tạo trong file SoDoQuanTriSql.cs. Bảng HR_Users dùng để lưu thông tin tài khoản, vai trò, hash mật khẩu, salt, trạng thái tài khoản và lần đăng nhập gần nhất. Bảng HR_AuditLogs dùng để lưu nhật ký thao tác. Trong KhoXacThuc.cs có hàm GhiNhatKyAsync để ghi lại các sự kiện như đăng nhập thành công, đăng nhập thất bại hoặc tài khoản bị khóa.",
    "Về phân quyền, phần này nằm chủ yếu trong ManHinhChinhViewModel.cs. Hệ thống chia vai trò gồm Admin, Giám đốc, Trưởng phòng và Nhân viên. Các biến CoQuyen... quyết định vai trò nào được truy cập phân hệ nào. Ví dụ Admin có quyền cài đặt tài khoản, Admin và Giám đốc có quyền xử lý bảng lương, Trưởng phòng thao tác trong phạm vi phòng ban, còn Nhân viên chỉ xem dữ liệu liên quan tới cá nhân.",
    "Phân quyền không chỉ nằm ở menu, mà còn nằm ở từng lệnh thao tác. Ví dụ thêm và lưu nhân viên cần quyền quản lý hồ sơ, duyệt nghỉ phép cần quyền duyệt nghỉ, tính lương cần quyền xử lý bảng lương. Vì vậy người dùng không có quyền thì không thể thao tác sai phạm vi.",
    "Ngoài ra, em phụ trách phần sao lưu và phục hồi dữ liệu. Trong ManHinhChinhViewModel.cs có hàm SaoLuuDuLieu để xuất dữ liệu ra file .hrmbackup.json và hàm PhucHoiDuLieu để đọc lại file backup. Model bản sao dữ liệu nằm trong BanSaoDuLieuNhanSu.cs. Chức năng này giúp bảo vệ dữ liệu khi cần bàn giao hoặc khôi phục.",
    "Về kiểm thử, nhóm em sử dụng MSTest. Các file test nằm trong project QuanLyNhanSuWpf.Tests. BaoMatMatKhauTests kiểm tra việc băm và xác minh mật khẩu. PhanQuyenPhamViTests kiểm tra phạm vi dữ liệu theo vai trò. NghiepVuNhanSuTests kiểm tra các quy tắc nghiệp vụ như tính kỳ đánh giá, tính số ngày nghỉ phép, tính lương và chấm công. CauHinhUngDungTests kiểm tra cấu hình kết nối và mật khẩu khởi tạo. SapXepChucVuTests kiểm tra sắp xếp chức vụ.",
    "Kết quả chạy test hiện tại là 17 trên 17 test đều passed. Điều này chứng minh các phần quan trọng như bảo mật, phân quyền và nghiệp vụ nhân sự đã được kiểm tra tự động.",
    "Về triển khai, nhóm em có script package-release.ps1 trong thư mục tools. Script này chạy kiểm thử, publish ứng dụng bản Windows 64-bit self-contained và nén thành file zip để bàn giao. File chạy bàn giao nằm trong thư mục artifacts với tên QuanLyNhanSuWpf-win-x64.zip.",
    "Về quản lý phiên bản, dự án đã được commit bằng Git và đẩy lên GitHub. Trên GitHub có nhánh codex/quan-ly-nhan-su-wpf để lưu mã nguồn, tài liệu và minh chứng bonus. Nhóm cũng có workflow GitHub Actions để build, test và publish trên môi trường Windows.",
    "Tóm lại, phần em phụ trách giúp hệ thống hoàn thiện ở góc độ vận hành thực tế: đăng nhập an toàn, phân quyền theo vai trò, ghi nhật ký, sao lưu phục hồi, kiểm thử tự động, đóng gói bàn giao và lưu trữ trên GitHub.",
]


def add_script_section(doc: Document, title: str, script: list[str]):
    add_heading(doc, title, 1)
    add_para(doc, "Phần dưới đây có thể đọc nguyên văn khi báo cáo:")
    for para in script:
        add_para(doc, para)


def add_demo_steps(doc: Document, rows):
    add_heading(doc, "Vừa nói vừa chỉ", 1)
    add_table(doc, ["Khi nói tới", "Mở/chỉ vào", "Ý cần nhấn mạnh"], rows)


def add_questions(doc: Document, rows):
    add_heading(doc, "Câu trả lời nhanh nếu thầy hỏi", 1)
    add_table(doc, ["Câu hỏi", "Trả lời ngắn"], rows)


def add_long_reading(doc: Document, sections):
    add_heading(doc, "Bản đọc chi tiết mở rộng", 1)
    add_para(
        doc,
        "Phần này dùng khi cần trình bày đầy đủ hơn. Người báo cáo có thể đọc theo thứ tự, hoặc chọn các mục phù hợp nếu thời gian bị rút ngắn.",
    )
    for title, paragraphs in sections:
        add_heading(doc, title, 2)
        for para in paragraphs:
            add_para(doc, para)


def add_report_checklist(doc: Document, rows):
    add_heading(doc, "Checklist khi đứng báo cáo", 1)
    add_table(doc, ["Việc cần làm", "Nói/chỉ như thế nào"], rows)


person1_extended = [
    (
        "1. Giới thiệu bài toán và phạm vi đề tài",
        [
            "Nếu trình bày đầy đủ hơn, em bắt đầu bằng việc nói rõ bài toán mà nhóm giải quyết. Trong doanh nghiệp, bộ phận nhân sự thường phải quản lý nhiều loại dữ liệu khác nhau như hồ sơ nhân viên, phòng ban, vị trí công việc, ứng viên, chấm công, nghỉ phép, đánh giá năng lực, bảng lương và báo cáo. Nếu quản lý bằng sổ sách hoặc nhiều file rời rạc thì dữ liệu dễ bị trùng, khó tra cứu và khó tổng hợp khi cần báo cáo.",
            "Vì vậy nhóm em chọn xây dựng phần mềm quản lý nhân sự tập trung. Phần mềm không chỉ nhập và xem danh sách nhân viên, mà còn mô phỏng một quy trình nhân sự tương đối hoàn chỉnh: từ tuyển dụng ứng viên, tiếp nhận thành nhân viên, quản lý phòng ban, chấm công, nghỉ phép, đánh giá, tính lương, xuất báo cáo và quản trị tài khoản.",
            "Phạm vi của đề tài là ứng dụng desktop chạy trên Windows, phù hợp với quy mô đồ án môn học. Nhóm tập trung vào việc thể hiện đúng quy trình Công nghệ phần mềm: phân tích yêu cầu, thiết kế, lập trình, kiểm thử và bàn giao.",
        ],
    ),
    (
        "2. Nghiên cứu tính khả thi",
        [
            "Ở tài liệu nghiên cứu tính khả thi, nhóm đánh giá dự án theo bốn khía cạnh. Thứ nhất là khả thi kỹ thuật: nhóm sử dụng .NET, WPF và SQL Server. Đây là các công nghệ có tài liệu rõ ràng, dễ cài trên Windows và phù hợp với ứng dụng quản lý nội bộ.",
            "Thứ hai là khả thi kinh tế: dự án có thể thực hiện bằng Visual Studio, .NET SDK, SQL Server Express hoặc LocalDB, Git và GitHub. Các công cụ này đều miễn phí hoặc phổ biến trong môi trường học tập, nên chi phí triển khai thấp.",
            "Thứ ba là khả thi thời gian: nhóm chia công việc theo giai đoạn. Trước hết là phân tích yêu cầu, sau đó thiết kế phần mềm và CSDL, tiếp theo lập trình các phân hệ, rồi kiểm thử và đóng gói bàn giao. Cách chia này giúp nhóm kiểm soát tiến độ.",
            "Thứ tư là khả thi vận hành: chương trình có tài khoản mẫu, dữ liệu mẫu, hướng dẫn cài đặt, hướng dẫn sử dụng và bản đóng gói. Vì vậy khi bàn giao, người nhận có thể chạy thử mà không cần hiểu toàn bộ mã nguồn ngay từ đầu.",
        ],
    ),
    (
        "3. Đặc tả yêu cầu và actor",
        [
            "Trong tài liệu đặc tả yêu cầu, nhóm xác định bốn nhóm người dùng chính. Admin là người quản trị hệ thống, có quyền quản lý tài khoản và cấu hình. Giám đốc là người theo dõi toàn bộ dữ liệu nhân sự và báo cáo. Trưởng phòng quản lý dữ liệu trong phạm vi phòng ban. Nhân viên là người dùng tự phục vụ, chủ yếu xem thông tin cá nhân, chấm công, nghỉ phép, đánh giá và phiếu lương của mình.",
            "Từ các actor đó, nhóm xác định các yêu cầu chức năng. Hệ thống cần đăng nhập, phân quyền, quản lý nhân viên, quản lý phòng ban, quản lý ứng viên, chấm công, nghỉ phép, đánh giá, bảng lương, báo cáo, thông báo nội bộ và sao lưu phục hồi dữ liệu.",
            "Ngoài yêu cầu chức năng, nhóm cũng xác định yêu cầu phi chức năng. Ví dụ giao diện phải dễ sử dụng bằng tiếng Việt, dữ liệu phải được lưu vào SQL Server, mật khẩu không lưu dạng văn bản thường, thao tác quan trọng nên có thông báo hoặc nhật ký, và chương trình cần có bản đóng gói để bàn giao.",
        ],
    ),
    (
        "4. Thiết kế phần mềm và thiết kế CSDL",
        [
            "Về thiết kế phần mềm, nhóm đi theo hướng tách giao diện và xử lý dữ liệu. Giao diện nằm trong các file XAML như LoginWindow.xaml và MainWindow.xaml. Phần điều phối dữ liệu, lệnh bấm, bộ lọc và phân quyền nằm trong ManHinhChinhViewModel.cs. Các lớp mô hình dữ liệu nằm trong MoHinh.cs.",
            "Về thiết kế CSDL, nhóm sử dụng SQL Server và đặt tên database mặc định là HRManagementDB. Các bảng chính gồm HR_Employees để lưu nhân viên, HR_Departments để lưu phòng ban, HR_JobPositions để lưu vị trí công việc, HR_Attendances để lưu chấm công, HR_LeaveRequests để lưu nghỉ phép, HR_Appraisals để lưu đánh giá, HR_Payslips để lưu phiếu lương, HR_Users để lưu tài khoản và HR_AuditLogs để lưu nhật ký.",
            "Thiết kế này giúp dữ liệu được tách theo nghiệp vụ, đồng thời vẫn liên kết được với nhau. Ví dụ nhân viên thuộc phòng ban và vị trí; chấm công, nghỉ phép, đánh giá và phiếu lương đều gắn với nhân viên; tài khoản người dùng được tạo theo vai trò và có thể liên hệ với nhân sự trong hệ thống.",
        ],
    ),
    (
        "5. Kết luận phần người 1",
        [
            "Khi kết thúc phần này, em nhấn mạnh rằng nhóm không làm chương trình trước rồi mới viết tài liệu, mà đã bám theo quy trình của môn học. Các tài liệu bắt buộc đều có: nghiên cứu khả thi, kế hoạch dự án, đặc tả yêu cầu, thiết kế phần mềm, thiết kế CSDL và test case. Đây là nền tảng để người 2 trình bày phần giao diện/chức năng và người 3 trình bày phần bảo mật/kiểm thử/triển khai.",
        ],
    ),
]


person2_extended = [
    (
        "1. Công nghệ giao diện và ý tưởng thiết kế",
        [
            "Phần giao diện của hệ thống được xây dựng bằng WPF trên .NET 10. WPF cho phép thiết kế giao diện bằng XAML, hỗ trợ data binding, DataGrid, style, control template và bố cục linh hoạt. Điều này phù hợp với một phần mềm quản lý có nhiều bảng dữ liệu và form nhập liệu.",
            "Ý tưởng thiết kế giao diện là chia hệ thống thành các phân hệ rõ ràng. Người dùng sau khi đăng nhập sẽ không phải tìm chức năng rải rác, mà đi theo thanh điều hướng bên trái. Vùng bên phải thay đổi nội dung theo phân hệ đang chọn. Cách thiết kế này giống các phần mềm quản trị nội bộ, giúp dễ demo và dễ sử dụng.",
            "Trong App.xaml, nhóm khai báo màu sắc, kiểu nút, ô nhập liệu, DataGrid và nạp theme PresentationFramework.Fluent. Vì vậy có thể trình bày rằng giao diện được xây dựng bằng WPF XAML, sử dụng Fluent theme và tùy biến lại theo phong cách quản trị nhân sự.",
        ],
    ),
    (
        "2. Màn hình đăng nhập",
        [
            "Màn hình đăng nhập là điểm vào đầu tiên của hệ thống. Bên trái có khối nhận diện HR và mô tả ngắn về nền tảng quản trị nhân sự. Bên phải là form đăng nhập gồm tên đăng nhập, mật khẩu, vùng hiển thị thông báo và nút đăng nhập, thoát.",
            "Điểm cần nói khi demo là màn hình đăng nhập không chỉ để nhập tài khoản, mà còn kết nối với phần xác thực. Khi bấm đăng nhập, code trong LoginWindow.xaml.cs gọi KhoXacThuc để kiểm tra tài khoản trong SQL Server. Nếu đăng nhập sai, giao diện hiển thị thông báo lỗi. Nếu đăng nhập đúng, chương trình mở MainWindow và truyền phiên đăng nhập sang màn hình chính.",
        ],
    ),
    (
        "3. Màn hình chính và thanh điều hướng",
        [
            "Màn hình chính được thiết kế trong MainWindow.xaml. Bố cục gồm sidebar bên trái và vùng nội dung bên phải. Sidebar có logo, thông tin phiên làm việc, vai trò người dùng, phạm vi dữ liệu và danh sách phân hệ. Đây là nơi người dùng chuyển giữa Tổng quan, Tuyển dụng, Hồ sơ nhân viên, Phòng ban, Chấm công, Nghỉ phép, Đánh giá, Bảng lương, Báo cáo và Cài đặt tài khoản.",
            "Mỗi mục trên sidebar được binding với lệnh ChonMucLenh trong ViewModel. Khi người dùng chọn một mục, ViewModel đổi trạng thái mục đang chọn và giao diện hiển thị đúng vùng nội dung tương ứng. Các mục cũng có ràng buộc quyền, ví dụ Cài đặt tài khoản chỉ bật với Admin, Bảng lương tùy theo quyền xử lý bảng lương.",
        ],
    ),
    (
        "4. Dashboard tổng quan",
        [
            "Trang Tổng quan là màn hình giúp người quản lý xem nhanh tình hình hệ thống. Giao diện có các thẻ thống kê như tổng nhân sự, quân số hiện diện, số thông báo, số lượt chấm công, nghỉ phép và phiếu lương trong kỳ.",
            "Ngoài các con số, dashboard còn có biểu đồ. File BieuDoDashboard.cs tự vẽ các biểu đồ như biểu đồ tròn nhân sự, biểu đồ đường lương và biểu đồ cột ứng viên. Khi báo cáo, có thể chỉ vào phần này để chứng minh giao diện không chỉ là bảng dữ liệu mà còn có trực quan hóa.",
            "Dashboard cũng có trung tâm thông báo. Người dùng có thể xem thông báo mới, lọc thông báo chưa đọc và mở thông báo có tệp đính kèm.",
        ],
    ),
    (
        "5. Các phân hệ nghiệp vụ",
        [
            "Phân hệ Hồ sơ nhân viên dùng DataGrid để hiển thị danh sách và form để thêm/sửa nhân viên. Các trường quan trọng gồm mã nhân viên, họ tên, phòng ban, vị trí, ngày sinh, ngày vào làm, trạng thái làm việc, thông tin liên hệ, ngân hàng và căn cước.",
            "Phân hệ Phòng ban cho phép quản lý cơ cấu tổ chức. Người dùng có thể thêm phòng ban, sửa phòng ban, xóa phòng ban nếu chưa có nhân viên và gán trưởng phòng. Phần này liên quan trực tiếp đến phân quyền theo phạm vi phòng ban.",
            "Phân hệ Ứng viên mô phỏng quy trình tuyển dụng. Ứng viên có thể được thêm mới, chuyển giai đoạn và khi đạt yêu cầu thì chuyển thành nhân viên. Đây là luồng nghiệp vụ từ tuyển dụng sang hồ sơ nhân sự.",
            "Phân hệ Chấm công cho phép vào ca, ra ca, điều chỉnh công và xem trạng thái công. Phân hệ Nghỉ phép cho phép tạo đơn, duyệt hoặc từ chối. Hai phân hệ này ảnh hưởng tới quân số hiện diện và tính lương.",
            "Phân hệ Đánh giá cho phép tạo đánh giá năng lực theo kỳ, nhập điểm, nhận xét và chốt kết quả. Phân hệ Bảng lương dùng dữ liệu lương cơ bản, phụ cấp, khấu trừ, ngày công và nghỉ phép để tạo phiếu lương.",
            "Phân hệ Báo cáo cho phép xuất báo cáo hồ sơ, chấm công, nghỉ phép và lương. Khi người dùng bấm xuất, hệ thống tạo tài liệu theo định dạng được chọn như Word, Excel, PDF, PowerPoint hoặc văn bản.",
        ],
    ),
    (
        "6. Binding, ViewModel và trải nghiệm sử dụng",
        [
            "Điểm kỹ thuật quan trọng trong phần giao diện là data binding. Các bảng, form, bộ lọc, nút bấm và trạng thái hiển thị đều liên kết với ManHinhChinhViewModel.cs. Khi dữ liệu thay đổi, giao diện tự cập nhật thông qua INotifyPropertyChanged và ObservableCollection.",
            "Ví dụ khi chọn một nhân viên trong DataGrid, thuộc tính NhanVienDangChon thay đổi, các nút sửa/xóa/tính lương được làm mới trạng thái. Khi lọc phòng ban hoặc nhập từ khóa tìm kiếm, ICollectionView refresh lại danh sách hiển thị. Nhờ vậy giao diện phản hồi theo thao tác người dùng.",
            "Kết luận phần này là giao diện được thiết kế không chỉ đẹp mắt, mà bám theo luồng nghiệp vụ thật: xem tổng quan, nhập liệu, lọc dữ liệu, thao tác nghiệp vụ, nhận thông báo và xuất báo cáo.",
        ],
    ),
]


person3_extended = [
    (
        "1. Bảo mật đăng nhập và mật khẩu",
        [
            "Phần bảo mật đầu tiên là cách lưu mật khẩu. Trong file BaoMatMatKhau.cs, hệ thống có hàm BamMatKhau để tạo hash mật khẩu. Mật khẩu không lưu trực tiếp trong CSDL. Thay vào đó hệ thống tạo salt ngẫu nhiên cho từng tài khoản, rồi dùng PBKDF2-SHA256 để sinh hash.",
            "Khi người dùng đăng nhập, hệ thống dùng hàm XacMinhMatKhau. Hàm này lấy mật khẩu người dùng nhập vào, dùng lại salt và số vòng lặp đã lưu để tạo hash mới, sau đó so sánh với hash trong CSDL. Cách này an toàn hơn lưu mật khẩu thường vì ngay cả khi nhìn vào bảng tài khoản cũng không thấy mật khẩu gốc.",
            "Phần đăng nhập nằm trong KhoXacThuc.cs. Hàm DangNhapAsync thử kết nối SQL Server, đảm bảo có bảng tài khoản, lấy tài khoản theo tên đăng nhập, kiểm tra tài khoản có bị khóa không, xác minh mật khẩu, cập nhật lần đăng nhập và tạo phiên đăng nhập.",
        ],
    ),
    (
        "2. Tài khoản, vai trò và audit log",
        [
            "Bảng HR_Users được tạo trong SoDoQuanTriSql.cs. Bảng này lưu Username, FullName, RoleName, PasswordHash, PasswordSalt, PasswordIterations, IsActive, RequirePasswordChange, FailedLoginCount và LastLoginAt. Các trường này phục vụ đăng nhập, khóa mở tài khoản, reset mật khẩu và kiểm soát người dùng.",
            "Bảng HR_AuditLogs lưu nhật ký hệ thống. Khi đăng nhập thành công, đăng nhập thất bại hoặc tài khoản bị khóa, KhoXacThuc gọi GhiNhatKyAsync để ghi lại ActorUsername, ActionName, EntityName, EntityKey, Detail và MachineName. Đây là minh chứng cho yêu cầu vận hành thực tế: hệ thống có thể truy vết thao tác quan trọng.",
            "Tài khoản mặc định gồm Admin, Giám đốc, Trưởng phòng và Nhân viên. Ngoài ra TaiKhoanNhanSuSql.cs còn hỗ trợ đồng bộ tài khoản theo nhân viên, để mỗi nhân viên có tài khoản riêng theo mã nhân viên.",
        ],
    ),
    (
        "3. Phân quyền theo vai trò và phạm vi dữ liệu",
        [
            "Phân quyền nằm trong ManHinhChinhViewModel.cs. Các thuộc tính CoQuyenTuyenDung, CoQuyenQuanLyHoSoNhanVien, CoQuyenPhongBan, CoQuyenDieuChinhCong, CoQuyenDuyetNghiPhep, CoQuyenGhiNhanDanhGia, CoQuyenXuLyBangLuong, CoQuyenCaiDatTaiKhoan quyết định vai trò nào được dùng chức năng nào.",
            "Hàm LaVaiTro dùng để kiểm tra người dùng hiện tại có thuộc nhóm vai trò được phép hay không. Hàm CoQuyenTruyCap kiểm tra quyền vào từng phân hệ. Các command như ThemMoiLenh, LuuLenh, DuyetNghiPhepLenh, TinhLuongLenh, XuatBaoCaoLuongLenh cũng kiểm tra quyền trước khi cho thao tác.",
            "Không chỉ kiểm tra quyền chức năng, hệ thống còn kiểm tra phạm vi dữ liệu. Admin và Giám đốc có thể xem toàn hệ thống. Trưởng phòng xem theo phòng ban. Nhân viên chỉ xem dữ liệu liên quan tới bản thân. Phần này thể hiện qua LayNhanVienTrongPhamVi, CoTheXemNhanVien và CoTheXemTheoTen.",
        ],
    ),
    (
        "4. Sao lưu, phục hồi và báo cáo",
        [
            "Sao lưu và phục hồi nằm trong ManHinhChinhViewModel.cs. Hàm SaoLuuDuLieu tạo file .hrmbackup.json bằng cách serialize dữ liệu hiện có. Hàm PhucHoiDuLieu đọc file backup và phục hồi dữ liệu vào ứng dụng. Model dùng cho backup nằm trong BanSaoDuLieuNhanSu.cs.",
            "Xuất báo cáo nằm trong BoXuatOffice.cs và các hàm tạo tài liệu trong ManHinhChinhViewModel.cs. Hệ thống hỗ trợ các định dạng .docx, .xlsx, .pdf, .pptx và .txt. Điều này giúp phần mềm có sản phẩm đầu ra rõ ràng, không chỉ lưu dữ liệu trong màn hình.",
        ],
    ),
    (
        "5. Kiểm thử tự động",
        [
            "Project QuanLyNhanSuWpf.Tests chứa các test tự động bằng MSTest. BaoMatMatKhauTests kiểm tra cùng một mật khẩu nhưng sinh salt khác nhau và chỉ xác minh đúng với mật khẩu gốc. PhanQuyenPhamViTests kiểm tra nhân viên, trưởng phòng và giám đốc nhìn thấy dữ liệu khác nhau theo đúng vai trò.",
            "NghiepVuNhanSuTests kiểm tra các quy tắc quan trọng như tạo kỳ đánh giá theo quý, tính số ngày trong tháng, tính phiếu lương theo công và nghỉ phép, tính số năm tròn, tạo mã nhân viên tiếp theo, kiểm tra nghỉ phép giao ngày/giao khoảng và trạng thái chấm công.",
            "CauHinhUngDungTests kiểm tra cấu hình kết nối và mật khẩu khởi tạo. SapXepChucVuTests kiểm tra thứ tự cấp bậc chức vụ. Kết quả hiện tại là 17/17 test passed, có thể mở terminal và chạy dotnet test để chứng minh.",
        ],
    ),
    (
        "6. Triển khai, GitHub và bàn giao",
        [
            "Phần đóng gói nằm trong tools/package-release.ps1. Script chạy test trước, sau đó dotnet publish bản win-x64 self-contained và nén thư mục publish thành file QuanLyNhanSuWpf-win-x64.zip. Người nhận chỉ cần giải nén, chạy QuanLyNhanSuWpf.exe và đăng nhập bằng tài khoản mẫu.",
            "Dự án đã được quản lý bằng Git và đẩy lên GitHub. Nhánh sử dụng để nộp là codex/quan-ly-nhan-su-wpf. Git log có các commit thể hiện quá trình bổ sung minh chứng bonus, bật theme Fluent và tạo kịch bản báo cáo.",
            "Khi bàn giao, nhóm có thể gửi ba thứ: link GitHub chứa source và tài liệu, file zip chạy chương trình trong artifacts, và thư mục BaoCaoTheoMauGoc chứa toàn bộ tài liệu bắt buộc/bonus.",
        ],
    ),
]


def create_person1():
    doc = new_doc("Kịch bản báo cáo - Người 1", "Phân tích yêu cầu, kế hoạch dự án và thiết kế tổng thể")
    add_script_section(doc, "Lời đọc nguyên văn", person1_script)
    add_long_reading(doc, person1_extended)
    add_demo_steps(
        doc,
        [
            ("Nghiên cứu khả thi", "BaoCao_NghienCuuTinhKhaThi_QuanLyNhanSuWpf.doc", "Dự án khả thi về kỹ thuật, chi phí, thời gian và vận hành."),
            ("Kế hoạch dự án", "KeHoach_ThucHienDuAn_QuanLyNhanSuWpf.doc", "Có phân chia giai đoạn và phân công nhóm."),
            ("Đặc tả yêu cầu", "TaiLieu_DacTaYeuCau_QuanLyNhanSuWpf.doc", "Actor, chức năng, yêu cầu phi chức năng."),
            ("Thiết kế phần mềm", "TaiLieu_ThietKePhanMem_QuanLyNhanSuWpf.doc", "Kiến trúc WPF, lớp dữ liệu, luồng nghiệp vụ."),
            ("Thiết kế CSDL", "TaiLieu_ThietKeCSDL_QuanLyNhanSuWpf.doc", "Database HRManagementDB và các bảng HR_* trong SQL Server."),
            ("Mô hình trong code", "QuanLyNhanSuWpf/MoHinh.cs", "Các entity chính: nhân viên, phòng ban, chấm công, nghỉ phép, lương."),
        ],
    )
    add_questions(
        doc,
        [
            ("Vì sao chọn đề tài này?", "Vì nghiệp vụ nhân sự rõ ràng, có đủ yêu cầu phân tích, thiết kế, lập trình và kiểm thử."),
            ("Dự án khả thi ở đâu?", "Trong tài liệu nghiên cứu tính khả thi, đánh giá kỹ thuật, kinh tế, thời gian và vận hành."),
            ("Thiết kế CSDL thể hiện ở đâu?", "Trong tài liệu thiết kế CSDL và trong code tạo bảng SQL Server."),
        ],
    )
    add_report_checklist(
        doc,
        [
            ("Mở thư mục tài liệu", "Chỉ vào BaoCaoTheoMauGoc để chứng minh đủ tài liệu bắt buộc."),
            ("Nói về CSDL", "Mở TaiLieu_ThietKeCSDL hoặc SoDoQuanTriSql.cs nếu thầy hỏi minh chứng."),
            ("Nối sang người 2", "Kết luận rằng tài liệu yêu cầu là cơ sở để xây dựng giao diện và chức năng."),
        ],
    )
    return doc


def create_person2():
    doc = new_doc("Kịch bản báo cáo - Người 2", "Thiết kế giao diện WPF và xây dựng chức năng nghiệp vụ")
    add_script_section(doc, "Lời đọc nguyên văn", person2_script)
    add_long_reading(doc, person2_extended)
    add_demo_steps(
        doc,
        [
            ("Công nghệ giao diện", "QuanLyNhanSuWpf/App.xaml", "WPF XAML, style chung, theme PresentationFramework.Fluent."),
            ("Màn hình đăng nhập", "QuanLyNhanSuWpf/LoginWindow.xaml", "Bố cục 2 vùng, form tài khoản/mật khẩu, thông báo lỗi."),
            ("Màn hình chính", "QuanLyNhanSuWpf/MainWindow.xaml", "Sidebar phân hệ bên trái, vùng nội dung bên phải."),
            ("Binding và lệnh bấm", "QuanLyNhanSuWpf/ManHinhChinhViewModel.cs", "Dữ liệu, bộ lọc, command và cập nhật giao diện."),
            ("Dashboard", "QuanLyNhanSuWpf/BieuDoDashboard.cs", "Biểu đồ tròn, đường, cột cho tổng quan nhân sự."),
            ("Thông báo", "MainWindow.xaml và ManHinhChinhViewModel.cs", "Tạo/lọc/đánh dấu đã đọc/đính kèm tệp."),
        ],
    )
    add_questions(
        doc,
        [
            ("Em dùng công nghệ gì?", "WPF trên .NET 10, giao diện XAML, xử lý bằng C#, binding qua ViewModel."),
            ("Em thiết kế giao diện thế nào?", "Thiết kế theo phân hệ nghiệp vụ, sidebar điều hướng, DataGrid danh sách, form chi tiết và dashboard."),
            ("Có dùng Fluent UI không?", "Có nạp theme PresentationFramework.Fluent trong App.xaml và tùy biến style theo Fluent Design."),
            ("UI/UX Figma ở đâu?", "Wireframe import Figma nằm trong thư mục uiux, gồm màn hình đăng nhập, tổng quan và quản lý nhân viên."),
        ],
    )
    add_report_checklist(
        doc,
        [
            ("Demo đăng nhập", "Mở LoginWindow hoặc chạy app, nói màn hình đăng nhập kết nối xác thực."),
            ("Demo sidebar", "Chỉ menu phân hệ trong MainWindow.xaml hoặc trên app."),
            ("Demo chức năng", "Chọn Nhân viên, Chấm công, Nghỉ phép, Bảng lương, Báo cáo để thầy thấy luồng đầy đủ."),
            ("Nói kỹ thuật", "Chỉ App.xaml, ManHinhChinhViewModel.cs, BieuDoDashboard.cs."),
        ],
    )
    return doc


def create_person3():
    doc = new_doc("Kịch bản báo cáo - Người 3", "Bảo mật, phân quyền, kiểm thử, triển khai và bàn giao")
    add_script_section(doc, "Lời đọc nguyên văn", person3_script)
    add_long_reading(doc, person3_extended)
    add_demo_steps(
        doc,
        [
            ("Bảo mật mật khẩu", "QuanLyNhanSuWpf/BaoMatMatKhau.cs", "PBKDF2-SHA256, salt riêng, xác minh hash."),
            ("Đăng nhập", "QuanLyNhanSuWpf/KhoXacThuc.cs", "DangNhapAsync kiểm tra tài khoản, mật khẩu, trạng thái."),
            ("Audit log", "QuanLyNhanSuWpf/SoDoQuanTriSql.cs và KhoXacThuc.cs", "Bảng HR_AuditLogs và hàm GhiNhatKyAsync."),
            ("Phân quyền", "QuanLyNhanSuWpf/ManHinhChinhViewModel.cs", "CoQuyen..., LaVaiTro, CoQuyenTruyCap, CoTheXemNhanVien."),
            ("Backup/restore", "ManHinhChinhViewModel.cs và BanSaoDuLieuNhanSu.cs", "Sao lưu .hrmbackup.json và phục hồi dữ liệu."),
            ("Test", "QuanLyNhanSuWpf.Tests", "Bảo mật, phân quyền, nghiệp vụ, cấu hình, chức vụ."),
            ("Đóng gói", "tools/package-release.ps1", "dotnet publish win-x64 và Compress-Archive tạo zip."),
            ("GitHub", "git log và link nhánh codex/quan-ly-nhan-su-wpf", "Có commit và đã đẩy lên GitHub."),
        ],
    )
    add_questions(
        doc,
        [
            ("Mật khẩu có an toàn không?", "Không lưu plain text, dùng PBKDF2-SHA256 và salt riêng."),
            ("Phân quyền nằm ở đâu?", "Trong ManHinhChinhViewModel.cs qua các biến CoQuyen và hàm kiểm tra vai trò."),
            ("Test được những gì?", "Bảo mật mật khẩu, phân quyền, tính lương, nghỉ phép, chấm công, cấu hình và chức vụ."),
            ("Bàn giao chạy thế nào?", "Giải nén QuanLyNhanSuWpf-win-x64.zip, chạy QuanLyNhanSuWpf.exe, đăng nhập admin/Admin@2026!."),
        ],
    )
    add_report_checklist(
        doc,
        [
            ("Mở bảo mật", "Chỉ BaoMatMatKhau.cs: PBKDF2-SHA256, salt, xác minh mật khẩu."),
            ("Mở phân quyền", "Chỉ ManHinhChinhViewModel.cs: CoQuyen, LaVaiTro, CoQuyenTruyCap."),
            ("Chạy test", "Nói lệnh dotnet test và kết quả 17/17 passed."),
            ("Bàn giao", "Chỉ package-release.ps1 và file zip artifacts/QuanLyNhanSuWpf-win-x64.zip."),
        ],
    )
    return doc


def create_combined():
    doc = new_doc("Kịch bản báo cáo phần mềm cho 3 người", "Đọc nguyên văn khi bảo vệ đồ án QuanLyNhanSuWpf")
    add_heading(doc, "Phân chia phần báo cáo", 1)
    add_table(
        doc,
        ["Người", "Nội dung chính", "Minh chứng"],
        [
            ("Người 1", "Phân tích yêu cầu, kế hoạch, thiết kế phần mềm và CSDL", "Các tài liệu trong BaoCaoTheoMauGoc"),
            ("Người 2", "Thiết kế giao diện WPF và chức năng nghiệp vụ", "LoginWindow.xaml, MainWindow.xaml, App.xaml, ViewModel"),
            ("Người 3", "Bảo mật, phân quyền, kiểm thử, triển khai và GitHub", "BaoMatMatKhau.cs, KhoXacThuc.cs, Tests, package-release.ps1"),
        ],
    )
    add_script_section(doc, "Người 1 - Lời đọc nguyên văn", person1_script)
    add_long_reading(doc, person1_extended)
    add_script_section(doc, "Người 2 - Lời đọc nguyên văn", person2_script)
    add_long_reading(doc, person2_extended)
    add_script_section(doc, "Người 3 - Lời đọc nguyên văn", person3_script)
    add_long_reading(doc, person3_extended)
    add_heading(doc, "Câu chốt chung cho nhóm", 1)
    add_para(
        doc,
        "Tóm lại, nhóm em đã hoàn thành các yêu cầu bắt buộc gồm nghiên cứu khả thi, kế hoạch dự án, đặc tả yêu cầu, thiết kế, test case và chương trình. Phần bonus có Git/GitHub, thiết kế CSDL, hướng dẫn sử dụng, hướng dẫn cài đặt, giao diện theo Fluent Design, UI/UX Figma-ready và notification nội bộ.",
    )
    return doc


def save_doc(doc: Document, name: str):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / name
    doc.save(path)
    return path


def main():
    outputs = [
        save_doc(create_person1(), "Nguoi_1_PhanTich_YeuCau_ThietKe.docx"),
        save_doc(create_person2(), "Nguoi_2_GiaoDien_ChucNang.docx"),
        save_doc(create_person3(), "Nguoi_3_BaoMat_KiemThu_TrienKhai.docx"),
        save_doc(create_combined(), "KichBan_BaoCao_TongHop_3Nguoi.docx"),
    ]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
